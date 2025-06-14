#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys
import os

print(f"Python路径: {sys.executable}")
print(f"Python版本: {sys.version}")

try:
    import docx
    print(f"python-docx版本: {docx.__version__}")
    print("✅ python-docx库已正确安装")
except ImportError as e:
    print(f"❌ python-docx库导入失败: {e}")
    sys.exit(1)

# 测试创建新文档
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph("测试文档")
    print("✅ 可以创建新的docx文档")
except Exception as e:
    print(f"❌ 创建文档失败: {e}")

# 测试读取系统中的docx文件
test_files = [
    "/Users/acheng/Downloads/law-exam-assistant/uploads/刑法_1749471200053.docx",
    "/Users/acheng/Downloads/law-exam-assistant/民法客观题.docx"
]

for file_path in test_files:
    if os.path.exists(file_path):
        print(f"\n测试文件: {file_path}")
        try:
            doc = Document(file_path)
            print(f"✅ 成功打开，包含 {len(doc.paragraphs)} 个段落")
        except Exception as e:
            print(f"❌ 打开失败: {e}")
            print(f"错误类型: {type(e).__name__}")