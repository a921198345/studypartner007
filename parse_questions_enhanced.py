#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
增强版的题目解析脚本 - 处理2013-2016年的特殊格式
主要改进：
1. 识别【答案】标记
2. 处理选项格式中的空格问题
3. 更智能的选项提取
4. 更准确的答案提取
"""

import re
import json
import sys
import os
import argparse
from datetime import datetime

def clean_text(text):
    """智能清理文本中的异常空白"""
    if not text:
        return text
    
    # 1. 替换全角空格为半角空格
    text = text.replace('　', ' ')
    
    # 2. 替换多个连续空格为单个空格
    text = re.sub(r' +', ' ', text)
    
    # 3. 清理行内的异常空白（保留正常的段落结构）
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # 去除行首尾空白
        line = line.strip()
        
        # 处理选项中的异常空白，如 "A .选项" -> "A.选项"
        line = re.sub(r'([A-D])\s+\.', r'\1.', line)
        
        # 处理标点符号前的异常空白
        line = re.sub(r'\s+([，。！？；：、）】》"''])', r'\1', line)
        
        # 处理标点符号后的异常空白（但保留句号后的正常空格）
        line = re.sub(r'([（【《"''])\s+', r'\1', line)
        
        cleaned_lines.append(line)
    
    # 4. 处理多余的空行（连续3个或以上换行符）
    text = '\n'.join(cleaned_lines)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text

def smart_fix_options(options_dict):
    """智能修复选项中的问题"""
    fixed_options = {}
    
    for key, value in options_dict.items():
        # 清理选项文本
        value = clean_text(value)
        
        # 如果选项为空或只有空白，标记为异常
        if not value or value.isspace():
            value = "[选项内容缺失]"
        
        fixed_options[key] = value
    
    return fixed_options

def validate_question_format(text):
    """智能验证题目格式 - 更宽松的验证规则"""
    issues = []
    suggestions = []  # 新增建议列表
    
    # 先清理文本
    cleaned_text = clean_text(text)
    
    # 检查题号格式 - 支持多种格式
    question_id_patterns = [
        r'【\d{4,}】',      # 支持4位或更多数字
        r'\[\d+\]',        # 方括号格式
        r'（\d+）',         # 圆括号格式
        r'^\d+[\.\.、：:]'   # 数字开头格式
    ]
    
    has_question_id = any(re.search(pattern, cleaned_text) for pattern in question_id_patterns)
    if not has_question_id:
        issues.append("未找到有效的题号格式")
    
    # 检查选项数量 - 更智能的检测
    # 分段检查，避免跨页问题
    option_count = 0
    found_options = set()
    
    # 检查各种选项格式
    option_patterns = [
        r'([A-D])\s*[\.．、：:]',  # 修改：允许选项字母后有空格
        r'^([A-D])\s+[^A-Z]',        # A 选项内容
        r'\n([A-D])[．、]\s*[^A-Z]'  # 处理特殊格式
    ]
    
    for pattern in option_patterns:
        matches = re.findall(pattern, cleaned_text, re.MULTILINE)
        found_options.update(matches)
    
    option_count = len(found_options)
    
    if option_count < 4:
        missing = set(['A', 'B', 'C', 'D']) - found_options
        if option_count >= 2:  # 如果找到了至少2个选项
            suggestions.append(f"可能缺少选项: {', '.join(sorted(missing))} (可能在其他页面)")
        else:
            issues.append(f"选项不完整：找到{option_count}个选项，缺少: {', '.join(sorted(missing))}")
    
    # 检查解析标记 - 支持多种格式
    analysis_patterns = ['【解析】', '解析：', '解析:', '分析：', '答案解释：']
    has_analysis = any(pattern in cleaned_text for pattern in analysis_patterns)
    
    if not has_analysis:
        # 检查是否有解析性内容
        if any(keyword in cleaned_text for keyword in ['因为', '根据', '由于', '所以', '因此', '故']):
            suggestions.append("未找到标准解析标记，但存在解析性内容")
        else:
            issues.append("缺少解析内容")
    
    # 检查答案标识 - 扩展答案模式，特别是【答案】格式
    answer_patterns = [
        "【答案】", "答案", "正确答案", "故选", "应选", "本题选",
        "因此选", "所以选", "正确的是", "A正确", "B正确", 
        "C正确", "D正确", "选A", "选B", "选C", "选D",
        "A项是正确", "B项是正确", "C项是正确", "D项是正确"
    ]
    
    has_answer = any(pattern in cleaned_text for pattern in answer_patterns)
    
    # 检查是否在末尾有单独的字母
    if not has_answer:
        if re.search(r'[。，]\s*[A-D]\s*[。]?\s*$', cleaned_text):
            has_answer = True
    
    if not has_answer:
        # 检查是否引用其他题目
        if re.search(r'同\s*【?\d+】?\s*题|参见\s*【?\d+】?\s*题', cleaned_text):
            suggestions.append("答案可能引用了其他题目")
        else:
            suggestions.append("未找到明确的答案标识（可能在其他位置）")
    
    # 只返回严重问题作为issues，建议不算错误
    return issues
   
def parse_question(text):
    """解析单个题目文本，严格要求4个选项，智能处理空白和特殊格式"""
    # 先清理文本
    text = clean_text(text)
    
    # 智能匹配题号 - 支持多种格式
    question_id = None
    question_id_patterns = [
        (r'【(\d{4,8})】', 'bracket'),      # 标准格式
        (r'\[(\d+)\]', 'square'),          # 方括号
        (r'（(\d+)）', 'paren'),            # 圆括号
        (r'^(\d+)[\.\.、：:]', 'number')     # 数字开头
    ]
    
    for pattern, format_type in question_id_patterns:
        match = re.search(pattern, text, re.MULTILINE)
        if match:
            question_id = match.group(1)
            # 智能标准化题号
            if len(question_id) == 9:  # 如果是9位，直接使用
                pass
            elif len(question_id) < 8:
                if len(question_id) >= 4 and question_id[:4].isdigit() and int(question_id[:4]) >= 2000:
                    question_id = question_id.ljust(8, '0')
                else:
                    question_id = question_id.zfill(8)
            elif len(question_id) > 9:
                # 超过9位，可能是特殊格式，保留原样
                pass
            break
    
    if not question_id:
        return None
    
    # 提取年份
    year = int(question_id[:4])
    
    # 分离题干、选项和解析
    parts = text.split('【解析】')
    if len(parts) != 2:
        return None
    
    question_part = parts[0]
    analysis = parts[1].strip()
    
    # 清理解析文本
    analysis = clean_text(analysis)
    
    # 提取题干（题号之后到第一个选项之前的文本）
    # 改进的正则表达式，处理选项字母后有空格的情况
    question_text_match = re.search(r'】(.*?)(?=[A-D]\s*[\.．、：:])', question_part, re.DOTALL)
    if not question_text_match:
        return None
    
    question_text = question_text_match.group(1).strip()
    question_text = clean_text(question_text)
    
    # 智能提取选项 - 处理跨页和多种格式
    options_dict = {}
    
    # 改进的选项提取方法 - 特别处理2016年的格式
    # 首先尝试提取所有可能的选项内容
    if year >= 2016:
        # 2016年后的题目经常在解析部分才有选项内容
        # 合并题干和解析来查找选项
        full_content_for_options = text
    else:
        full_content_for_options = question_part
    
    # 方法1: 使用更灵活的正则表达式
    option_patterns = [
        # 处理 "A  ．内容" 或 "A．内容" 格式
        r'([A-D])\s*[\.．、：:]\s*([^\n]*?)(?=\n[A-D]\s*[\.．、：:]|【解析】|【答案】|\n\n|$)',
        # 处理 "A 内容" 格式（适用于解析部分的选项说明）
        r'\n([A-D])\s+项?[是为]?[正确错误][的，。]([^\n]*?)(?=\n[A-D]\s+项?[是为]?[正确错误]|【答案】|$)',
        # 处理简单的 "A 内容" 格式
        r'(?:^|\n)([A-D])\s+([^\n]*?)(?=\n[A-D]\s+|【解析】|【答案】|\n\n|$)',
        # 处理同一行的选项 "A.选项1 B.选项2 C.选项3 D.选项4"
        r'([A-D])\s*[\.．、：:]\s*([^A-Z]*?)(?=[A-D]\s*[\.．、：:]|【解析】|【答案】|$)',
    ]
    
    for pattern in option_patterns:
        matches = re.findall(pattern, full_content_for_options, re.DOTALL | re.MULTILINE)
        for opt, content in matches:
            content = clean_text(content)
            # 过滤掉选项说明（如"项是正确的"）
            if content and opt not in options_dict and not re.match(r'^项?[是为]?[正确错误]', content):
                options_dict[opt] = content
    
    # 方法2: 如果选项不完整，逐行查找
    if len(options_dict) < 4:
        lines = question_part.split('\n')
        current_option = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            # 检查是否是新选项的开始
            option_match = re.match(r'^([A-D])\s*[\.．、：:]\s*(.*)$', line)
            if option_match:
                # 保存前一个选项
                if current_option and current_option not in options_dict:
                    content = ' '.join(current_content).strip()
                    if content:
                        options_dict[current_option] = clean_text(content)
                
                # 开始新选项
                current_option = option_match.group(1)
                current_content = [option_match.group(2)]
            elif current_option and line and not re.match(r'^[A-D]\s*[\.．、：:]', line):
                # 这是当前选项的延续内容
                current_content.append(line)
        
        # 保存最后一个选项
        if current_option and current_option not in options_dict:
            content = ' '.join(current_content).strip()
            if content:
                options_dict[current_option] = clean_text(content)
    
    # 智能修复选项
    options_dict = smart_fix_options(options_dict)
    
    # 再次确认有4个选项
    if len(options_dict) != 4 or set(options_dict.keys()) != {'A', 'B', 'C', 'D'}:
        return None
               
    # 初始化 correct_answer 变量
    correct_answer = ""
                   
    # 智能查找答案 - 支持更多格式，特别是【答案】格式
    answer_patterns = [
        # 【答案】格式 - 优先级最高
        r'【答案】\s*[：:：]?\s*([A-D]+)[。\s]*',
        r'【答案】([A-D]+)[。\s]*',
        
        # 显式答案标记
        r'答案\s*[：:：]\s*([A-D]+)',
        r'正确答案\s*[是为：:：]\s*([A-D]+)',
        
        # 常见答案表述
        r'故\s*选\s*([A-D]+)',
        r'应\s*选\s*([A-D]+)',
        r'本题\s*选\s*([A-D]+)',
        r'因此\s*选\s*([A-D]+)',
        r'所以\s*选\s*([A-D]+)',
        r'综上.*?选\s*([A-D]+)',
        
        # 选项正确性描述
        r'([A-D])\s*项?\s*(?:正确|对|符合|满足)',
        r'正确的?是\s*([A-D])',
        r'([A-D])\s*(?:正确|对)(?:[。，]|$)',
        
        # 末尾答案
        r'[。，]\s*([A-D]+)\s*[。]?\s*$',
        
        # 处理题目引用
        r'同\s*【?\d+】?\s*题[，。]?\s*(?:答案是)?\s*([A-D])',
        r'参见\s*【?\d+】?\s*题[，。]?\s*(?:答案是)?\s*([A-D])',
        
        # 其他变体
        r'选择\s*([A-D])',
        r'应选择\s*([A-D])',
        r'正确选项[是为：:]\s*([A-D])',
        
        # 特殊格式：在解析末尾单独的字母
        r'^\s*([A-D]+)\s*$',  # 单独一行只有一个或多个字母
    ]
    
    # 优先在解析部分查找
    if analysis:
        for pattern in answer_patterns:
            answer_match = re.search(pattern, analysis, re.IGNORECASE | re.MULTILINE)
            if answer_match:
                correct_answer = answer_match.group(1).strip().upper()
                break
    
    # 如果解析中没找到，在全文查找
    if not correct_answer:
        for pattern in answer_patterns:
            answer_match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if answer_match:
                correct_answer = answer_match.group(1).strip().upper()
                break

    # 如果使用上面的方法仍然没找到答案，尝试最后一段文本
    if not correct_answer:
        # 获取解析的最后一段
        last_lines = analysis.strip().split('\n')[-3:]  # 获取最后3行
        for line in reversed(last_lines):
            line = line.strip()
            # 查找单个或多个字母
            if re.match(r'^[A-D]+$', line):
                correct_answer = line
                break
            # 或者在短行中查找字母
            elif len(line) < 10:
                match = re.search(r'([A-D]+)', line)
                if match:
                    correct_answer = match.group(1)
                    break

    # 根据答案长度判断题型
    if correct_answer and len(correct_answer) > 1:
        question_type = 2  # 多选
    else:
        question_type = 1  # 单选
    
    return {
        'question_id': question_id,
        'year': year,
        'question_text': question_text,
        'options': json.dumps(options_dict, ensure_ascii=False),
        'correct_answer': correct_answer,
        'analysis': analysis,
        'question_type': question_type
    }

def extract_questions_from_doc(file_path, subject, validate=True):
    """从Word文档中提取所有题目，严格验证格式，智能处理空白"""
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        print(f"[调试] 文件大小: {file_size} bytes", file=sys.stderr)
        
        # 导入python-docx
        try:
            from docx import Document
        except ImportError:
            raise Exception("未安装python-docx库，请执行: pip install python-docx")
            
        # 尝试打开文档
        print(f"[调试] 尝试打开文档: {file_path}", file=sys.stderr)
        doc = Document(file_path)
        print(f"[调试] 文档打开成功", file=sys.stderr)
        
    except Exception as e:
        # 如果是"Package not found"错误，说明文件格式有问题
        if "Package not found" in str(e):
            raise Exception(f"文档格式错误，请确保是标准的.docx文件: {e}")
        else:
            raise Exception(f"无法打开文档: {e}")
    
    full_text = ""
    
    # 将所有段落合并为一个文本
    para_count = len(doc.paragraphs)
    print(f"[调试] 文档包含 {para_count} 个段落", file=sys.stderr)
    
    for i, para in enumerate(doc.paragraphs):
        full_text += para.text + "\n"
        if i % 100 == 0:
            print(f"[调试] 已处理 {i}/{para_count} 个段落", file=sys.stderr)
    
    # 先清理整个文档的文本
    full_text = clean_text(full_text)
    
    # 按题号分割文本
    question_pattern = r'【\d+】.*?(?=【\d+】|$)'
    questions_text = re.findall(question_pattern, full_text, re.DOTALL)
    
    print(f"[调试] 找到 {len(questions_text)} 个潜在题目", file=sys.stderr)
    
    parsed_questions = []
    format_issues = {}
    auto_fixed_count = 0  # 记录自动修复的数量
    
    for i, q_text in enumerate(questions_text):
        if i % 50 == 0:
            print(f"[调试] 正在解析题目 {i}/{len(questions_text)}", file=sys.stderr)
            
        # 验证格式
        issues = validate_question_format(q_text)
        
        # 获取题号
        question_id_match = re.search(r'【(\d+)】', q_text)
        question_id = question_id_match.group(1) if question_id_match else "未知题号"
        
        if issues:
            # 记录格式问题
            format_issues[question_id] = {
                "text": q_text[:200] + "...",
                "issues": issues
            }
            continue  # 跳过格式有问题的题目
        
        # 解析题目
        parsed = parse_question(q_text)
        if parsed:
            parsed['subject'] = subject
            
            # 检查是否有自动修复的内容
            if "[选项内容缺失]" in parsed['options']:
                auto_fixed_count += 1
                parsed['auto_fixed'] = True
            
            parsed_questions.append(parsed)
        else:
            # 解析失败也记录为格式问题
            format_issues[question_id] = {
                "text": q_text[:200] + "...",
                "issues": ["解析失败，可能是格式不规范"]
            }
       
    return parsed_questions, format_issues, auto_fixed_count

def apply_manual_answers(questions):
    """应用手动指定的答案规则"""
    manual_answers = {
        "20210201": "B",
    }
    
    for q in questions:
        if q['question_id'] in manual_answers:
            q['correct_answer'] = manual_answers[q['question_id']]
    
    return questions
   
def main():
    parser = argparse.ArgumentParser(description='增强版法考题目解析(处理2013-2016特殊格式)')
    parser.add_argument('file_path', help='Word文件路径')
    parser.add_argument('subject', help='学科名称')
    parser.add_argument('--output-json', action='store_true', help='以JSON格式输出结果')
    args = parser.parse_args()
    
    file_path = args.file_path
    subject = args.subject
    output_json = args.output_json
    
    if not os.path.exists(file_path):
        error_msg = f"错误: 文件 '{file_path}' 不存在"
        if output_json:
            print(json.dumps({"error": error_msg}, ensure_ascii=False))
        else:
            print(error_msg)
        return
    
    try:
        # 始终验证格式
        questions, format_issues, auto_fixed_count = extract_questions_from_doc(file_path, subject, validate=True)
        questions = apply_manual_answers(questions)
        
        # 准备输出信息
        result = {
            "success": True,
            "message": f"成功解析 {len(questions)} 个题目",
            "total_questions": len(questions) + len(format_issues),
            "parsed_questions": len(questions),
            "format_issues": format_issues,
            "questions": questions,
            "auto_fixed_count": auto_fixed_count
        }
        
        # 添加统计信息
        if auto_fixed_count > 0:
            result["message"] += f"（自动修复了{auto_fixed_count}个异常空白问题）"
        
        if format_issues:
            result["message"] += f"，{len(format_issues)} 个题目存在格式错误"
            
            # 统计错误类型
            error_types = {}
            for q_id, issue_data in format_issues.items():
                for issue in issue_data['issues']:
                    if issue not in error_types:
                        error_types[issue] = []
                    error_types[issue].append(q_id)
            
            result["error_summary"] = error_types
        
        # 输出结果
        if output_json:
            print(json.dumps(result, ensure_ascii=False))
        else:
            print(f"共解析出 {len(questions)} 个题目")
            if auto_fixed_count > 0:
                print(f"自动修复了 {auto_fixed_count} 个异常空白问题")
            if format_issues:
                print(f"\n发现 {len(format_issues)} 个题目存在格式问题：")
                for q_id, issue in list(format_issues.items())[:10]:
                    print(f"\n题目 {q_id}:")
                    for err in issue['issues']:
                        print(f"  - {err}")
                if len(format_issues) > 10:
                    print(f"\n... 还有 {len(format_issues) - 10} 个题目存在问题")
    
    except Exception as e:
        error_msg = f"程序出错: {e}"
        if output_json:
            print(json.dumps({"error": error_msg}, ensure_ascii=False))
        else:
            print(error_msg)
            import traceback
            traceback.print_exc()
   
if __name__ == "__main__":
    main()