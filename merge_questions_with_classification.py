#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合并完整题目内容与科目分类
将提取的真题内容与之前的8科目分类结果合并
"""

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def load_questions():
    """加载提取的完整题目"""
    with open('/Users/acheng/Downloads/law-exam-assistant/2024_law_exam_enhanced_questions.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data['questions']

def load_answers():
    """加载答案解析"""
    answers = {}
    try:
        # 尝试从答案解析文档中提取答案
        doc = Document('/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx')
        current_num = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith('正确答案：') or '正确答案：' in text:
                # 提取题号和答案
                import re
                num_match = re.search(r'(\d+)\.\s*正确答案[：:]\s*([ABCD]+)', text)
                if num_match:
                    num = int(num_match.group(1))
                    answer = num_match.group(2)
                    answers[num] = answer
    except Exception as e:
        print(f"加载答案时出错: {e}")
    
    return answers

def get_subject_classification():
    """获取科目分类（使用之前的分类结果）"""
    # 这里使用之前生成的分类结果
    return {
        '民法': list(range(1, 29)),  # 1-28
        '商经知': list(range(29, 51)),  # 29-50
        '理论法': list(range(51, 64)),  # 51-63
        '民事诉讼法': list(range(64, 76)),  # 64-75
        '刑法': list(range(76, 88)),  # 76-87
        '行政法': list(range(88, 92)),  # 88-91
        '刑事诉讼法': list(range(92, 95)),  # 92-94
        '三国法': list(range(95, 97))  # 95-96
    }

def create_enhanced_subject_document(subject, question_numbers, all_questions, answers, output_dir):
    """创建包含完整题目的科目文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(f'2024年法考客观题 - {subject}（完整真题版）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph()
    stats.add_run(f'共 {len(question_numbers)} 道题').bold = True
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(20)
    
    # 分隔线
    doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 按题号排序添加题目
    valid_questions = 0
    for num in sorted(question_numbers):
        # 查找对应的题目
        question = None
        for q in all_questions:
            if q['number'] == num:
                question = q
                break
        
        if question:
            valid_questions += 1
            
            # 添加题目编号
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f'第{num}题')
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_para.paragraph_format.space_after = Pt(6)
            
            # 添加题干
            stem_para = doc.add_paragraph()
            stem_para.add_run(question['stem'])
            stem_para.paragraph_format.space_after = Pt(8)
            
            # 添加选项
            for option_key in ['A', 'B', 'C', 'D']:
                if option_key in question['options']:
                    option_para = doc.add_paragraph()
                    option_para.add_run(f"{option_key}. {question['options'][option_key]}")
                    option_para.paragraph_format.left_indent = Pt(20)
                    option_para.paragraph_format.space_after = Pt(4)
            
            # 添加正确答案（如果有）
            if num in answers:
                answer_para = doc.add_paragraph()
                answer_run = answer_para.add_run(f"正确答案：{answers[num]}")
                answer_run.bold = True
                answer_run.font.color.rgb = RGBColor(0, 128, 0)
                answer_para.paragraph_format.space_before = Pt(6)
            
            # 添加分隔线
            separator = doc.add_paragraph('-' * 50)
            separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            separator.paragraph_format.space_before = Pt(12)
            separator.paragraph_format.space_after = Pt(12)
    
    # 保存文档
    output_path = os.path.join(output_dir, f'{subject}（完整题目版）.docx')
    doc.save(output_path)
    print(f"已生成 {subject} 文档，包含 {valid_questions}/{len(question_numbers)} 道题")

def main():
    """主函数"""
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx_complete'
    
    print("2024年法考题完整内容合并系统")
    print("=" * 60)
    
    # 确保输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 加载数据
    print("\n加载题目数据...")
    questions = load_questions()
    print(f"加载了 {len(questions)} 道题目")
    
    print("\n加载答案数据...")
    answers = load_answers()
    print(f"加载了 {len(answers)} 道题目的答案")
    
    # 获取科目分类
    classification = get_subject_classification()
    
    # 生成各科目文档
    print("\n生成科目文档...")
    total_classified = 0
    for subject, question_numbers in classification.items():
        if question_numbers:
            create_enhanced_subject_document(subject, question_numbers, questions, answers, output_dir)
            total_classified += len(question_numbers)
    
    # 生成统计报告
    print("\n" + "=" * 60)
    print("完成统计：")
    print(f"- 总题目数：{len(questions)}")
    print(f"- 已分类题目：{total_classified}")
    print(f"- 有答案的题目：{len(answers)}")
    print(f"- 输出目录：{output_dir}")
    
    # 生成报告文件
    report = {
        'total_questions': len(questions),
        'classified_questions': total_classified,
        'questions_with_answers': len(answers),
        'subjects': {subject: len(nums) for subject, nums in classification.items()},
        'output_directory': output_dir
    }
    
    with open(os.path.join(output_dir, 'merge_report.json'), 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    print(f"\n✓ 合并完成！")
    print(f"完整题目版文档已保存到: {output_dir}")

if __name__ == "__main__":
    main()