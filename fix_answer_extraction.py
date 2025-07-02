#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
修正答案提取的脚本
基于对2023年卷一答案.docx的分析结果，提供正确的答案提取方法
"""

import re
from docx import Document

def extract_answers_from_docx(docx_path):
    """
    从答案文档中提取答案
    
    根据分析结果，答案格式为：
    1.正确答案： A
    2.正确答案： D
    ...
    """
    print(f"正在提取答案文档: {docx_path}")
    
    try:
        doc = Document(docx_path)
        answers = {}
        
        # 精确的答案匹配模式
        # 基于分析结果，答案格式是 "数字.正确答案： 字母"
        answer_pattern = r'^(\d+)\.正确答案：\s*([ABCD]+)$'
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            match = re.match(answer_pattern, text)
            if match:
                question_num = int(match.group(1))
                answer = match.group(2)
                answers[question_num] = answer
                print(f"题目 {question_num}: {answer}")
        
        print(f"总共提取到 {len(answers)} 个答案")
        print(f"题号范围: {min(answers.keys())} - {max(answers.keys())}")
        
        # 检查是否有多选题
        single_choice = sum(1 for ans in answers.values() if len(ans) == 1)
        multi_choice = sum(1 for ans in answers.values() if len(ans) > 1)
        print(f"单选题: {single_choice} 个")
        print(f"多选题: {multi_choice} 个")
        
        return answers
        
    except Exception as e:
        print(f"错误: {str(e)}")
        return {}

def test_regex_patterns():
    """测试不同的正则表达式模式"""
    
    # 测试文本样例（基于实际分析结果）
    test_texts = [
        "1.正确答案： A",
        "2.正确答案： D", 
        "50.正确答案： ABCD",
        "51.正确答案： BC",
        "100.正确答案： BCD",
        "1. 正确答案：A",  # 可能的变体
        "1.正确答案：A",   # 没有空格的变体
    ]
    
    # 测试不同的模式
    patterns = [
        r'^(\d+)\.正确答案：\s*([ABCD]+)$',        # 主要模式
        r'^(\d+)\.正确答案[：:]\s*([ABCD]+)$',     # 兼容中英文冒号
        r'^(\d+)\.\s*正确答案[：:]\s*([ABCD]+)$',  # 题号后可能有空格
        r'(\d+)\.正确答案[：:]\s*([ABCD]+)',       # 更宽松的匹配
    ]
    
    print("正则表达式模式测试:")
    print("=" * 60)
    
    for i, pattern in enumerate(patterns, 1):
        print(f"\n模式 {i}: {pattern}")
        print("-" * 40)
        
        for text in test_texts:
            match = re.match(pattern, text)
            if match:
                print(f"✓ '{text}' -> 题号: {match.group(1)}, 答案: {match.group(2)}")
            else:
                print(f"✗ '{text}' -> 不匹配")

def create_fixed_merge_function():
    """创建修正后的合并函数代码"""
    
    code = '''
def extract_answers_from_docx(answer_docx_path):
    """
    从答案文档中提取答案 - 修正版本
    
    基于对实际答案文档的分析，答案格式为：
    1.正确答案： A
    2.正确答案： D
    等等
    """
    answers = {}
    
    try:
        from docx import Document
        doc = Document(answer_docx_path)
        
        # 修正后的答案匹配模式
        answer_pattern = r'^(\d+)\.正确答案[：:]\s*([ABCD]+)$'
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            match = re.match(answer_pattern, text)
            if match:
                question_num = int(match.group(1))
                answer = match.group(2)
                answers[question_num] = answer
        
        print(f"从答案文档提取到 {len(answers)} 个答案")
        return answers
        
    except Exception as e:
        print(f"提取答案时出错: {str(e)}")
        return {}

# 在你的合并脚本中，替换原来的答案提取函数即可
'''
    
    return code

if __name__ == "__main__":
    print("2023年答案文档格式分析和修正方案")
    print("=" * 80)
    
    # 1. 测试正则表达式
    test_regex_patterns()
    
    print("\n" + "=" * 80)
    
    # 2. 实际提取答案
    answer_docx_path = "/Users/acheng/Downloads/2023 年卷一答案.docx"
    answers = extract_answers_from_docx(answer_docx_path)
    
    print("\n" + "=" * 80)
    print("修正后的合并函数代码:")
    print("=" * 80)
    print(create_fixed_merge_function())