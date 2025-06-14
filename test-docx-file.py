#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import sys
from docx import Document

def test_docx_file(file_path):
    """测试docx文件是否能正常打开"""
    print(f"测试文件: {file_path}")
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return False
    
    # 检查文件大小
    file_size = os.path.getsize(file_path)
    print(f"文件大小: {file_size} bytes ({file_size/1024:.2f} KB)")
    
    # 检查文件权限
    if not os.access(file_path, os.R_OK):
        print("❌ 文件没有读取权限")
        return False
    
    print("✅ 文件存在且有读取权限")
    
    # 尝试打开文档
    try:
        print("尝试打开文档...")
        doc = Document(file_path)
        print("✅ 文档打开成功")
        
        # 读取段落数量
        para_count = len(doc.paragraphs)
        print(f"文档包含 {para_count} 个段落")
        
        # 读取前5个段落
        print("\n前5个段落内容:")
        for i, para in enumerate(doc.paragraphs[:5]):
            text = para.text.strip()
            if text:
                print(f"  段落{i+1}: {text[:50]}...")
        
        return True
        
    except Exception as e:
        print(f"❌ 打开文档失败: {e}")
        print(f"错误类型: {type(e).__name__}")
        
        # 如果是Package not found错误，可能是文件格式问题
        if "Package not found" in str(e):
            print("\n可能的原因:")
            print("1. 文件不是标准的.docx格式")
            print("2. 文件可能是.doc格式(旧版Word)")
            print("3. 文件可能已损坏")
            print("4. 文件可能是其他软件生成的非标准docx")
        
        return False

if __name__ == "__main__":
    # 测试最近的刑法文件
    test_file = "uploads/刑法_1749471200053.docx"
    if os.path.exists(test_file):
        print("=== 测试已知可用的文件 ===")
        test_docx_file(test_file)
        print("\n")
    
    # 如果提供了命令行参数，测试指定文件
    if len(sys.argv) > 1:
        print("=== 测试指定文件 ===")
        test_docx_file(sys.argv[1])