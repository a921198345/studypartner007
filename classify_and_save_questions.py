import os
import re
import docx
import openai
import time
from docx import Document

# --- 配置 ---
# 使用OpenRouter API
openai.api_key = "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"
openai.base_url = "https://openrouter.ai/api/v1"

SOURCE_DOCX = '23年法考题.docx'
OUTPUT_DIR = 'docx'
SUBJECTS = ["民法", "刑法", "民事诉讼法", "刑事诉讼法", "行政法", "商经知", "三国法", "理论法"]
GPT_MODEL = "openai/gpt-3.5-turbo" # OpenRouter格式的模型名称

# --- 辅助函数 ---

def classify_question(question_text):
    """使用OpenAI API对问题进行分类"""
    if not question_text or not question_text.strip():
        return "未知"
    try:
        # print(f"正在分类问题: {question_text[:50]}...")
        completion = openai.chat.completions.create(
            model=GPT_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": f"你是一个专业的法律考试题目分类助手。请根据下面给出的问题内容，将其准确分类到以下八个科目中：{', '.join(SUBJECTS)}。请只返回科目名称，不要包含任何多余的解释或标点符号。"
                },
                {
                    "role": "user",
                    "content": question_text
                }
            ],
            temperature=0.0,
        )
        subject = completion.choices[0].message.content.strip()
        # print(f"分类结果: {subject}")
        # 增加一个小的延迟，避免过于频繁地调用API
        time.sleep(1) 
        return subject if subject in SUBJECTS else "未知"
    except Exception as e:
        print(f"调用API时发生错误: {e}")
        # 如果API调用失败，可以等待一段时间再重试
        time.sleep(5)
        return "未知"

def parse_questions(doc_path):
    """从DOCX文件中解析问题列表"""
    print(f"正在从 {doc_path} 文件中解析问题...")
    try:
        doc = Document(doc_path)
        questions = []
        current_question = ""
        # 正则表达式匹配问题编号，例如 "1."、"2."、"10."
        question_pattern = re.compile(r'^\d{1,3}\.')

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 如果段落以数字和点开头，我们认为它是一个新问题的开始
            if question_pattern.match(text):
                if current_question:
                    questions.append(current_question.strip())
                current_question = text
            else:
                current_question += "\n" + text
        
        # 不要忘记添加最后一个问题
        if current_question:
            questions.append(current_question.strip())
        
        print(f"成功解析出 {len(questions)} 个问题。")
        return questions
    except Exception as e:
        print(f"解析DOCX文件时出错: {e}")
        return []

# --- 主逻辑 ---

def main():
    """主执行函数"""
    print("开始执行题目分类任务...")
    
    # 1. 检查并创建输出目录
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        print(f"已创建输出目录: {OUTPUT_DIR}")

    # 2. 解析源文件中的所有问题
    all_questions = parse_questions(SOURCE_DOCX)
    if not all_questions:
        print("未能解析出任何问题，程序终止。")
        return

    # 3. 对每个问题进行分类
    print("开始对所有问题进行智能分类，这可能需要一些时间...")
    classified_data = {subject: [] for subject in SUBJECTS}
    unknown_questions = []

    total = len(all_questions)
    for i, question in enumerate(all_questions):
        print(f"正在处理第 {i+1}/{total} 个问题...")
        subject = classify_question(question)
        if subject != "未知":
            classified_data[subject].append(question)
        else:
            unknown_questions.append(question)
            print(f"警告：问题未能分类: {question[:50]}...")
    
    print("所有问题分类完成！")

    # 4. 将分类好的问题写入新的DOCX文件
    print("正在将分类结果写入新的Word文档...")
    for subject, questions in classified_data.items():
        if not questions:
            print(f"科目 '{subject}' 没有题目，跳过生成文件。")
            continue
            
        output_path = os.path.join(OUTPUT_DIR, f"{subject}.docx")
        new_doc = Document()
        new_doc.add_heading(subject, level=1)
        
        for q_text in questions:
            new_doc.add_paragraph(q_text)
            new_doc.add_paragraph() # 添加一个空段落作为分隔
            
        new_doc.save(output_path)
        print(f"已成功保存 {len(questions)} 个问题到: {output_path}")

    # 5. 如果有未分类的问题，也保存到一个单独的文件中
    if unknown_questions:
        output_path = os.path.join(OUTPUT_DIR, "未分类题目.docx")
        new_doc = Document()
        new_doc.add_heading("未分类题目", level=1)
        for q_text in unknown_questions:
            new_doc.add_paragraph(q_text)
            new_doc.add_paragraph()
        new_doc.save(output_path)
        print(f"警告：有 {len(unknown_questions)} 个问题未能分类，已保存到: {output_path}")

    print("\n所有任务已成功完成！")

if __name__ == "__main__":
    main() 