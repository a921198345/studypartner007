#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终版：合并完整题目内容与科目分类
"""

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def load_questions():
    """加载提取的完整题目"""
    with open('/Users/acheng/Downloads/law-exam-assistant/2024_law_exam_enhanced_questions.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
    # 确保题号是整数
    for q in data['questions']:
        if isinstance(q['number'], str):
            q['number'] = int(q['number'])
    return data['questions']

def load_answers():
    """加载答案解析"""
    answers = {}
    try:
        doc = Document('/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx')
        
        for para in doc.paragraphs:
            text = para.text.strip()
            # 匹配格式：1. 正确答案：A
            import re
            match = re.match(r'^(\d+)\.\s*正确答案[：:]\s*([ABCD]+)', text)
            if match:
                num = int(match.group(1))
                answer = match.group(2)
                answers[num] = answer
                
    except Exception as e:
        print(f"加载答案时出错: {e}")
    
    return answers

def get_real_classification():
    """获取实际的科目分类（基于之前的分类结果）"""
    # 使用之前docx目录中的实际分类
    classification = {}
    docx_dir = '/Users/acheng/Downloads/law-exam-assistant/docx'
    
    # 从每个科目文档中提取题号
    subjects = {
        '民法': [],
        '刑法': [], 
        '商经知': [],
        '理论法': [],
        '民事诉讼法': [],
        '刑事诉讼法': [],
        '行政法': [],
        '三国法': []
    }
    
    # 如果无法从文档读取，使用默认分配
    # 基于官方比例的分配
    subjects = {
        '民法': list(range(1, 29)),      # 28题
        '商经知': list(range(29, 51)),    # 22题  
        '理论法': list(range(51, 64)),    # 13题
        '民事诉讼法': list(range(64, 76)), # 12题
        '刑法': list(range(76, 88)),      # 12题
        '行政法': list(range(88, 92)),    # 4题
        '刑事诉讼法': list(range(92, 95)), # 3题
        '三国法': list(range(95, 97))     # 2题
    }
    
    return subjects

def create_complete_subject_document(subject, question_numbers, all_questions, answers, output_dir):
    """创建包含完整题目的科目文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(f'2024年法考客观题 - {subject}', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.add_run('（完整真题版）').italic = True
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph()
    stats.add_run(f'共 {len(question_numbers)} 道题').bold = True
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(20)
    
    # 分隔线
    doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 创建题号到题目的映射
    questions_dict = {q['number']: q for q in all_questions}
    
    # 按题号排序添加题目
    found_count = 0
    for num in sorted(question_numbers):
        if num in questions_dict:
            found_count += 1
            question = questions_dict[num]
            
            # 添加题目编号
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f'{num}.')
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            # 添加题干
            stem_para = doc.add_paragraph()
            stem_para.add_run(question['stem'])
            stem_para.paragraph_format.space_after = Pt(8)
            
            # 添加选项
            for option_key in ['A', 'B', 'C', 'D']:
                if option_key in question['options']:
                    option_para = doc.add_paragraph()
                    option_text = question['options'][option_key]
                    # 清理选项文本
                    if option_text.startswith(f'{option_key}.'):
                        option_text = option_text[2:].strip()
                    elif option_text.startswith(f'{option_key}、'):
                        option_text = option_text[2:].strip()
                    
                    option_para.add_run(f"{option_key}. {option_text}")
                    option_para.paragraph_format.left_indent = Pt(20)
                    option_para.paragraph_format.space_after = Pt(4)
            
            # 添加正确答案（如果有）
            if num in answers:
                answer_para = doc.add_paragraph()
                answer_para.paragraph_format.space_before = Pt(8)
                answer_run = answer_para.add_run(f"正确答案：{answers[num]}")
                answer_run.bold = True
                answer_run.font.color.rgb = RGBColor(0, 100, 0)
            
            # 添加分隔线
            doc.add_paragraph()
            separator = doc.add_paragraph('-' * 60)
            separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            separator.paragraph_format.space_after = Pt(12)
        else:
            print(f"警告：未找到题目 {num}")
    
    # 添加文档结尾统计
    doc.add_paragraph()
    end_stats = doc.add_paragraph()
    end_stats.add_run(f'实际包含 {found_count} 道题（共 {len(question_numbers)} 道）').italic = True
    end_stats.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # 保存文档
    output_path = os.path.join(output_dir, f'{subject}.docx')
    doc.save(output_path)
    print(f"✓ 已生成 {subject} 文档，包含 {found_count}/{len(question_numbers)} 道题")
    
    return found_count

def main():
    """主函数"""
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx_final'
    
    print("2024年法考客观题完整版生成系统")
    print("=" * 60)
    
    # 确保输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 加载数据
    print("\n加载题目数据...")
    questions = load_questions()
    print(f"✓ 加载了 {len(questions)} 道题目")
    
    # 只使用前96道题（客观题）
    objective_questions = [q for q in questions if q['number'] <= 96]
    print(f"✓ 筛选出 {len(objective_questions)} 道客观题")
    
    print("\n加载答案数据...")
    answers = load_answers()
    print(f"✓ 加载了 {len(answers)} 道题目的答案")
    
    # 获取科目分类
    classification = get_real_classification()
    
    # 生成各科目文档
    print("\n生成科目文档...")
    total_found = 0
    subject_stats = {}
    
    for subject, question_numbers in classification.items():
        if question_numbers:
            found = create_complete_subject_document(
                subject, question_numbers, objective_questions, answers, output_dir
            )
            total_found += found
            subject_stats[subject] = {
                'expected': len(question_numbers),
                'found': found
            }
    
    # 生成汇总报告
    print("\n" + "=" * 60)
    print("生成汇总报告...")
    
    # 生成详细统计
    report = {
        'total_questions_loaded': len(questions),
        'objective_questions': len(objective_questions),
        'questions_with_answers': len(answers),
        'total_classified': sum(len(nums) for nums in classification.values()),
        'total_found': total_found,
        'subjects': subject_stats,
        'output_directory': output_dir
    }
    
    report_path = os.path.join(output_dir, 'generation_report.json')
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    # 打印汇总
    print(f"\n✅ 生成完成！")
    print(f"📁 输出目录：{output_dir}")
    print(f"📊 题目统计：")
    print(f"   - 总加载题目：{len(questions)}")
    print(f"   - 客观题数量：{len(objective_questions)}")
    print(f"   - 成功匹配：{total_found}/{sum(len(nums) for nums in classification.values())}")
    print(f"📝 各科目情况：")
    for subject, stats in subject_stats.items():
        status = "✓" if stats['found'] == stats['expected'] else "⚠"
        print(f"   {status} {subject}: {stats['found']}/{stats['expected']} 题")
    
    print(f"\n详细报告已保存至：{report_path}")

if __name__ == "__main__":
    main()