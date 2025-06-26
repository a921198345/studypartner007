#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将法考题按科目分类
根据题目内容智能分类到8个科目：民法、刑法、民事诉讼法、刑事诉讼法、行政法、商经知、三国法、理论法
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 科目关键词映射
SUBJECT_KEYWORDS = {
    '民法': [
        '民法', '合同', '物权', '债权', '侵权', '婚姻', '继承', '人格权', '担保', '抵押', '质押', 
        '买卖', '租赁', '借款', '赠与', '保证', '定金', '违约', '损害赔偿', '所有权', '用益物权',
        '占有', '善意取得', '不当得利', '无因管理', '婚姻家庭', '离婚', '抚养', '赡养', '遗嘱',
        '法定继承', '遗赠', '著作权', '专利权', '商标权', '知识产权', '隐私权', '名誉权'
    ],
    '刑法': [
        '刑法', '犯罪', '刑罚', '故意', '过失', '正当防卫', '紧急避险', '犯罪构成', '主观要件',
        '客观要件', '共同犯罪', '主犯', '从犯', '教唆犯', '累犯', '自首', '立功', '缓刑', '假释',
        '盗窃', '抢劫', '诈骗', '贪污', '受贿', '行贿', '挪用', '职务侵占', '故意杀人', '故意伤害',
        '强奸', '绑架', '非法拘禁', '敲诈勒索', '寻衅滋事', '聚众斗殴', '毒品', '走私', '伪造',
        '危险驾驶', '交通肇事', '渎职', '玩忽职守', '徇私枉法', '量刑', '死刑', '无期徒刑', '有期徒刑',
        '拘役', '管制', '罚金', '没收财产', '剥夺政治权利'
    ],
    '民事诉讼法': [
        '民事诉讼', '民诉', '起诉', '受理', '管辖', '级别管辖', '地域管辖', '移送管辖', '管辖权异议',
        '原告', '被告', '第三人', '诉讼代理人', '举证责任', '证据', '证明标准', '调解', '撤诉',
        '缺席判决', '简易程序', '普通程序', '二审', '上诉', '再审', '审判监督', '执行', '强制执行',
        '财产保全', '先予执行', '诉讼时效', '期间', '送达', '开庭', '法庭辩论', '判决', '裁定',
        '调解书', '支付令', '公示催告', '诉讼费'
    ],
    '刑事诉讼法': [
        '刑事诉讼', '刑诉', '立案', '侦查', '起诉', '审判', '执行', '管辖', '回避', '辩护', '辩护人',
        '犯罪嫌疑人', '被告人', '被害人', '证人', '鉴定', '勘验', '检查', '搜查', '扣押', '逮捕',
        '拘留', '取保候审', '监视居住', '拘传', '强制措施', '侦查终结', '审查起诉', '公诉', '自诉',
        '简易程序', '普通程序', '二审程序', '死刑复核', '审判监督程序', '执行程序', '刑事和解',
        '认罪认罚', '速裁程序', '缺席审判'
    ],
    '行政法': [
        '行政法', '行政行为', '行政许可', '行政处罚', '行政强制', '行政复议', '行政诉讼', '行政主体',
        '行政相对人', '具体行政行为', '抽象行政行为', '行政立法', '行政决定', '行政命令', '行政指导',
        '行政合同', '行政给付', '行政征收', '行政征用', '行政确认', '行政监督', '行政责任', '国家赔偿',
        '行政机关', '公务员', '行政程序', '听证', '告知', '说明理由', '政府信息公开', '信访'
    ],
    '商经知': [
        '商法', '经济法', '公司法', '公司', '股东', '董事', '监事', '股权', '股份', '有限责任公司',
        '股份有限公司', '合伙', '合伙企业', '个人独资', '破产', '清算', '重整', '和解', '证券',
        '股票', '债券', '基金', '期货', '保险', '保险合同', '投保人', '被保险人', '受益人', '保险金',
        '票据', '汇票', '本票', '支票', '背书', '承兑', '保证', '追索权', '信托', '信托财产',
        '反垄断', '反不正当竞争', '消费者权益', '产品质量', '食品安全', '环境保护', '劳动法',
        '劳动合同', '社会保险', '工伤', '税法', '增值税', '所得税', '房地产', '建设工程', '招投标'
    ],
    '三国法': [
        '国际法', '国际公法', '国际私法', '国际经济法', '国家主权', '国家领土', '领海', '领空',
        '外交关系', '领事关系', '条约', '国际习惯', '国际组织', '联合国', '国际法院', '国际责任',
        '国际争端', '战争法', '国际人权法', '引渡', '庇护', '国籍', '外国人', '冲突规范', '准据法',
        '法律冲突', '反致', '公共秩序保留', '法律规避', '识别', '涉外合同', '涉外侵权', '涉外婚姻',
        '涉外继承', '国际商事仲裁', '司法协助', '外国判决承认执行', 'WTO', '国际贸易', '国际投资',
        '国际金融', '国际税收', '海商法', '海事', '海损', '共同海损', '海上保险', '提单', '租船合同'
    ],
    '理论法': [
        '法理学', '法理', '法的概念', '法的本质', '法的特征', '法的作用', '法的价值', '正义', '自由',
        '秩序', '法的要素', '法律规则', '法律原则', '法律概念', '法的渊源', '法律体系', '法的效力',
        '法律关系', '权利', '义务', '法律事实', '法律行为', '法的运行', '立法', '执法', '司法',
        '守法', '法律监督', '法律解释', '法律推理', '法律论证', '法治', '依法治国', '宪法', '宪法学',
        '国家制度', '根本制度', '基本制度', '选举制度', '国家机构', '全国人大', '国务院', '监察委',
        '人民法院', '人民检察院', '基本权利', '公民权利', '政治权利', '人身自由', '财产权', '社会权利',
        '法制史', '中国法制史', '外国法制史', '司法制度', '法律职业道德', '职业道德', '职业操守'
    ]
}

def classify_question(question_text):
    """根据题目内容判断所属科目"""
    # 统计每个科目的关键词匹配数
    subject_scores = {}
    
    for subject, keywords in SUBJECT_KEYWORDS.items():
        score = 0
        for keyword in keywords:
            if keyword in question_text:
                # 根据关键词长度给不同权重
                if len(keyword) >= 4:
                    score += 3
                elif len(keyword) >= 3:
                    score += 2
                else:
                    score += 1
        
        if score > 0:
            subject_scores[subject] = score
    
    # 如果没有匹配到任何关键词，返回None
    if not subject_scores:
        return None
    
    # 返回得分最高的科目
    return max(subject_scores.items(), key=lambda x: x[1])[0]

def extract_questions_with_answers(doc_path):
    """从合并文档中提取题目和答案"""
    doc = Document(doc_path)
    questions = {}
    current_question = None
    current_content = []
    
    # 匹配题号的正则表达式 - 更宽松的匹配
    question_pattern = re.compile(r'^(\d+)\.(.*)$')  # 匹配 "1.xxx" 格式
    separator_pattern = re.compile(r'^-{10,}$')  # 匹配至少10个连字符
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 跳过标题行
        if '法考' in text and '题' in text and len(text) < 30:
            continue
        
        # 检查是否是分隔线
        if separator_pattern.match(text) or text == '-' * 30:
            # 保存当前题目
            if current_question and current_content:
                questions[current_question] = '\n'.join(current_content)
                current_question = None
                current_content = []
            continue
        
        # 检查是否是新题目（题号在行首）
        match = question_pattern.match(text)
        if match:
            # 检查是否是答案行（通常包含"正确答案"）
            if '正确答案' in text:
                # 这是答案部分，添加到当前内容
                if current_question:
                    current_content.append(text)
            else:
                # 这是新题目
                # 保存前一题
                if current_question and current_content:
                    questions[current_question] = '\n'.join(current_content)
                
                # 开始新题目
                current_question = int(match.group(1))
                current_content = [text]
        
        # 继续收集内容
        elif current_question and text:
            current_content.append(text)
    
    # 保存最后一题
    if current_question and current_content:
        questions[current_question] = '\n'.join(current_content)
    
    return questions

def create_subject_document(subject, questions, output_dir):
    """为特定科目创建文档"""
    doc = Document()
    
    # 设置文档的默认字体
    doc.styles['Normal'].font.name = '宋体'
    
    # 添加标题
    title = doc.add_heading(f'2023年法考真题 - {subject}', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph(f'共 {len(questions)} 道题')
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(12)
    
    # 添加分隔线
    doc.add_paragraph("═" * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 添加题目
    for q_num, content in sorted(questions.items()):
        # 题目内容
        doc.add_paragraph(content)
        
        # 添加分隔线
        separator = doc.add_paragraph("─" * 80)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(12)
        separator.paragraph_format.space_after = Pt(12)
    
    # 保存文档
    filename = f"{subject}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filename

def main():
    # 输入文件路径
    input_file = '/Users/acheng/Downloads/law-exam-assistant/23年法考题.docx'
    
    # 检查文件是否存在
    if not os.path.exists(input_file):
        print(f"错误：找不到文件 {input_file}")
        print("请确认文件路径是否正确")
        return
    
    # 创建输出目录
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx'
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"开始处理文件: {input_file}")
    print(f"输出目录: {output_dir}")
    
    # 提取所有题目
    print("\n提取题目中...")
    all_questions = extract_questions_with_answers(input_file)
    print(f"共提取到 {len(all_questions)} 道题目")
    
    # 按科目分类
    subject_questions = {subject: {} for subject in SUBJECT_KEYWORDS.keys()}
    unclassified = {}
    
    print("\n开始分类...")
    for q_num, content in all_questions.items():
        # 只用题目部分进行分类（不包括答案）
        question_part = content.split('【答案解析】')[0]
        subject = classify_question(question_part)
        
        if subject:
            subject_questions[subject][q_num] = content
        else:
            unclassified[q_num] = content
    
    # 生成各科目文档
    print("\n生成科目文档...")
    for subject, questions in subject_questions.items():
        if questions:
            filename = create_subject_document(subject, questions, output_dir)
            print(f"{subject}: {len(questions)} 道题 -> {filename}")
    
    # 如果有未分类的题目，创建单独文档
    if unclassified:
        filename = create_subject_document("未分类", unclassified, output_dir)
        print(f"未分类: {len(unclassified)} 道题 -> {filename}")
    
    # 输出统计信息
    print("\n分类统计:")
    total_classified = sum(len(q) for q in subject_questions.values())
    print(f"已分类: {total_classified} 道")
    print(f"未分类: {len(unclassified)} 道")
    print(f"总计: {len(all_questions)} 道")
    
    print("\n处理完成！")

if __name__ == "__main__":
    main()