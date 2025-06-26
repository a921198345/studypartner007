#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于答案解析生成完整的法考试题文档
由于原题目文档格式问题，主要基于答案解析来重构完整的试题
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os

def extract_complete_answers(doc_path):
    """从答案解析文档中提取完整信息"""
    doc = Document(doc_path)
    questions_data = {}
    current_question = None
    current_content = []
    
    print(f"开始解析答案文档: {doc_path}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目（匹配 "数字. 正确答案：X" 格式）
        answer_match = re.match(r'^(\d+)\.\s*正确答案[：:]\s*([A-Z])', text)
        if answer_match:
            # 保存前一题
            if current_question is not None:
                questions_data[current_question]['content'] = '\n'.join(current_content)
            
            # 开始新题目
            current_question = int(answer_match.group(1))
            correct_answer = answer_match.group(2)
            
            questions_data[current_question] = {
                'answer': correct_answer,
                'content': '',
                'explanation': []
            }
            current_content = []
            
        elif current_question is not None:
            # 收集当前题目的解析内容
            if text == '【答案解析】':
                continue  # 跳过标题
            elif text.startswith('综上所述，本题答案为'):
                # 这是总结，结束当前题目
                questions_data[current_question]['content'] = '\n'.join(current_content)
                current_content = []
            else:
                current_content.append(text)
    
    # 保存最后一题
    if current_question is not None:
        questions_data[current_question]['content'] = '\n'.join(current_content)
    
    print(f"提取了 {len(questions_data)} 道题的答案解析")
    return questions_data

def extract_question_text_from_answers(answer_content):
    """从答案解析中尝试提取或推断题目内容"""
    lines = answer_content.split('\n')
    question_indicators = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 查找包含选项分析的内容，可能包含题目信息
        if re.search(r'[ABCD]\s*项?[：:]', line) or re.search(r'[ABCD]\s*选项', line):
            question_indicators.append(line)
    
    return question_indicators

def create_question_from_analysis(num, data):
    """基于答案解析创建题目内容"""
    answer = data['answer']
    content = data['content']
    
    # 分析内容，尝试重构题目
    lines = content.split('\n')
    
    # 构建题目描述
    question_text = f"第 {num} 题（基于答案解析重构）\n"
    question_text += f"正确答案：{answer}\n\n"
    
    # 添加解析内容
    explanation_parts = []
    options_found = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 检查是否是选项分析
        if re.search(r'^[ABCD]\s*项?[：:]', line):
            if not options_found:
                explanation_parts.append("【选项分析】")
                options_found = True
            explanation_parts.append(f"• {line}")
        elif re.search(r'^[ABCD]\s*选项', line):
            if not options_found:
                explanation_parts.append("【选项分析】")
                options_found = True
            explanation_parts.append(f"• {line}")
        else:
            explanation_parts.append(line)
    
    return question_text, explanation_parts

def add_question_section(doc, num, question_text, explanation_parts, answer):
    """添加题目部分到文档"""
    
    # 题目标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"第 {num} 题")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色
    title_para.paragraph_format.space_before = Pt(16)
    title_para.paragraph_format.space_after = Pt(8)
    
    # 答案信息
    answer_para = doc.add_paragraph()
    answer_run = answer_para.add_run(f"正确答案：{answer}")
    answer_run.font.size = Pt(12)
    answer_run.bold = True
    answer_run.font.color.rgb = RGBColor(128, 0, 0)  # 深红色
    answer_para.paragraph_format.space_after = Pt(8)
    
    # 解析标题
    analysis_title = doc.add_paragraph()
    analysis_run = analysis_title.add_run("【答案解析】")
    analysis_run.font.size = Pt(12)
    analysis_run.bold = True
    analysis_run.font.color.rgb = RGBColor(0, 128, 0)  # 深绿色
    analysis_title.paragraph_format.space_before = Pt(8)
    analysis_title.paragraph_format.space_after = Pt(6)
    
    # 解析内容
    for part in explanation_parts:
        if part.strip():
            if part.startswith('•'):
                # 选项分析，使用缩进
                option_para = doc.add_paragraph(part)
                option_para.paragraph_format.left_indent = Pt(20)
                option_para.paragraph_format.space_after = Pt(4)
            elif part.startswith('【') and part.endswith('】'):
                # 小标题
                subtitle_para = doc.add_paragraph()
                subtitle_run = subtitle_para.add_run(part)
                subtitle_run.bold = True
                subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
                subtitle_para.paragraph_format.space_before = Pt(6)
                subtitle_para.paragraph_format.space_after = Pt(4)
            else:
                # 普通解析内容
                explanation_para = doc.add_paragraph(part)
                explanation_para.paragraph_format.space_after = Pt(4)
    
    # 分隔线
    separator = doc.add_paragraph("─" * 60)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    separator.paragraph_format.space_before = Pt(12)
    separator.paragraph_format.space_after = Pt(16)

def create_complete_document(questions_data, output_path):
    """创建完整的题目文档"""
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    # 文档标题
    title = doc.add_heading('2024年法考客观题（完整版）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('真题答案及详细解析')
    subtitle_run.font.size = Pt(14)
    subtitle_run.bold = True
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    
    # 说明文字
    note = doc.add_paragraph()
    note.add_run("说明：本文档基于2024年法考答案解析重新整理，每题包含正确答案和详细解析。").italic = True
    note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    note.paragraph_format.space_after = Pt(10)
    
    # 统计信息
    stats = doc.add_paragraph()
    stats.add_run(f"共收录 {len(questions_data)} 道题目，每题均含答案解析")
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(30)
    
    # 分隔线
    doc.add_paragraph("═" * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    print(f"\n开始生成完整文档...")
    print(f"总题数: {len(questions_data)}")
    
    # 按题号顺序添加题目
    for i, num in enumerate(sorted(questions_data.keys()), 1):
        data = questions_data[num]
        
        if i % 10 == 0:
            print(f"处理进度: {i}/{len(questions_data)}")
        
        # 创建题目内容
        question_text, explanation_parts = create_question_from_analysis(num, data)
        
        # 添加到文档
        add_question_section(doc, num, question_text, explanation_parts, data['answer'])
    
    # 保存文档
    doc.save(output_path)
    print(f"\n文档生成完成！")
    print(f"保存路径: {output_path}")

def main():
    """主函数"""
    answers_file = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/2024年法考完整题目-基于解析版.docx'
    
    if not os.path.exists(answers_file):
        print(f"错误：找不到答案文件 {answers_file}")
        return
    
    print("2024年法考题目重构工具")
    print("=" * 50)
    print("基于答案解析重新生成完整的题目文档")
    print("=" * 50)
    
    # 提取答案数据
    questions_data = extract_complete_answers(answers_file)
    
    if not questions_data:
        print("错误：没有提取到题目数据")
        return
    
    # 生成完整文档
    create_complete_document(questions_data, output_file)
    
    print(f"\n处理完成！共生成 {len(questions_data)} 道题目")
    print(f"文档已保存至：{output_file}")

if __name__ == "__main__":
    main()