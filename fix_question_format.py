import re
import sys
import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def validate_question_format(text):
    """验证题目格式并返回问题列表"""
    issues = []
    
    # 检查题号格式
    if not re.search(r'【\d{8}】', text):
        issues.append("题号格式不正确，应为【8位数字】")
    
    # 检查选项数量
    options = re.findall(r'[A-D]\.', text)
    if len(options) != 4:
        issues.append(f"选项数量不正确，应有4个选项(A-D)，实际有{len(options)}个")
    
    # 检查选项格式统一性
    if not all(option in text for option in ["A.", "B.", "C.", "D."]):
        issues.append("选项标识不完整或格式不统一")
    
    # 检查解析标记
    if "【解析】" not in text:
        issues.append("缺少【解析】标记")
    
    # 检查是否有答案标识
    has_answer = any(pattern in text for pattern in [
        "答案：", "正确答案", "故选", "本题答案", "【答案】"
    ])
    if not has_answer:
        issues.append("解析中缺少明确的答案标识")
    
    # 检查多余空行
    if "\n\n\n" in text:
        issues.append("存在多余空行")
    
    return issues

def fix_question_format(text):
    """修复题目格式问题"""
    fixed_text = text
    fixed_issues = []
    
    # 1. 修复题号格式
    question_id_match = re.search(r'[【\[（(](\d+)[】\]）)]', fixed_text)
    if question_id_match:
        question_id = question_id_match.group(1)
        # 确保题号是8位数字
        if len(question_id) != 8:
            # 如果少于8位，在前面补0
            if len(question_id) < 8:
                new_id = question_id.zfill(8)
                fixed_text = re.sub(r'[【\[（(]' + question_id + r'[】\]）)]', f'【{new_id}】', fixed_text)
                fixed_issues.append(f"题号从{question_id}修复为{new_id}")
            # 如果超过8位，取最后8位
            else:
                new_id = question_id[-8:]
                fixed_text = re.sub(r'[【\[（(]' + question_id + r'[】\]）)]', f'【{new_id}】', fixed_text)
                fixed_issues.append(f"题号从{question_id}修复为{new_id}")
    else:
        # 如果完全没有题号，尝试生成一个
        year = "20240101"  # 默认题号
        fixed_text = f"【{year}】" + fixed_text
        fixed_issues.append(f"添加缺失的题号：{year}")
    
    # 2. 修复选项格式
    options = re.findall(r'([A-D])[\.、．:： ](.*?)(?=[A-D][\.、．:： ]|【解析】|$)', fixed_text, re.DOTALL)
    
    # 如果找到的选项少于4个，尝试找出所有可能的选项文本
    if len(options) < 4:
        # 尝试用其他分隔符查找
        potential_options = re.findall(r'([A-D])[^A-Z]+(.*?)(?=[A-D][^A-Z]+|【解析】|$)', fixed_text, re.DOTALL)
        
        if len(potential_options) >= len(options):
            options = potential_options
    
    # 统一选项格式
    standard_options = {}
    for opt in options:
        if len(opt) >= 2:
            letter, content = opt[0], opt[1].strip()
            standard_options[letter] = content
    
    # 确保有ABCD四个选项
    required_options = ['A', 'B', 'C', 'D']
    for letter in required_options:
        if letter not in standard_options:
            standard_options[letter] = f"选项{letter}内容（自动添加）"
            fixed_issues.append(f"添加缺失的{letter}选项")
    
    # 替换选项部分
    # 首先找到题干和选项的分界点
    question_parts = re.split(r'([A-D][\.、．:： ])', fixed_text, 1)
    if len(question_parts) > 1:
        question_text = question_parts[0]
        
        # 构建新的选项文本
        options_text = ""
        for letter in required_options:
            options_text += f"{letter}. {standard_options[letter]}\n"
        
        # 查找解析部分
        analysis_match = re.search(r'【解析】(.*?)$', fixed_text, re.DOTALL)
        analysis_text = ""
        if analysis_match:
            analysis_text = "【解析】" + analysis_match.group(1).strip()
        else:
            # 如果没有解析部分，尝试创建一个
            analysis_text = "【解析】本题暂无解析。答案：A"
            fixed_issues.append("添加缺失的解析部分")
        
        # 重构题目
        fixed_text = question_text + "\n" + options_text + "\n" + analysis_text
    
    # 3. 修复解析部分
    if "【解析】" not in fixed_text:
        # 尝试查找可能的解析标记
        analysis_match = re.search(r'(解析|分析|答案解释)[：:](.*?)$', fixed_text, re.DOTALL)
        if analysis_match:
            analysis_content = analysis_match.group(2).strip()
            # 替换为标准格式
            fixed_text = fixed_text.replace(analysis_match.group(0), f"【解析】{analysis_content}")
            fixed_issues.append("修复解析标记格式")
    
    # 4. 确保解析中有答案标识
    if not any(pattern in fixed_text for pattern in ["答案：", "正确答案", "故选", "本题答案", "【答案】"]):
        # 如果解析中没有明确的答案标识，添加一个默认答案
        analysis_part = re.search(r'【解析】(.*?)$', fixed_text, re.DOTALL)
        if analysis_part:
            original_analysis = analysis_part.group(1)
            new_analysis = original_analysis + "\n答案：A"  # 添加默认答案
            fixed_text = fixed_text.replace(analysis_part.group(0), f"【解析】{new_analysis}")
            fixed_issues.append("添加缺失的答案标识")
    
    # 5. 修复多余空行
    fixed_text = re.sub(r'\n{3,}', '\n\n', fixed_text)
    
    return fixed_text, fixed_issues

def extract_and_fix_questions(input_file_path, output_file_path, subject):
    """从Word文档中提取所有题目，修复格式问题，并创建新的Word文档"""
    input_doc = Document(input_file_path)
    full_text = ""
    
    # 将所有段落合并为一个文本
    for para in input_doc.paragraphs:
        full_text += para.text + "\n"
    
    # 按题号分割文本
    question_pattern = r'【\d+】.*?(?=【\d+】|$)'
    questions_text = re.findall(question_pattern, full_text, re.DOTALL)
    
    # 如果没有找到题目，尝试其他分隔方式
    if not questions_text:
        # 尝试其他可能的题号格式
        alt_question_pattern = r'[【\[（(]\d+[】\]）)].*?(?=[【\[（(]\d+[】\]）)]|$)'
        questions_text = re.findall(alt_question_pattern, full_text, re.DOTALL)
    
    # 如果仍然没有找到题目，尝试按段落分割
    if not questions_text:
        # 按段落分割，假设每个题目占多个段落
        paragraphs = full_text.split('\n\n')
        current_question = ""
        questions_text = []
        
        for para in paragraphs:
            if re.search(r'^[0-9]+[\.、．:： ]', para.strip()):  # 如果段落以数字开头，可能是新题目
                if current_question:
                    questions_text.append(current_question)
                current_question = para + "\n"
            else:
                current_question += para + "\n"
        
        if current_question:
            questions_text.append(current_question)
    
    # 创建新文档
    output_doc = Document()
    
    # 记录修复情况
    all_fixes = {}
    fixed_count = 0
    total_count = len(questions_text)
    
    # 处理每个题目
    for i, q_text in enumerate(questions_text):
        # 验证并修复格式
        issues = validate_question_format(q_text)
        
        if issues:
            fixed_text, fixed_issues = fix_question_format(q_text)
            
            # 记录修复情况
            question_id_match = re.search(r'【(\d+)】', fixed_text)
            question_id = question_id_match.group(1) if question_id_match else f"未知题号_{i+1}"
            
            all_fixes[question_id] = {
                "original_issues": issues,
                "fixed_issues": fixed_issues
            }
            
            fixed_count += 1
            
            # 将修复后的题目添加到新文档
            para = output_doc.add_paragraph(fixed_text)
        else:
            # 如果没有问题，直接添加原始题目
            para = output_doc.add_paragraph(q_text)
        
        # 添加段落间距
        if i < len(questions_text) - 1:
            output_doc.add_paragraph()
    
    # 保存新文档
    output_doc.save(output_file_path)
    
    # 打印修复统计
    print(f"总共处理 {total_count} 道题目，修复了 {fixed_count} 道题目的格式问题。")
    
    # 返回修复报告
    return {
        "total_questions": total_count,
        "fixed_questions": fixed_count,
        "fixes": all_fixes
    }

def main():
    if len(sys.argv) < 3:
        print("用法: python fix_question_format.py <输入文件路径> <输出文件路径> <学科名称>")
        print("例如: python fix_question_format.py 民法原始.docx 民法修复.docx 民法")
        return
    
    input_file_path = sys.argv[1]
    output_file_path = sys.argv[2]
    subject = sys.argv[3] if len(sys.argv) > 3 else "未知学科"
    
    if not os.path.exists(input_file_path):
        print(f"错误: 文件 '{input_file_path}' 不存在")
        return
    
    try:
        print(f"正在处理文件: {input_file_path}...")
        fix_report = extract_and_fix_questions(input_file_path, output_file_path, subject)
        
        print(f"修复完成，已将结果保存到: {output_file_path}")
        print(f"修复报告: {json.dumps(fix_report, ensure_ascii=False)}")
        
    except Exception as e:
        print(f"程序出错: {e}")

if __name__ == "__main__":
    main() 