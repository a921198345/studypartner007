#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2023年法考题目分类脚本

根据官方科目分值占比对2023年法考题目进行分类：
- 理论法（法治思想+法理学+宪法+法史+司法制度）：约20%（约40题）
- 商经知（商法+经济法+环境法+劳动法+知识产权）：约29%（约58题）
- 民法：约20%（约40题）
- 民事诉讼法：约15%（约30题）
- 刑法：约12%（约24题）
- 刑事诉讼法：约8%（约16题）
- 行政法：约8%（约16题）
- 三国法（国际法+国际私法+国际经济法）：约10%（约20题）
"""

import os
import re
import json
from docx import Document
from collections import defaultdict
import jieba
import jieba.posseg as pseg

# 初始化jieba
jieba.initialize()

class LawExamClassifier:
    def __init__(self):
        # 官方科目分值占比
        self.official_ratios = {
            '理论法': 0.20,  # 40题
            '商经知': 0.29,  # 58题
            '民法': 0.20,    # 40题
            '民事诉讼法': 0.15,  # 30题
            '刑法': 0.12,    # 24题
            '刑事诉讼法': 0.08,  # 16题
            '行政法': 0.08,  # 16题
            '三国法': 0.10   # 20题
        }
        
        # 科目关键词库
        self.subject_keywords = {
            '理论法': {
                '法治思想': ['新时代', '中国特色社会主义', '习近平', '全面依法治国', '法治建设', '法治体系', '法治国家', '法治政府', '法治社会'],
                '法理学': ['法的概念', '法的特征', '法的作用', '法的价值', '法的效力', '法的渊源', '法的适用', '法律关系', '法律责任', '法律解释'],
                '宪法': ['宪法', '人民代表大会', '国家主席', '国务院', '监察委员会', '人民法院', '人民检察院', '国家机构', '公民权利', '基本权利'],
                '法史': ['法制史', '中国法制史', '外国法制史', '传统法律', '法律发展', '历史沿革'],
                '司法制度': ['司法制度', '法官', '检察官', '律师', '公证', '仲裁', '司法考试', '司法改革']
            },
            '商经知': {
                '商法': ['公司法', '合伙企业法', '个人独资企业法', '企业破产法', '证券法', '保险法', '票据法', '海商法', '公司', '股东', '董事', '监事'],
                '经济法': ['反垄断法', '反不正当竞争法', '消费者权益保护法', '产品质量法', '价格法', '银行业监督管理法', '垄断', '竞争', '消费者'],
                '环境法': ['环境保护法', '大气污染防治法', '水污染防治法', '固体废物污染防治法', '环境影响评价法', '环境', '污染', '生态'],
                '劳动法': ['劳动法', '劳动合同法', '劳动争议调解仲裁法', '社会保险法', '劳动', '工资', '工作时间', '劳动合同', '社会保险'],
                '知识产权': ['专利法', '商标法', '著作权法', '专利', '商标', '著作权', '知识产权', '版权', '发明', '实用新型']
            },
            '民法': {
                '民法总则': ['民事主体', '民事权利', '民事法律行为', '代理', '诉讼时效', '期间', '自然人', '法人', '非法人组织'],
                '物权': ['物权', '所有权', '用益物权', '担保物权', '占有', '抵押', '质押', '留置', '不动产', '动产'],
                '合同': ['合同', '债权', '债务', '违约', '违约责任', '合同履行', '合同解除', '合同终止', '买卖', '租赁'],
                '人格权': ['人格权', '生命权', '身体权', '健康权', '姓名权', '肖像权', '名誉权', '荣誉权', '隐私权'],
                '婚姻家庭': ['婚姻', '家庭', '夫妻', '离婚', '子女', '抚养', '赡养', '继承', '遗产', '遗嘱'],
                '侵权责任': ['侵权', '侵权责任', '过错', '损害赔偿', '人身损害', '财产损害', '精神损害']
            },
            '民事诉讼法': ['起诉', '受理', '审理', '判决', '裁定', '执行', '当事人', '代理人', '证据', '举证', '质证', '认证', '管辖', '回避', '期间', '送达'],
            '刑法': {
                '刑法总则': ['犯罪', '刑罚', '犯罪构成', '主观方面', '客观方面', '犯罪主体', '犯罪客体', '故意', '过失', '未遂', '中止', '共同犯罪'],
                '刑法分则': ['危害公共安全罪', '破坏社会主义市场经济秩序罪', '侵犯公民人身权利罪', '侵犯财产罪', '妨害社会管理秩序罪', '危害国防利益罪', '贪污贿赂罪', '渎职罪']
            },
            '刑事诉讼法': ['立案', '侦查', '起诉', '审判', '执行', '公安机关', '人民检察院', '辩护', '代理', '强制措施', '搜查', '扣押', '逮捕', '取保候审'],
            '行政法': ['行政行为', '行政处罚', '行政许可', '行政强制', '行政复议', '行政诉讼', '行政机关', '行政相对人', '行政法规', '规章'],
            '三国法': {
                '国际法': ['国际法', '国际条约', '国际组织', '国家主权', '领土', '外交', '领事', '战争', '和平'],
                '国际私法': ['国际私法', '冲突法', '法律适用', '涉外民事关系', '准据法', '法律选择'],
                '国际经济法': ['国际经济法', '国际贸易法', '国际投资法', 'WTO', '世界贸易组织', '关税', '反倾销']
            }
        }
    
    def extract_questions_from_docx(self, file_path):
        """从docx文件中提取题目"""
        try:
            doc = Document(file_path)
            questions = []
            current_question = ""
            question_num = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检测题目编号
                if re.match(r'^\d+[.、]', text):
                    if current_question:
                        questions.append({
                            'number': question_num,
                            'content': current_question.strip(),
                            'original_text': current_question.strip()
                        })
                    
                    question_num += 1
                    current_question = text
                else:
                    current_question += " " + text
            
            # 添加最后一个题目
            if current_question:
                questions.append({
                    'number': question_num,
                    'content': current_question.strip(),
                    'original_text': current_question.strip()
                })
            
            print(f"成功提取 {len(questions)} 道题目")
            return questions
            
        except Exception as e:
            print(f"读取文档失败: {e}")
            return []
    
    def classify_question(self, question_text):
        """对单个题目进行分类"""
        scores = defaultdict(float)
        
        # 对每个科目计算匹配分数
        for subject, keywords in self.subject_keywords.items():
            if isinstance(keywords, dict):
                # 多层级关键词
                for category, kw_list in keywords.items():
                    for keyword in kw_list:
                        if keyword in question_text:
                            scores[subject] += 1
            else:
                # 单层级关键词
                for keyword in keywords:
                    if keyword in question_text:
                        scores[subject] += 1
        
        # 特殊规则增强
        scores = self.apply_special_rules(question_text, scores)
        
        # 返回得分最高的科目
        if scores:
            return max(scores.items(), key=lambda x: x[1])[0]
        else:
            return '未分类'
    
    def apply_special_rules(self, question_text, scores):
        """应用特殊分类规则"""
        # 强化关键词匹配
        strong_indicators = {
            '民法': ['民事', '合同', '物权', '债权', '侵权', '婚姻', '继承', '民法典'],
            '刑法': ['犯罪', '刑罚', '犯罪构成', '故意', '过失', '刑法'],
            '刑事诉讼法': ['侦查', '起诉', '审判', '辩护', '证据', '刑事诉讼'],
            '民事诉讼法': ['起诉', '受理', '审理', '判决', '民事诉讼'],
            '行政法': ['行政', '行政行为', '行政处罚', '行政许可', '行政复议'],
            '商经知': ['公司', '合伙', '破产', '证券', '保险', '专利', '商标', '劳动'],
            '理论法': ['宪法', '法理', '法治', '司法制度', '法制史'],
            '三国法': ['国际法', '国际条约', '涉外', '国际私法', '国际经济法']
        }
        
        for subject, keywords in strong_indicators.items():
            for keyword in keywords:
                if keyword in question_text:
                    scores[subject] += 2  # 强化匹配得分
        
        return scores
    
    def classify_all_questions(self, questions):
        """对所有题目进行分类"""
        classified = {subject: [] for subject in self.official_ratios.keys()}
        classified['未分类'] = []
        
        for question in questions:
            subject = self.classify_question(question['content'])
            classified[subject].append(question)
        
        return classified
    
    def balance_classification(self, classified_questions, total_questions=200):
        """根据官方比例平衡分类结果"""
        # 计算目标数量
        target_counts = {}
        for subject, ratio in self.official_ratios.items():
            target_counts[subject] = int(total_questions * ratio)
        
        print("\n=== 分类平衡调整 ===")
        print("官方目标分布:")
        for subject, count in target_counts.items():
            current_count = len(classified_questions[subject])
            print(f"{subject}: 目标{count}题, 当前{current_count}题, 差值{count-current_count}")
        
        # 简单的重新分配逻辑（可以进一步优化）
        balanced_result = {}
        for subject in self.official_ratios.keys():
            balanced_result[subject] = classified_questions[subject][:target_counts[subject]]
        
        return balanced_result
    
    def generate_reports(self, classified_questions, output_dir="docx_2023"):
        """生成分类报告和各科目文档"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 统计信息
        total_questions = sum(len(questions) for questions in classified_questions.values())
        stats = {}
        
        print(f"\n=== 2023年法考题目分类结果 ===")
        print(f"总题目数: {total_questions}")
        print(f"{'科目':<15} {'题目数':<8} {'实际比例':<10} {'官方比例':<10} {'差值'}")
        print("-" * 60)
        
        for subject, questions in classified_questions.items():
            count = len(questions)
            actual_ratio = count / total_questions if total_questions > 0 else 0
            official_ratio = self.official_ratios.get(subject, 0)
            difference = actual_ratio - official_ratio
            
            stats[subject] = {
                'count': count,
                'actual_ratio': actual_ratio,
                'official_ratio': official_ratio,
                'difference': difference
            }
            
            print(f"{subject:<15} {count:<8} {actual_ratio:.2%} {official_ratio:.2%} {difference:+.2%}")
            
            # 生成各科目文档
            if count > 0:
                self.create_subject_document(subject, questions, f"{output_dir}/{subject}.docx")
        
        # 生成统计报告
        self.create_statistics_report(stats, total_questions, f"{output_dir}/分类统计报告.json")
        
        return stats
    
    def create_subject_document(self, subject, questions, file_path):
        """创建科目文档"""
        try:
            doc = Document()
            doc.add_heading(f'2023年法考 - {subject}', 0)
            doc.add_paragraph(f'题目总数: {len(questions)}')
            doc.add_paragraph('')
            
            for i, question in enumerate(questions, 1):
                doc.add_paragraph(f"{i}. {question['content']}")
                doc.add_paragraph('')
            
            doc.save(file_path)
            print(f"已生成: {file_path}")
            
        except Exception as e:
            print(f"生成文档失败 {file_path}: {e}")
    
    def create_statistics_report(self, stats, total_questions, file_path):
        """创建统计报告"""
        report = {
            'total_questions': total_questions,
            'classification_date': '2025-06-26',
            'official_ratios': self.official_ratios,
            'classification_results': stats,
            'summary': {
                'largest_deviation': max(stats.values(), key=lambda x: abs(x['difference']))['difference'] if stats else 0,
                'most_accurate_subject': min(stats.items(), key=lambda x: abs(x[1]['difference']))[0] if stats else None
            }
        }
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        print(f"已生成统计报告: {file_path}")


def main():
    """主函数"""
    classifier = LawExamClassifier()
    
    # 文档路径
    docx_path = "23年法考题.docx"
    
    if not os.path.exists(docx_path):
        print(f"错误: 找不到文档 {docx_path}")
        return
    
    print("开始分析2023年法考题目...")
    
    # 1. 提取题目
    questions = classifier.extract_questions_from_docx(docx_path)
    if not questions:
        print("未能提取到题目，程序退出")
        return
    
    # 2. 分类题目
    classified = classifier.classify_all_questions(questions)
    
    # 3. 平衡分类（可选）
    # balanced = classifier.balance_classification(classified, len(questions))
    
    # 4. 生成报告
    stats = classifier.generate_reports(classified)
    
    print(f"\n分类完成！共处理 {len(questions)} 道题目")
    print("生成的文件:")
    print("- docx_2023/各科目.docx")
    print("- docx_2023/分类统计报告.json")


if __name__ == "__main__":
    main()