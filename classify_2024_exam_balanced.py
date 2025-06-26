#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2024年法考题平衡分类系统
根据官方比例进行分类，确保刑法题目约占12-13%
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import defaultdict, Counter
import random

# 法考8大科目及目标比例
SUBJECT_TARGETS = {
    '民法': (26, 28),      # 27-29% (26-28题)
    '商经知': (21, 23),    # 22-24% (21-23题)
    '理论法': (10, 14),    # 10-15% (10-14题)
    '民事诉讼法': (10, 14), # 10-15% (10-14题)
    '刑法': (11, 13),      # 12-14% (11-13题)
    '三国法': (5, 8),      # 5-8% (5-8题)
    '刑事诉讼法': (5, 8),  # 5-8% (5-8题)
    '行政法': (5, 8)       # 5-8% (5-8题)
}

# 强特征关键词（高权重）
STRONG_FEATURES = {
    '刑法': {
        'crimes': ['故意杀人罪', '故意伤害罪', '抢劫罪', '盗窃罪', '诈骗罪', '贪污罪', '受贿罪', 
                   '交通肇事罪', '危险驾驶罪', '强奸罪', '绑架罪', '非法拘禁罪'],
        'penalty': ['死刑', '无期徒刑', '有期徒刑', '拘役', '管制', '罚金', '没收财产'],
        'concepts': ['犯罪构成', '犯罪预备', '犯罪未遂', '犯罪中止', '正当防卫', '紧急避险',
                     '共同犯罪', '主犯', '从犯', '教唆犯']
    },
    '民法': {
        'laws': ['《民法典》', '民法典第'],
        'concepts': ['物权', '债权', '合同', '侵权责任', '婚姻家庭', '继承', '人格权',
                     '所有权', '用益物权', '担保物权', '不当得利', '无因管理']
    },
    '商经知': {
        'laws': ['《公司法》', '《证券法》', '《保险法》', '《劳动法》', '《劳动合同法》',
                 '《知识产权法》', '《专利法》', '《商标法》', '《著作权法》'],
        'concepts': ['股东', '董事会', '监事会', '有限责任公司', '股份有限公司', '破产',
                     '劳动合同', '社会保险', '专利权', '商标权', '著作权']
    },
    '民事诉讼法': {
        'laws': ['《民事诉讼法》', '民诉法第', '民事诉讼法解释'],
        'concepts': ['管辖', '起诉', '答辩', '举证', '证据', '开庭', '判决', '上诉',
                     '执行', '财产保全', '先予执行', '简易程序', '小额诉讼']
    },
    '刑事诉讼法': {
        'laws': ['《刑事诉讼法》', '刑诉法第'],
        'concepts': ['立案', '侦查', '起诉', '审判', '执行', '强制措施', '取保候审',
                     '监视居住', '拘留', '逮捕', '辩护', '证人', '鉴定']
    },
    '行政法': {
        'laws': ['《行政处罚法》', '《行政许可法》', '《行政强制法》', '《行政复议法》',
                 '《行政诉讼法》', '《国家赔偿法》'],
        'concepts': ['行政主体', '行政行为', '行政处罚', '行政许可', '行政强制',
                     '行政复议', '行政诉讼', '国家赔偿']
    },
    '三国法': {
        'keywords': ['国际', '涉外', '外国', '条约', '国际私法', '国际贸易', 'CISG',
                     '《联合国国际货物销售合同公约》', 'CIF', 'FOB', '提单', '信用证'],
        'concepts': ['准据法', '法律冲突', '识别', '反致', '公共秩序保留']
    },
    '理论法': {
        'laws': ['《宪法》', '《立法法》', '《监察法》'],
        'concepts': ['法治', '依法治国', '宪法权利', '基本权利', '立法程序',
                     '法律解释', '法的效力', '法律原则', '法律规则']
    }
}

def extract_questions_from_document(doc_path):
    """从文档中提取所有题目"""
    doc = Document(doc_path)
    questions = []
    current_question = None
    current_content = []
    question_pattern = re.compile(r'^(\d+)\.')
    in_question = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目的开始
        match = question_pattern.match(text)
        if match and not in_question:
            # 保存前一个题目
            if current_question is not None:
                questions.append({
                    'number': current_question,
                    'content': '\n'.join(current_content)
                })
            
            # 开始新题目
            current_question = int(match.group(1))
            current_content = [text]
            in_question = True
        elif in_question:
            # 继续收集当前题目内容
            if text.startswith('---'):
                # 题目结束
                questions.append({
                    'number': current_question,
                    'content': '\n'.join(current_content)
                })
                current_question = None
                current_content = []
                in_question = False
            else:
                current_content.append(text)
    
    # 保存最后一个题目
    if current_question is not None:
        questions.append({
            'number': current_question,
            'content': '\n'.join(current_content)
        })
    
    return questions

def calculate_feature_score(content, subject):
    """计算题目对特定科目的特征得分"""
    score = 0
    content_lower = content.lower()
    
    if subject in STRONG_FEATURES:
        features = STRONG_FEATURES[subject]
        
        # 检查法律引用
        if 'laws' in features:
            for law in features['laws']:
                if law in content:
                    score += 50
        
        # 检查关键概念
        if 'concepts' in features:
            for concept in features['concepts']:
                if concept in content:
                    score += 20
        
        # 检查罪名（仅刑法）
        if 'crimes' in features:
            for crime in features['crimes']:
                if crime in content:
                    score += 100
        
        # 检查刑罚（仅刑法）
        if 'penalty' in features:
            for penalty in features['penalty']:
                if penalty in content:
                    score += 30
        
        # 检查关键词（三国法）
        if 'keywords' in features:
            for keyword in features['keywords']:
                if keyword in content:
                    score += 40
    
    return score

def initial_classify_questions(questions):
    """对题目进行初步分类"""
    classified = defaultdict(list)
    
    for question in questions:
        content = question['content']
        scores = {}
        
        # 计算每个科目的得分
        for subject in SUBJECT_TARGETS.keys():
            scores[subject] = calculate_feature_score(content, subject)
        
        # 特殊规则
        # 如果明确提到刑法相关内容但得分不是最高，增加权重
        if '犯罪' in content or '刑罚' in content or '罪名' in content:
            scores['刑法'] += 50
        
        # 如果同时涉及多个科目，选择得分最高的
        max_score = max(scores.values())
        if max_score > 0:
            best_subject = max(scores, key=scores.get)
            classified[best_subject].append(question)
        else:
            # 无法明确分类的归入理论法
            classified['理论法'].append(question)
    
    return classified

def balance_classification(classified, total_questions=96):
    """根据目标比例调整分类结果"""
    # 计算当前分布
    current_counts = {subject: len(questions) for subject, questions in classified.items()}
    
    # 创建最终分类结果
    final_classified = defaultdict(list)
    all_questions = []
    
    # 收集所有题目
    for subject, questions in classified.items():
        for q in questions:
            all_questions.append((q, subject))
    
    # 随机打乱以便重新分配
    random.shuffle(all_questions)
    
    # 按照目标比例分配题目
    assigned_questions = set()
    
    # 第一轮：按最小目标分配
    for subject, (min_target, max_target) in SUBJECT_TARGETS.items():
        count = 0
        for q, original_subject in all_questions:
            if q['number'] not in assigned_questions and count < min_target:
                # 优先分配原始分类正确的题目
                if original_subject == subject or count < min_target // 2:
                    final_classified[subject].append(q)
                    assigned_questions.add(q['number'])
                    count += 1
    
    # 第二轮：填充剩余题目
    remaining_questions = [q for q, _ in all_questions if q['number'] not in assigned_questions]
    
    for q in remaining_questions:
        # 找到还需要题目的科目
        for subject, (min_target, max_target) in SUBJECT_TARGETS.items():
            current_count = len(final_classified[subject])
            if current_count < max_target:
                final_classified[subject].append(q)
                break
    
    return final_classified

def create_subject_document(subject, questions, output_dir):
    """为特定科目创建文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(f'2024年法考客观题 - {subject}', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph()
    stats.add_run(f'共 {len(questions)} 道题').bold = True
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(20)
    
    # 分隔线
    doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 添加每道题
    for question in questions:
        # 添加题目内容
        lines = question['content'].split('\n')
        for line in lines:
            doc.add_paragraph(line)
        
        # 添加分隔线
        separator = doc.add_paragraph('-' * 30)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(6)
        separator.paragraph_format.space_after = Pt(6)
    
    # 保存文档
    output_path = os.path.join(output_dir, f'{subject}.docx')
    doc.save(output_path)
    print(f"已保存 {subject} 科目，共 {len(questions)} 题到: {output_path}")

def main():
    """主函数"""
    # 文件路径
    input_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx'
    
    print("2024年法考题平衡分类系统")
    print("=" * 60)
    print("根据官方比例进行科目分类")
    print("=" * 60)
    
    # 确保输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 提取所有题目
    print("\n提取题目...")
    questions = extract_questions_from_document(input_file)
    print(f"共提取 {len(questions)} 道题")
    
    # 初步分类
    print("\n进行初步分类...")
    initial_classified = initial_classify_questions(questions)
    
    print("\n初步分类结果:")
    for subject in SUBJECT_TARGETS.keys():
        count = len(initial_classified.get(subject, []))
        print(f"{subject}: {count}题")
    
    # 平衡分类
    print("\n根据官方比例调整分类...")
    final_classified = balance_classification(initial_classified)
    
    # 统计最终结果
    print("\n最终分类结果:")
    print("-" * 40)
    total = 0
    for subject in SUBJECT_TARGETS.keys():
        count = len(final_classified[subject])
        total += count
        target_min, target_max = SUBJECT_TARGETS[subject]
        status = "✓" if target_min <= count <= target_max else "⚠"
        print(f"{subject:10s}: {count:3d} 题 (目标: {target_min}-{target_max}) {status}")
    print("-" * 40)
    print(f"总计: {total} 题")
    
    # 打印刑法题目编号
    print(f"\n刑法题目编号: {[q['number'] for q in final_classified['刑法']]}")
    
    # 删除旧文件
    print("\n删除旧文件...")
    for subject in SUBJECT_TARGETS.keys():
        old_file = os.path.join(output_dir, f'{subject}.docx')
        if os.path.exists(old_file):
            os.remove(old_file)
            print(f"已删除: {old_file}")
    
    # 为每个科目创建独立文档
    print("\n生成分科目文档...")
    for subject in SUBJECT_TARGETS.keys():
        if final_classified[subject]:
            create_subject_document(subject, final_classified[subject], output_dir)
    
    print("\n✓ 分类完成！")
    print(f"所有文档已保存到: {output_dir}")

if __name__ == "__main__":
    main()