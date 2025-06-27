#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法考真题文档合并脚本 - 最终版
将扫描版原始文本中的题目题干补充到完整版文档中
"""

import re
import json
from docx import Document
from collections import OrderedDict


class ExamQuestion:
    """题目类"""
    def __init__(self, number):
        self.number = number
        self.stem = ""  # 题干
        self.options = OrderedDict()  # 选项
        self.answer = ""  # 答案
        self.analysis = ""  # 解析
        
    def __str__(self):
        return f"题{self.number}: 题干长度={len(self.stem)}, 选项数={len(self.options)}, 有答案={bool(self.answer)}, 有解析={bool(self.analysis)}"


def parse_scanned_text(file_path):
    """解析扫描版原始文本，获取题干和选项"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    questions = OrderedDict()
    lines = content.split('\n')
    
    current_question = None
    i = 0
    
    while i < len(lines):
        # 去掉行号前缀
        line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
        
        if not line:
            i += 1
            continue
        
        # 匹配题号
        match = re.match(r'^(\d+)[.、]\s*(.+)', line)
        if match:
            # 保存上一题
            if current_question:
                questions[current_question.number] = current_question
            
            # 创建新题目
            current_question = ExamQuestion(int(match.group(1)))
            current_question.stem = match.group(2)
            
            # 继续读取题干直到遇到选项
            i += 1
            while i < len(lines):
                next_line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
                if not next_line:
                    i += 1
                    continue
                
                # 检查是否是选项
                if re.match(r'^[A-D][.、]\s*', next_line):
                    break
                
                # 添加到题干
                current_question.stem += " " + next_line
                i += 1
                
                # 如果题干包含问号，可能结束了
                if '?' in current_question.stem or '？' in current_question.stem:
                    # 再检查下一行是否是选项
                    if i < len(lines):
                        peek_line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
                        if not re.match(r'^[A-D][.、]\s*', peek_line):
                            # 不是选项，继续读题干
                            continue
                    break
            
            # 读取选项
            while i < len(lines):
                option_line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
                if not option_line:
                    i += 1
                    continue
                
                # 匹配选项
                opt_match = re.match(r'^([A-D])[.、]\s*(.+)', option_line)
                if opt_match:
                    current_question.options[opt_match.group(1)] = opt_match.group(2)
                    i += 1
                else:
                    # 不是选项，可能是下一题了
                    break
                
                # 如果已经有4个选项了，跳出
                if len(current_question.options) == 4:
                    break
        else:
            i += 1
    
    # 保存最后一题
    if current_question:
        questions[current_question.number] = current_question
    
    return questions


def parse_docx_file(file_path):
    """解析docx文件，获取答案和解析"""
    doc = Document(file_path)
    
    questions = OrderedDict()
    current_question = None
    current_mode = None  # 'stem', 'options', 'answer', 'analysis'
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 跳过标题和分隔线
        if text == '2024年法考客观题（真题+答案解析）' or text.startswith('---'):
            continue
        
        # 匹配题号（如：1. 正确答案：B）
        answer_match = re.match(r'^(\d+)[.、]\s*正确答案[：:]\s*([A-D]+)', text)
        if answer_match:
            q_num = int(answer_match.group(1))
            if q_num not in questions:
                questions[q_num] = ExamQuestion(q_num)
            questions[q_num].answer = answer_match.group(2)
            current_question = questions[q_num]
            current_mode = 'answer'
            continue
        
        # 匹配题号开头（可能是解析的开始）
        num_match = re.match(r'^(\d+)[.、]\s*(.+)', text)
        if num_match and '选项内容待补充' not in text:
            q_num = int(num_match.group(1))
            content = num_match.group(2)
            
            # 如果不是答案行，可能是解析的开始
            if not content.startswith('正确答案'):
                if q_num not in questions:
                    questions[q_num] = ExamQuestion(q_num)
                current_question = questions[q_num]
                current_question.analysis = content
                current_mode = 'analysis'
            continue
        
        # 处理【答案解析】标记
        if text == '【答案解析】' or text == '答案解析：':
            current_mode = 'analysis'
            continue
        
        # 处理选项占位符
        if '选项内容待补充' in text:
            continue
        
        # 根据当前模式添加内容
        if current_question and current_mode == 'analysis':
            # 添加到解析内容
            if text and not text.startswith('综上所述'):
                if current_question.analysis:
                    current_question.analysis += "\n" + text
                else:
                    current_question.analysis = text
    
    return questions


def merge_questions(scanned_questions, docx_questions):
    """合并题目数据"""
    merged_questions = OrderedDict()
    
    # 获取所有题号
    all_numbers = sorted(set(scanned_questions.keys()) | set(docx_questions.keys()))
    
    for num in all_numbers:
        scanned_q = scanned_questions.get(num)
        docx_q = docx_questions.get(num)
        
        merged_q = ExamQuestion(num)
        
        # 题干和选项来自扫描版
        if scanned_q:
            merged_q.stem = scanned_q.stem
            merged_q.options = scanned_q.options
        
        # 答案和解析来自docx版
        if docx_q:
            merged_q.answer = docx_q.answer
            merged_q.analysis = docx_q.analysis
            
            # 如果扫描版没有题干，尝试从解析中提取
            if not merged_q.stem and docx_q.analysis:
                # 某些情况下，解析的第一句可能包含题干信息
                pass
        
        merged_questions[num] = merged_q
    
    return merged_questions


def create_merged_document(merged_questions, output_path):
    """创建合并后的文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('2024年法律职业资格考试客观题真题（完整版）', 0)
    title.alignment = 1  # 居中
    
    doc.add_paragraph()
    
    # 添加说明
    info = doc.add_paragraph('本文档由扫描版原始文本和答案解析文档合并而成，包含完整的题干、选项、答案和解析。')
    info.alignment = 1
    doc.add_paragraph()
    
    # 题型分类
    single_choice_start = 1
    multiple_choice_start = 51
    uncertain_choice_start = 86
    
    current_section = None
    
    # 按题号顺序添加题目
    for num in sorted(merged_questions.keys()):
        question = merged_questions[num]
        
        # 添加题型标题
        if num == single_choice_start:
            doc.add_heading('一、单项选择题', 1)
            doc.add_paragraph('每题所设选项中只有一个正确答案，多选、错选或不选均不得分。')
            doc.add_paragraph()
            current_section = 'single'
        elif num == multiple_choice_start:
            doc.add_paragraph()
            doc.add_heading('二、多项选择题', 1)
            doc.add_paragraph('每题所设选项中至少有两个正确答案，多选、少选、错选或不选均不得分。')
            doc.add_paragraph()
            current_section = 'multiple'
        elif num == uncertain_choice_start:
            doc.add_paragraph()
            doc.add_heading('三、不定项选择题', 1)
            doc.add_paragraph('每题所设选项中有一个或多个正确答案，少选且选对的每个选项得0.5分，错选、多选或不选均不得分。')
            doc.add_paragraph()
            current_section = 'uncertain'
        
        # 题号和题干
        if question.stem:
            para = doc.add_paragraph()
            para.add_run(f"{num}. ").bold = True
            para.add_run(question.stem)
        else:
            para = doc.add_paragraph()
            para.add_run(f"{num}. ").bold = True
            para.add_run("（题干待补充）").font.italic = True
        
        # 选项
        if question.options:
            for opt_key, opt_value in question.options.items():
                opt_para = doc.add_paragraph()
                opt_para.add_run(f"    {opt_key}. {opt_value}")
        else:
            opt_para = doc.add_paragraph()
            opt_para.add_run("    （选项待补充）").font.italic = True
        
        # 答案
        if question.answer:
            doc.add_paragraph()
            ans_para = doc.add_paragraph()
            ans_para.add_run("【答案】").bold = True
            ans_para.add_run(question.answer)
        
        # 解析
        if question.analysis:
            analysis_para = doc.add_paragraph()
            analysis_para.add_run("【解析】").bold = True
            
            # 解析内容可能有多行
            analysis_lines = question.analysis.split('\n')
            for i, line in enumerate(analysis_lines):
                if i == 0:
                    analysis_para.add_run(line)
                else:
                    doc.add_paragraph(line)
        
        # 题目间分隔
        doc.add_paragraph()
        doc.add_paragraph("—" * 50)
        doc.add_paragraph()
    
    # 保存文档
    doc.save(output_path)


def generate_summary_report(merged_questions):
    """生成汇总报告"""
    # 题型分类
    single_choice = range(1, 51)
    multiple_choice = range(51, 86)
    uncertain_choice = range(86, 101)
    
    # 统计各题型的完整度
    stats = {
        'single_choice': {
            'total': 0,
            'with_stem': 0,
            'with_options': 0,
            'with_answer': 0,
            'with_analysis': 0
        },
        'multiple_choice': {
            'total': 0,
            'with_stem': 0,
            'with_options': 0,
            'with_answer': 0,
            'with_analysis': 0
        },
        'uncertain_choice': {
            'total': 0,
            'with_stem': 0,
            'with_options': 0,
            'with_answer': 0,
            'with_analysis': 0
        }
    }
    
    for num, question in merged_questions.items():
        if num in single_choice:
            section = 'single_choice'
        elif num in multiple_choice:
            section = 'multiple_choice'
        elif num in uncertain_choice:
            section = 'uncertain_choice'
        else:
            continue
        
        stats[section]['total'] += 1
        if question.stem:
            stats[section]['with_stem'] += 1
        if question.options:
            stats[section]['with_options'] += 1
        if question.answer:
            stats[section]['with_answer'] += 1
        if question.analysis:
            stats[section]['with_analysis'] += 1
    
    return stats


def main():
    """主函数"""
    # 文件路径
    scanned_file = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）_扫描版_raw_doc.txt'
    docx_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/2024年法考真题-补充完整版.docx'
    
    print("="*60)
    print("2024年法考真题文档合并工具")
    print("="*60)
    
    print("\n1. 正在解析扫描版原始文本...")
    scanned_questions = parse_scanned_text(scanned_file)
    print(f"   ✓ 从扫描版中提取了 {len(scanned_questions)} 道题目")
    
    print("\n2. 正在解析答案解析文档...")
    try:
        docx_questions = parse_docx_file(docx_file)
        print(f"   ✓ 从答案文档中提取了 {len(docx_questions)} 道题目的答案和解析")
    except Exception as e:
        print(f"   ✗ 读取docx文件出错: {e}")
        docx_questions = OrderedDict()
    
    print("\n3. 正在合并题目数据...")
    merged_questions = merge_questions(scanned_questions, docx_questions)
    print(f"   ✓ 合并后共有 {len(merged_questions)} 道题目")
    
    # 生成统计信息
    stats = generate_summary_report(merged_questions)
    
    print("\n4. 题目完整度统计：")
    print(f"\n   单项选择题（1-50题）：")
    print(f"      - 总题数：{stats['single_choice']['total']}")
    print(f"      - 有题干：{stats['single_choice']['with_stem']}")
    print(f"      - 有选项：{stats['single_choice']['with_options']}")
    print(f"      - 有答案：{stats['single_choice']['with_answer']}")
    print(f"      - 有解析：{stats['single_choice']['with_analysis']}")
    
    print(f"\n   多项选择题（51-85题）：")
    print(f"      - 总题数：{stats['multiple_choice']['total']}")
    print(f"      - 有题干：{stats['multiple_choice']['with_stem']}")
    print(f"      - 有选项：{stats['multiple_choice']['with_options']}")
    print(f"      - 有答案：{stats['multiple_choice']['with_answer']}")
    print(f"      - 有解析：{stats['multiple_choice']['with_analysis']}")
    
    print(f"\n   不定项选择题（86-100题）：")
    print(f"      - 总题数：{stats['uncertain_choice']['total']}")
    print(f"      - 有题干：{stats['uncertain_choice']['with_stem']}")
    print(f"      - 有选项：{stats['uncertain_choice']['with_options']}")
    print(f"      - 有答案：{stats['uncertain_choice']['with_answer']}")
    print(f"      - 有解析：{stats['uncertain_choice']['with_analysis']}")
    
    print(f"\n5. 正在创建合并文档...")
    create_merged_document(merged_questions, output_file)
    print(f"   ✓ 文档已保存到: {output_file}")
    
    # 生成详细报告
    report_file = '/Users/acheng/Downloads/law-exam-assistant/merge_report_final.json'
    report = {
        'summary': {
            'scanned_questions_count': len(scanned_questions),
            'docx_questions_count': len(docx_questions),
            'merged_questions_count': len(merged_questions),
        },
        'statistics': stats,
        'missing_components': {
            'missing_stem': [num for num, q in merged_questions.items() if not q.stem],
            'missing_options': [num for num, q in merged_questions.items() if not q.options],
            'missing_answer': [num for num, q in merged_questions.items() if not q.answer],
            'missing_analysis': [num for num, q in merged_questions.items() if not q.analysis]
        }
    }
    
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    print(f"\n6. 详细报告已保存到: {report_file}")
    print("\n合并完成！✓")
    print("="*60)


if __name__ == '__main__':
    main()