#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法考真题文档合并脚本 V2
将扫描版原始文本中的题目题干补充到完整版文档中
"""

import re
import json
from docx import Document
from collections import OrderedDict


class ExamQuestion:
    """题目类"""
    def __init__(self, number):
        self.number = number
        self.stem = ""  # 题干
        self.options = OrderedDict()  # 选项
        self.answer = ""  # 答案
        self.analysis = ""  # 解析
        
    def to_text(self):
        """转换为文本格式"""
        text = f"{self.number}. {self.stem}\n"
        for opt_key, opt_value in self.options.items():
            text += f"{opt_key}. {opt_value}\n"
        if self.answer:
            text += f"答案：{self.answer}\n"
        if self.analysis:
            text += f"解析：{self.analysis}\n"
        return text


def parse_scanned_text(file_path):
    """解析扫描版原始文本"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    questions = OrderedDict()
    lines = content.split('\n')
    
    current_question = None
    i = 0
    
    while i < len(lines):
        # 去掉行号前缀
        line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
        
        if not line:
            i += 1
            continue
        
        # 匹配题号
        match = re.match(r'^(\d+)[.、]\s*(.+)', line)
        if match:
            # 保存上一题
            if current_question:
                questions[current_question.number] = current_question
            
            # 创建新题目
            current_question = ExamQuestion(int(match.group(1)))
            current_question.stem = match.group(2)
            
            # 继续读取题干直到遇到选项或问号
            i += 1
            while i < len(lines):
                next_line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
                if not next_line:
                    i += 1
                    continue
                
                # 检查是否是选项
                if re.match(r'^[A-D][.、]\s*', next_line):
                    break
                
                # 添加到题干
                current_question.stem += " " + next_line
                i += 1
                
                # 如果题干包含问号，可能结束了
                if '?' in current_question.stem or '？' in current_question.stem:
                    break
            
            # 读取选项
            while i < len(lines):
                option_line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
                if not option_line:
                    i += 1
                    continue
                
                # 匹配选项
                opt_match = re.match(r'^([A-D])[.、]\s*(.+)', option_line)
                if opt_match:
                    current_question.options[opt_match.group(1)] = opt_match.group(2)
                    i += 1
                else:
                    # 不是选项，可能是下一题了
                    break
        else:
            i += 1
    
    # 保存最后一题
    if current_question:
        questions[current_question.number] = current_question
    
    return questions


def parse_docx_file(file_path):
    """解析docx文件"""
    doc = Document(file_path)
    
    questions = OrderedDict()
    current_question = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 匹配题号
        match = re.match(r'^(\d+)[.、]\s*(.+)?', text)
        if match:
            # 保存上一题
            if current_question:
                questions[current_question.number] = current_question
            
            # 创建新题目
            current_question = ExamQuestion(int(match.group(1)))
            if match.group(2):
                # 检查是否直接是选项
                opt_match = re.match(r'^([A-D])[.、]\s*(.+)', match.group(2))
                if opt_match:
                    current_question.options[opt_match.group(1)] = opt_match.group(2)
                else:
                    current_question.stem = match.group(2)
        elif current_question:
            # 处理当前题目的内容
            # 匹配选项
            opt_match = re.match(r'^([A-D])[.、]\s*(.+)', text)
            if opt_match:
                current_question.options[opt_match.group(1)] = opt_match.group(2)
            # 匹配答案
            elif text.startswith('答案：') or text.startswith('答案:'):
                current_question.answer = text.replace('答案：', '').replace('答案:', '').strip()
            # 匹配解析
            elif text.startswith('解析：') or text.startswith('解析:'):
                current_question.analysis = text.replace('解析：', '').replace('解析:', '').strip()
            else:
                # 可能是题干的延续或解析的延续
                if not current_question.options and not current_question.answer:
                    current_question.stem += " " + text
                elif current_question.answer and not current_question.analysis:
                    pass  # 答案后面，解析前面的内容忽略
                else:
                    current_question.analysis += " " + text
    
    # 保存最后一题
    if current_question:
        questions[current_question.number] = current_question
    
    return questions


def merge_questions(scanned_questions, docx_questions):
    """合并题目，优先使用扫描版的题干"""
    merged_questions = OrderedDict()
    
    # 获取所有题号
    all_numbers = sorted(set(scanned_questions.keys()) | set(docx_questions.keys()))
    
    for num in all_numbers:
        scanned_q = scanned_questions.get(num)
        docx_q = docx_questions.get(num)
        
        if scanned_q and docx_q:
            # 两个版本都有，合并
            merged_q = ExamQuestion(num)
            
            # 优先使用扫描版的题干（更完整）
            merged_q.stem = scanned_q.stem if scanned_q.stem else docx_q.stem
            
            # 选项：如果扫描版有就用扫描版的，否则用docx版的
            if scanned_q.options:
                merged_q.options = scanned_q.options
            else:
                merged_q.options = docx_q.options
            
            # 答案和解析使用docx版的（如果有）
            merged_q.answer = docx_q.answer if docx_q.answer else scanned_q.answer
            merged_q.analysis = docx_q.analysis if docx_q.analysis else scanned_q.analysis
            
            merged_questions[num] = merged_q
        elif scanned_q:
            # 只有扫描版
            merged_questions[num] = scanned_q
        elif docx_q:
            # 只有docx版
            merged_questions[num] = docx_q
    
    return merged_questions


def create_merged_document(merged_questions, output_path):
    """创建合并后的文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('2024年法律职业资格考试客观题真题（完整版）', 0)
    title.alignment = 1  # 居中
    
    doc.add_paragraph()
    
    # 添加说明
    doc.add_paragraph('说明：本文档由扫描版原始文本和完整版文档合并而成，确保每道题都有完整的题干、选项、答案和解析。')
    doc.add_paragraph()
    
    # 分类标题
    doc.add_heading('一、单项选择题', 1)
    doc.add_paragraph()
    
    # 按题号顺序添加题目
    for num in sorted(merged_questions.keys()):
        question = merged_questions[num]
        
        # 题号和题干
        para = doc.add_paragraph()
        para.add_run(f"{num}. ").bold = True
        para.add_run(question.stem)
        
        # 选项
        for opt_key, opt_value in question.options.items():
            opt_para = doc.add_paragraph(style='List Bullet')
            opt_para.add_run(f"{opt_key}. {opt_value}")
        
        # 答案
        if question.answer:
            ans_para = doc.add_paragraph()
            ans_para.add_run("答案：").bold = True
            ans_para.add_run(question.answer)
        
        # 解析
        if question.analysis:
            analysis_para = doc.add_paragraph()
            analysis_para.add_run("解析：").bold = True
            analysis_para.add_run(question.analysis)
        
        # 题目间空行
        doc.add_paragraph()
        
        # 多选题分界线
        if num == 50:
            doc.add_paragraph()
            doc.add_heading('二、多项选择题', 1)
            doc.add_paragraph()
        
        # 不定项选择题分界线
        if num == 85:
            doc.add_paragraph()
            doc.add_heading('三、不定项选择题', 1)
            doc.add_paragraph()
    
    # 保存文档
    doc.save(output_path)


def main():
    """主函数"""
    # 文件路径
    scanned_file = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）_扫描版_raw_doc.txt'
    docx_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/2024年法考真题-补充完整版.docx'
    
    print("正在解析扫描版原始文本...")
    scanned_questions = parse_scanned_text(scanned_file)
    print(f"从扫描版中提取了 {len(scanned_questions)} 道题目")
    
    print("\n正在解析完整版文档...")
    try:
        docx_questions = parse_docx_file(docx_file)
        print(f"从完整版中读取了 {len(docx_questions)} 道题目")
    except Exception as e:
        print(f"读取docx文件出错: {e}")
        docx_questions = OrderedDict()
    
    print("\n正在合并题目...")
    merged_questions = merge_questions(scanned_questions, docx_questions)
    print(f"合并后共有 {len(merged_questions)} 道题目")
    
    # 统计信息
    questions_with_stem = sum(1 for q in merged_questions.values() if q.stem)
    questions_with_options = sum(1 for q in merged_questions.values() if q.options)
    questions_with_answer = sum(1 for q in merged_questions.values() if q.answer)
    questions_with_analysis = sum(1 for q in merged_questions.values() if q.analysis)
    
    print(f"\n统计信息：")
    print(f"  有题干的题目：{questions_with_stem}")
    print(f"  有选项的题目：{questions_with_options}")
    print(f"  有答案的题目：{questions_with_answer}")
    print(f"  有解析的题目：{questions_with_analysis}")
    
    print(f"\n正在创建合并文档...")
    create_merged_document(merged_questions, output_file)
    print(f"文档已保存到: {output_file}")
    
    # 生成详细报告
    report_file = '/Users/acheng/Downloads/law-exam-assistant/merge_report_v2.json'
    report = {
        'scanned_questions_count': len(scanned_questions),
        'docx_questions_count': len(docx_questions),
        'merged_questions_count': len(merged_questions),
        'statistics': {
            'questions_with_stem': questions_with_stem,
            'questions_with_options': questions_with_options,
            'questions_with_answer': questions_with_answer,
            'questions_with_analysis': questions_with_analysis
        },
        'missing_components': {
            'missing_stem': [num for num, q in merged_questions.items() if not q.stem],
            'missing_options': [num for num, q in merged_questions.items() if not q.options],
            'missing_answer': [num for num, q in merged_questions.items() if not q.answer],
            'missing_analysis': [num for num, q in merged_questions.items() if not q.analysis]
        }
    }
    
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    print(f"\n详细报告已保存到: {report_file}")


if __name__ == '__main__':
    main()