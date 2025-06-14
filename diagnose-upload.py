#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys
import os
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import hashlib

def diagnose_file(filepath):
    """诊断上传的文件问题"""
    print(f"诊断文件: {filepath}")
    print(f"文件是否存在: {os.path.exists(filepath)}")
    
    if os.path.exists(filepath):
        # 获取文件信息
        file_size = os.path.getsize(filepath)
        print(f"文件大小: {file_size} bytes")
        
        # 计算文件MD5
        with open(filepath, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()
        print(f"文件MD5: {file_hash}")
        
        # 检查文件头
        with open(filepath, 'rb') as f:
            header = f.read(4)
            print(f"文件头(前4字节): {header.hex()}")
            
            # DOCX文件应该以PK开头（ZIP格式）
            if header[:2] == b'PK':
                print("✓ 文件头正确（ZIP/DOCX格式）")
            else:
                print("✗ 文件头不正确，可能不是有效的DOCX文件")
        
        # 尝试打开文件
        try:
            doc = Document(filepath)
            print(f"✓ 成功打开文档")
            print(f"段落数量: {len(doc.paragraphs)}")
            
            # 显示前几个段落
            print("\n前5个段落内容:")
            for i, para in enumerate(doc.paragraphs[:5]):
                text = para.text.strip()
                if text:
                    print(f"  段落{i+1}: {text[:50]}...")
                    
        except PackageNotFoundError as e:
            print(f"✗ PackageNotFoundError: {e}")
            print("  文件可能已损坏或不是有效的DOCX文件")
        except Exception as e:
            print(f"✗ 其他错误: {type(e).__name__}: {e}")
    
    print("\n建议:")
    print("1. 确保文件是通过二进制模式正确保存的")
    print("2. 检查文件上传过程中是否有编码转换")
    print("3. 确保FormData正确处理了二进制文件")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python diagnose-upload.py <文件路径>")
        sys.exit(1)
    
    diagnose_file(sys.argv[1])