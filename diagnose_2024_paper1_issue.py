#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2024年卷一完整版文档问题诊断脚本
专门诊断为什么只能读取到4个题目的问题
"""

import os
import re
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
import json

def analyze_document_structure(file_path):
    """深度分析文档结构"""
    print(f"\n🔍 正在分析文档: {file_path}")
    print("=" * 60)
    
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return None
    
    try:
        doc = Document(file_path)
        print(f"✅ 文档加载成功")
        
        analysis = {
            'file_info': {
                'path': file_path,
                'size': os.path.getsize(file_path),
                'exists': True
            },
            'document_structure': {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'sections': len(doc.sections)
            },
            'content_analysis': {
                'question_patterns': [],
                'paragraph_samples': [],
                'potential_questions': [],
                'all_question_numbers': [],
                'problematic_patterns': []
            }
        }
        
        print(f"📊 基本信息:")
        print(f"   - 段落总数: {len(doc.paragraphs)}")
        print(f"   - 表格总数: {len(doc.tables)}")
        print(f"   - 节数: {len(doc.sections)}")
        print(f"   - 文件大小: {os.path.getsize(file_path):,} 字节")
        
        # 分析段落内容
        print(f"\n📝 段落内容分析:")
        question_count = 0
        question_numbers = []
        
        for i, para in enumerate(doc.paragraphs[:50]):  # 分析前50个段落
            text = para.text.strip()
            if not text:
                continue
                
            # 检查是否是题目
            question_match = re.match(r'^(\d+)[．.\s]*(.+)', text)
            if question_match:
                q_num = int(question_match.group(1))
                if 1 <= q_num <= 100:  # 题目编号范围
                    question_count += 1
                    question_numbers.append(q_num)
                    analysis['content_analysis']['potential_questions'].append({
                        'paragraph_index': i,
                        'question_number': q_num,
                        'text_preview': text[:100] + "..." if len(text) > 100 else text,
                        'full_text_length': len(text)
                    })
                    
                    if question_count <= 10:  # 显示前10题
                        print(f"   题目 {q_num}: {text[:80]}...")
            
            # 记录段落样本
            if i < 20:  # 前20个段落样本
                analysis['content_analysis']['paragraph_samples'].append({
                    'index': i,
                    'text_preview': text[:100] + "..." if len(text) > 100 else text,
                    'length': len(text),
                    'is_empty': len(text) == 0
                })
        
        analysis['content_analysis']['all_question_numbers'] = sorted(question_numbers)
        
        print(f"\n🎯 题目识别结果:")
        print(f"   - 通过基础模式识别的题目数: {question_count}")
        print(f"   - 题目编号范围: {min(question_numbers) if question_numbers else 'N/A'} - {max(question_numbers) if question_numbers else 'N/A'}")
        
        if question_numbers:
            missing_numbers = []
            for i in range(1, 101):
                if i not in question_numbers:
                    missing_numbers.append(i)
            
            if missing_numbers:
                print(f"   - 缺失的题目编号: {missing_numbers[:20]}{'...' if len(missing_numbers) > 20 else ''}")
                analysis['content_analysis']['missing_numbers'] = missing_numbers
        
        # 深度文本模式分析
        print(f"\n🔬 深度模式分析:")
        
        all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # 各种题目模式
        patterns = [
            (r'\d+[．.]\s*[^0-9]', '标准题目模式 (数字+点+内容)'),
            (r'\d+\s*[．.]\s*[^0-9]', '带空格题目模式'),
            (r'^\d+[．.]', '行首数字模式'),
            (r'第\d+题', '第X题模式'),
            (r'\d+、', '数字顿号模式'),
            (r'\n\d+\s', '换行数字模式')
        ]
        
        for pattern, desc in patterns:
            matches = re.findall(pattern, all_text, re.MULTILINE)
            print(f"   - {desc}: 找到 {len(matches)} 个匹配")
            if matches:
                analysis['content_analysis']['question_patterns'].append({
                    'pattern': pattern,
                    'description': desc,
                    'match_count': len(matches),
                    'samples': matches[:5]
                })
        
        # 检查特殊字符和编码问题
        print(f"\n🔤 字符编码检查:")
        special_chars = set()
        for para in doc.paragraphs:
            for char in para.text:
                if ord(char) > 127:  # 非ASCII字符
                    special_chars.add(char)
        
        common_special = ['．', '。', '，', '、', '：', '；', '？', '！']
        found_special = [char for char in common_special if char in special_chars]
        print(f"   - 发现的特殊标点: {found_special}")
        
        # 检查文档格式问题
        print(f"\n📋 文档格式检查:")
        
        # 检查表格内容
        if doc.tables:
            print(f"   - 发现 {len(doc.tables)} 个表格，可能题目在表格中")
            for i, table in enumerate(doc.tables[:3]):  # 检查前3个表格
                print(f"     表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列")
                for row_idx, row in enumerate(table.rows[:5]):  # 每个表格看前5行
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        print(f"       行{row_idx+1}: {row_text[:100]}...")
        
        # 保存分析结果
        output_file = file_path.replace('.docx', '_diagnostic_report.json')
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, ensure_ascii=False, indent=2)
        
        print(f"\n💾 详细分析报告已保存到: {output_file}")
        return analysis
        
    except Exception as e:
        print(f"❌ 分析过程中出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def test_different_parsing_methods(file_path):
    """测试不同的解析方法"""
    print(f"\n🧪 测试不同解析方法:")
    print("=" * 40)
    
    try:
        doc = Document(file_path)
        
        # 方法1: 直接段落解析
        print("方法1: 直接段落解析")
        method1_questions = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^\d+[．.\s]', text):
                method1_questions.append(text[:50] + "...")
        print(f"   找到 {len(method1_questions)} 个可能的题目")
        
        # 方法2: 包含表格的解析
        print("\n方法2: 包含表格的解析")
        method2_questions = []
        
        # 段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^\d+[．.\s]', text):
                method2_questions.append(text[:50] + "...")
        
        # 表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if re.match(r'^\d+[．.\s]', text):
                            method2_questions.append(f"[表格]{text[:50]}...")
        
        print(f"   找到 {len(method2_questions)} 个可能的题目")
        
        # 方法3: 正则表达式扫描全文
        print("\n方法3: 正则表达式扫描全文")
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"
        
        # 多种正则模式
        patterns = [
            r'(\d+)[．.]\s*([^0-9\n]{10,})',  # 标准模式
            r'(\d+)\s*[．.]\s*([^0-9\n]{10,})',  # 带空格
            r'^(\d+)[．.]\s*(.+)$',  # 行首
        ]
        
        method3_questions = []
        for pattern in patterns:
            matches = re.findall(pattern, all_text, re.MULTILINE)
            for match in matches:
                if len(match) >= 2:
                    q_num = int(match[0])
                    if 1 <= q_num <= 100:
                        method3_questions.append(f"{q_num}. {match[1][:50]}...")
        
        method3_questions = list(set(method3_questions))  # 去重
        print(f"   找到 {len(method3_questions)} 个可能的题目")
        
        # 显示每种方法的前几个结果
        for i, (method_name, questions) in enumerate([
            ("方法1", method1_questions),
            ("方法2", method2_questions), 
            ("方法3", method3_questions)
        ], 1):
            print(f"\n{method_name}的前5个结果:")
            for j, q in enumerate(questions[:5]):
                print(f"   {j+1}. {q}")
        
        return {
            'method1': len(method1_questions),
            'method2': len(method2_questions),
            'method3': len(method3_questions)
        }
        
    except Exception as e:
        print(f"❌ 测试过程中出错: {str(e)}")
        return None

def main():
    """主函数"""
    print("🔧 2024年卷一完整版文档问题诊断工具")
    print("=" * 60)
    
    file_path = "/Users/acheng/Downloads/law-exam-assistant/2024年卷一完整版_修复版.docx"
    
    # 1. 基础文档分析
    analysis = analyze_document_structure(file_path)
    
    # 2. 测试不同解析方法
    parsing_results = test_different_parsing_methods(file_path)
    
    # 3. 总结和建议
    print(f"\n📋 诊断总结:")
    print("=" * 40)
    
    if analysis and analysis['content_analysis']['potential_questions']:
        total_found = len(analysis['content_analysis']['potential_questions'])
        print(f"✅ 基础分析找到 {total_found} 个题目")
        
        if total_found < 100:
            print(f"⚠️  题目数量不足，应该有100题")
            print(f"💡 可能的原因:")
            print(f"   - 文档格式与预期不符")
            print(f"   - 题目存储在表格中")
            print(f"   - 使用了特殊的编号格式")
            print(f"   - 读取脚本的正则表达式需要调整")
    else:
        print(f"❌ 基础分析未找到题目，这很不正常")
    
    if parsing_results:
        print(f"\n📊 不同解析方法结果对比:")
        for method, count in parsing_results.items():
            print(f"   {method}: {count} 个题目")
    
    print(f"\n🔧 建议的解决步骤:")
    print(f"1. 检查详细分析报告: 2024年卷一完整版_修复版_diagnostic_report.json")
    print(f"2. 根据报告调整读取脚本的正则表达式")
    print(f"3. 如果题目在表格中，需要修改脚本以解析表格内容")
    print(f"4. 考虑使用更宽松的题目识别模式")

if __name__ == "__main__":
    main()