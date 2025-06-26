#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2024年法考题智能分类系统
将题目按照8个法考科目分类并保存到独立文档
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import defaultdict

# 法考8大科目
SUBJECTS = ['民法', '刑法', '行政法', '民事诉讼法', '刑事诉讼法', '商经知', '三国法', '理论法']

# 法条引用映射（最高优先级）
LAW_CITATIONS = {
    '刑法': [
        '《刑法》', '刑法第', '《中华人民共和国刑法》', '刑法条文', '刑法典',
        '《刑法修正案》', '刑法分则', '刑法总则'
    ],
    '民法': [
        '《民法典》', '民法典第', '《中华人民共和国民法典》', '《民法通则》',
        '《物权法》', '《合同法》', '《侵权责任法》', '《婚姻法》', '《继承法》'
    ],
    '民事诉讼法': [
        '《民事诉讼法》', '民诉法第', '《中华人民共和国民事诉讼法》', 
        '《最高人民法院关于适用〈中华人民共和国民事诉讼法〉》',
        '民事诉讼法解释', '民诉解释', '《民事诉讼法解释》'
    ],
    '刑事诉讼法': [
        '《刑事诉讼法》', '刑诉法第', '《中华人民共和国刑事诉讼法》',
        '《最高人民法院关于适用〈中华人民共和国刑事诉讼法〉》',
        '刑事诉讼法解释', '刑诉解释'
    ],
    '行政法': [
        '《行政处罚法》', '《行政许可法》', '《行政强制法》', '《行政复议法》',
        '《行政诉讼法》', '《国家赔偿法》', '《公务员法》', '《政府信息公开条例》'
    ],
    '商经知': [
        '《公司法》', '《证券法》', '《保险法》', '《票据法》', '《破产法》',
        '《合伙企业法》', '《个人独资企业法》', '《反垄断法》', '《消费者权益保护法》',
        '《产品质量法》', '《劳动法》', '《劳动合同法》', '《知识产权法》', '《专利法》',
        '《商标法》', '《著作权法》', '《反不正当竞争法》', '《企业破产法》',
        '《建设工程施工合同解释》'
    ],
    '三国法': [
        '《联合国国际货物销售合同公约》', '《国际私法》', '《涉外民事关系法律适用法》',
        '《海商法》', '《维也纳外交关系公约》', '《维也纳领事关系公约》',
        'INCOTERMS', 'CIF', 'FOB', 'CIP', '《纽约公约》', '《海牙公约》'
    ],
    '理论法': [
        '《宪法》', '《立法法》', '《监察法》', '《选举法》', '《代表法》',
        '《法官法》', '《检察官法》', '《律师法》', '《仲裁法》', '宪法修正案',
        '《法理学》', '《法治理论》'
    ]
}

# 专业术语（高权重关键词）
PROFESSIONAL_TERMS = {
    '刑法': [
        '犯罪构成', '主观要件', '客观要件', '共同犯罪', '正当防卫', '紧急避险',
        '数罪并罚', '累犯', '自首', '立功', '缓刑', '假释', '死刑缓期',
        '不作为犯', '过失致死', '故意伤害', '故意杀人', '盗窃罪', '抢劫罪',
        '诈骗罪', '贪污罪', '受贿罪', '职务侵占', '挪用公款', '渎职罪',
        '交通肇事罪', '危险驾驶罪', '寻衅滋事', '聚众斗殴', '非法拘禁',
        '强奸罪', '绑架罪', '敲诈勒索', '毒品犯罪', '走私罪', '洗钱罪'
    ],
    '民法': [
        '法律行为', '民事法律行为', '代理制度', '诉讼时效', '物权变动',
        '善意取得', '不动产登记', '用益物权', '担保物权', '债权债务',
        '不当得利', '无因管理', '缔约过失', '违约责任', '侵权责任',
        '婚姻家庭', '夫妻财产', '监护制度', '继承权', '遗嘱继承',
        '人格权', '肖像权', '隐私权', '名誉权', '收养关系', '离婚协议',
        '买卖合同', '租赁合同', '赠与合同', '所有权', '用益物权', '担保物权',
        '相邻关系', '共有', '遗产继承', '法定继承', '遗嘱', '遗赠'
    ],
    '民事诉讼法': [
        '管辖权异议', '级别管辖', '地域管辖', '专属管辖', '协议管辖',
        '举证责任', '证据规则', '财产保全', '先予执行', '简易程序',
        '小额诉讼', '缺席判决', '第三人撤销之诉', '案外人执行异议',
        '申请再审', '审判监督程序', '强制执行', '执行和解', '发回重审',
        '合议庭', '独任审判', '人民陪审员', '行为保全', '人身安全保护令'
    ],
    '刑事诉讼法': [
        '侦查程序', '审查起诉', '审判程序', '强制措施', '取保候审',
        '监视居住', '刑事拘留', '逮捕', '搜查', '扣押', '查封', '冻结',
        '辩护权', '法律援助', '证人保护', '技侦措施', '认罪认罚',
        '速裁程序', '简易程序', '死刑复核', '刑事和解', '缺席审判'
    ],
    '行政法': [
        '行政主体', '行政相对人', '具体行政行为', '抽象行政行为',
        '行政许可', '行政处罚', '行政强制', '行政复议', '行政诉讼',
        '政府信息公开', '行政程序', '听证程序', '国家赔偿', '公务员',
        '市场监督管理局', '行政复议决定', '维持决定'
    ],
    '商经知': [
        '股东大会', '董事会', '监事会', '股权转让', '公司治理',
        '有限责任公司', '股份有限公司', '合伙企业', '个人独资企业',
        '公司设立', '公司解散', '破产重整', '破产清算', '证券发行',
        '票据权利', '票据责任', '保险合同', '保险责任', '劳动合同',
        '社会保险', '工伤保险', '反垄断', '不正当竞争', '消费者权益',
        '股东分红', '股权回购', '破产管理人', '债权申报', '清算程序',
        '合伙协议', '执行事务合伙人', '普通合伙企业', '建设工程合同'
    ],
    '三国法': [
        '国际私法', '法律冲突', '准据法', '反致', '识别', '先决问题',
        '公共秩序保留', '法律规避', '涉外合同', '涉外侵权', '涉外婚姻',
        '涉外继承', '国际商事仲裁', '司法协助', '外国判决承认执行',
        '国际贸易术语', '提单', '共同海损', '海上保险', '租船合同',
        '涉外收养', '缔约国', '出口管制', '国际货物销售'
    ],
    '理论法': [
        '法治理论', '法律原则', '法律规则', '法的效力', '法的适用',
        '立法程序', '法律解释', '法律冲突', '宪法权利', '基本权利',
        '法治国家', '依法治国', '党内法规', '实证主义', '非实证主义',
        '形式法治', '实质法治', '法律体系', '法律渊源'
    ]
}

# 特征词汇（较低权重）
FEATURE_WORDS = {
    '刑法': ['犯罪', '刑罚', '罪名', '量刑', '追诉', '刑事责任'],
    '民法': ['权利', '义务', '民事责任', '合同', '侵权', '财产'],
    '民事诉讼法': ['诉讼', '审理', '判决', '裁定', '执行', '管辖'],
    '刑事诉讼法': ['侦查', '起诉', '审判', '辩护', '证据', '强制措施'],
    '行政法': ['行政', '政府', '处罚', '许可', '复议', '赔偿'],
    '商经知': ['公司', '股东', '企业', '破产', '保险', '劳动'],
    '三国法': ['国际', '涉外', '外国', '公约', '冲突', '准据法'],
    '理论法': ['法治', '宪法', '立法', '司法', '法理', '法律体系']
}

def extract_questions_from_document(doc_path):
    """从文档中提取所有题目"""
    doc = Document(doc_path)
    questions = []
    current_question = None
    current_content = []
    question_pattern = re.compile(r'^(\d+)\.')
    in_question = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目的开始
        match = question_pattern.match(text)
        if match and not in_question:  # 只在不在题目中时才开始新题目
            # 保存前一个题目
            if current_question is not None:
                questions.append({
                    'number': current_question,
                    'content': '\n'.join(current_content)
                })
            
            # 开始新题目
            current_question = int(match.group(1))
            current_content = [text]
            in_question = True
        elif in_question:
            # 继续收集当前题目内容，直到遇到分隔线
            if text.startswith('---'):
                # 题目结束，保存
                questions.append({
                    'number': current_question,
                    'content': '\n'.join(current_content)
                })
                current_question = None
                current_content = []
                in_question = False
            else:
                current_content.append(text)
    
    # 保存最后一个题目
    if current_question is not None:
        questions.append({
            'number': current_question,
            'content': '\n'.join(current_content)
        })
    
    return questions

def classify_question(question_content):
    """对单个题目进行分类"""
    scores = {subject: 0 for subject in SUBJECTS}
    
    # 1. 法条引用（最高权重：100分）
    for subject, citations in LAW_CITATIONS.items():
        for citation in citations:
            if citation in question_content:
                scores[subject] += 100
    
    # 2. 专业术语（高权重：10分）
    for subject, terms in PROFESSIONAL_TERMS.items():
        for term in terms:
            if term in question_content:
                scores[subject] += 10
    
    # 3. 特征词汇（低权重：1分）
    for subject, words in FEATURE_WORDS.items():
        for word in words:
            if word in question_content:
                scores[subject] += 1
    
    # 4. 特殊规则
    # 三国法特殊处理
    if any(keyword in question_content for keyword in ['涉外', '国际', '外国人', '缔约国', 'CISG']):
        scores['三国法'] += 50
    
    # 如果同时涉及多个科目，根据上下文进行调整
    if scores['民法'] > 0 and '诉讼' in question_content:
        if '审理' in question_content or '管辖' in question_content:
            scores['民事诉讼法'] += 50
    
    if scores['刑法'] > 0 and '诉讼' in question_content:
        if '侦查' in question_content or '起诉' in question_content:
            scores['刑事诉讼法'] += 50
    
    # 返回得分最高的科目
    if max(scores.values()) == 0:
        return '理论法'  # 默认分类
    
    return max(scores, key=scores.get)

def create_subject_document(subject, questions, output_dir):
    """为特定科目创建文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(f'2024年法考客观题 - {subject}', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph()
    stats.add_run(f'共 {len(questions)} 道题').bold = True
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(20)
    
    # 分隔线
    doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 添加每道题
    for question in questions:
        # 添加题目内容
        lines = question['content'].split('\n')
        for line in lines:
            doc.add_paragraph(line)
        
        # 添加分隔线
        separator = doc.add_paragraph('-' * 30)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(6)
        separator.paragraph_format.space_after = Pt(6)
    
    # 保存文档
    output_path = os.path.join(output_dir, f'{subject}.docx')
    doc.save(output_path)
    print(f"已保存 {subject} 科目，共 {len(questions)} 题到: {output_path}")

def main():
    """主函数"""
    # 文件路径
    input_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx'
    
    print("2024年法考题智能分类系统")
    print("=" * 60)
    
    # 确保输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 提取所有题目
    print("\n提取题目...")
    questions = extract_questions_from_document(input_file)
    print(f"共提取 {len(questions)} 道题")
    
    # 对每道题进行分类
    print("\n开始分类...")
    classified_questions = defaultdict(list)
    
    for question in questions:
        subject = classify_question(question['content'])
        classified_questions[subject].append(question)
        print(f"题目 {question['number']:3d} -> {subject}")
    
    # 统计结果
    print("\n分类结果统计:")
    print("-" * 40)
    total = 0
    for subject in SUBJECTS:
        count = len(classified_questions[subject])
        total += count
        print(f"{subject:10s}: {count:3d} 题")
    print("-" * 40)
    print(f"总计: {total} 题")
    
    # 为每个科目创建独立文档
    print("\n生成分科目文档...")
    for subject in SUBJECTS:
        if classified_questions[subject]:
            create_subject_document(subject, classified_questions[subject], output_dir)
    
    print("\n✓ 分类完成！")
    print(f"所有文档已保存到: {output_dir}")

if __name__ == "__main__":
    main()