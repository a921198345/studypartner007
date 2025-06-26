#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""诊断上传超时问题"""

import sys
import os
import time
import signal
import json
from docx import Document

def signal_handler(signum, frame):
    """处理信号"""
    print(f"[ERROR] 收到信号 {signum} ({signal.Signals(signum).name})")
    if signum == signal.SIGTERM:
        print("[ERROR] 收到SIGTERM信号，通常是因为超时或内存限制")
    sys.exit(1)

# 注册信号处理器
signal.signal(signal.SIGTERM, signal_handler)
signal.signal(signal.SIGINT, signal_handler)

def check_file_size(file_path):
    """检查文件大小"""
    if os.path.exists(file_path):
        size = os.path.getsize(file_path)
        print(f"[INFO] 文件大小: {size} bytes ({size/1024/1024:.2f} MB)")
        return size
    else:
        print(f"[ERROR] 文件不存在: {file_path}")
        return 0

def test_docx_read(file_path):
    """测试读取docx文件"""
    print("[INFO] 开始读取docx文件...")
    start_time = time.time()
    
    try:
        doc = Document(file_path)
        para_count = len(doc.paragraphs)
        print(f"[INFO] 成功读取文档，共有 {para_count} 个段落")
        
        # 计算文档总字符数
        total_chars = 0
        for para in doc.paragraphs:
            total_chars += len(para.text)
        
        elapsed = time.time() - start_time
        print(f"[INFO] 文档总字符数: {total_chars}")
        print(f"[INFO] 读取耗时: {elapsed:.2f} 秒")
        
        return True
    except Exception as e:
        elapsed = time.time() - start_time
        print(f"[ERROR] 读取文档失败: {e}")
        print(f"[ERROR] 失败耗时: {elapsed:.2f} 秒")
        return False

def test_parse_speed(file_path):
    """测试解析速度"""
    print("\n[INFO] 测试解析速度...")
    
    try:
        import parse_questions
        
        start_time = time.time()
        questions, format_issues = parse_questions.extract_questions_from_doc(
            file_path, "民法", validate=True
        )
        elapsed = time.time() - start_time
        
        print(f"[INFO] 解析完成，共找到 {len(questions)} 个题目")
        print(f"[INFO] 格式问题: {len(format_issues)} 个")
        print(f"[INFO] 解析耗时: {elapsed:.2f} 秒")
        
        # 输出前3个题目作为示例
        for i, q in enumerate(questions[:3]):
            print(f"\n[题目{i+1}] ID: {q['question_id']}")
            print(f"  题干: {q['question_text'][:50]}...")
            print(f"  选项数: {len(q['options'])}")
            print(f"  答案: {q['correct_answer']}")
        
        return True
    except Exception as e:
        print(f"[ERROR] 解析失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    if len(sys.argv) < 2:
        print("Usage: python diagnose-upload-timeout.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    print(f"[INFO] 诊断文件: {file_path}")
    print(f"[INFO] Python版本: {sys.version}")
    print(f"[INFO] 当前进程ID: {os.getpid()}")
    
    # 检查文件
    file_size = check_file_size(file_path)
    if file_size == 0:
        sys.exit(1)
    
    # 测试读取文档
    if not test_docx_read(file_path):
        sys.exit(1)
    
    # 测试解析速度
    test_parse_speed(file_path)
    
    print("\n[INFO] 诊断完成")

if __name__ == "__main__":
    main()