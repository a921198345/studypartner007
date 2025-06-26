#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终版本：合并2024年法考真题和答案解析
支持.doc和.docx格式，优化题目提取逻辑
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import subprocess
import tempfile

def convert_doc_to_docx(doc_path):
    """将.doc文件转换为.docx格式"""
    print(f"转换 .doc 文件为 .docx 格式...")
    
    # 尝试使用python-docx2txt
    try:
        import docx2txt
        text = docx2txt.process(doc_path)
        
        # 创建临时docx文件
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc = Document()
        
        # 将文本按行分割并添加到文档
        lines = text.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(temp_docx.name)
        return temp_docx.name
    except ImportError:
        print("需要安装 docx2txt 库来处理 .doc 文件")
        print("正在尝试安装...")
        subprocess.run(['pip', 'install', 'docx2txt'])
        return convert_doc_to_docx(doc_path)
    except Exception as e:
        print(f"转换错误: {e}")
        
        # 尝试textract作为备选
        try:
            import textract
            text = textract.process(doc_path).decode('utf-8')
            
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            doc = Document()
            
            lines = text.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line)
            
            doc.save(temp_docx.name)
            return temp_docx.name
        except:
            print("无法转换 .doc 文件，请手动转换为 .docx 格式")
            return None

def extract_questions_comprehensive(doc_path):
    """全面提取题目内容"""
    # 处理.doc文件
    if doc_path.endswith('.doc'):
        temp_path = convert_doc_to_docx(doc_path)
        if not temp_path:
            return {}
        doc_path = temp_path
    
    doc = Document(doc_path)
    questions = {}
    current_question = None
    current_content = []
    in_question = False
    
    # 更全面的题号匹配模式
    question_patterns = [
        re.compile(r'^(\d+)[．.、。]\s*(.*)$'),           # 1．或1.或1、或1。
        re.compile(r'^(\d+)\s+(.*)$'),                  # 1 开头
        re.compile(r'^\s*(\d+)[．.、。]\s*(.*)$'),       # 前面有空格
        re.compile(r'^第(\d+)题[：:]\s*(.*)$'),          # 第1题：
        re.compile(r'^题(\d+)[：:]\s*(.*)$'),            # 题1：
        re.compile(r'^【(\d+)】\s*(.*)$'),              # 【1】
    ]
    
    print(f"开始解析题目文档: {doc_path}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目
        is_new_question = False
        for pattern in question_patterns:
            match = pattern.match(text)
            if match:
                num = int(match.group(1))
                # 合理的题号范围（1-200）
                if 1 <= num <= 200:
                    # 保存前一题
                    if current_question is not None and current_content:
                        questions[current_question] = {
                            'content': '\n'.join(current_content),
                            'has_options': any('A.' in line or 'A、' in line or 'A）' in line 
                                             for line in current_content)
                        }
                    
                    # 开始新题目
                    current_question = num
                    current_content = [text]
                    in_question = True
                    is_new_question = True
                    break
        
        if not is_new_question and in_question:
            # 继续收集当前题目内容
            current_content.append(text)
            
            # 检查是否已经收集到选项D，可能题目结束
            if any(marker in text for marker in ['D.', 'D、', 'D）', 'D．']):
                # 可能是题目结束，但继续收集直到下一题
                pass
    
    # 保存最后一题
    if current_question is not None and current_content:
        questions[current_question] = {
            'content': '\n'.join(current_content),
            'has_options': any('A.' in line or 'A、' in line or 'A）' in line 
                             for line in current_content)
        }
    
    print(f"从题目文档提取到 {len(questions)} 道题")
    
    # 清理临时文件
    if doc_path.endswith('.docx') and 'temp' in doc_path:
        try:
            os.unlink(doc_path)
        except:
            pass
    
    return questions

def extract_answers_comprehensive(doc_path):
    """全面提取答案解析"""
    doc = Document(doc_path)
    answers = {}
    current_question = None
    current_answer = None
    current_content = []
    
    print(f"开始解析答案文档: {doc_path}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 匹配答案开始（多种格式）
        answer_patterns = [
            re.compile(r'^(\d+)[．.、]\s*正确答案[：:]\s*([A-Z])', re.IGNORECASE),
            re.compile(r'^(\d+)[．.、]\s*答案[：:]\s*([A-Z])', re.IGNORECASE),
            re.compile(r'^第(\d+)题[：:]\s*([A-Z])', re.IGNORECASE),
        ]
        
        is_new_answer = False
        for pattern in answer_patterns:
            match = pattern.match(text)
            if match:
                # 保存前一题的答案
                if current_question is not None:
                    answers[current_question] = {
                        'answer': current_answer,
                        'content': '\n'.join(current_content)
                    }
                
                # 开始新答案
                current_question = int(match.group(1))
                current_answer = match.group(2).upper()
                current_content = [text]
                is_new_answer = True
                break
        
        if not is_new_answer and current_question is not None:
            # 继续收集解析内容
            if not (text.startswith('综上所述') and '本题答案' in text):
                current_content.append(text)
    
    # 保存最后一题
    if current_question is not None:
        answers[current_question] = {
            'answer': current_answer,
            'content': '\n'.join(current_content)
        }
    
    print(f"从答案文档提取到 {len(answers)} 道题的答案解析")
    return answers

def format_question_with_answer(doc, num, question_data, answer_data):
    """格式化题目和答案"""
    
    # 题目编号（加粗、蓝色）
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"第 {num} 题")
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(8)
    
    # 题目内容
    if question_data:
        content_lines = question_data['content'].split('\n')
        for line in content_lines:
            if line.strip():
                # 检查是否是选项
                if re.match(r'^[A-D][．.、）]\s*', line.strip()):
                    # 选项缩进显示
                    option_para = doc.add_paragraph()
                    option_para.paragraph_format.left_indent = Pt(20)
                    option_para.add_run(line.strip())
                else:
                    # 题干正常显示
                    doc.add_paragraph(line.strip())
    else:
        # 没有题目内容时的提示
        no_question = doc.add_paragraph()
        no_question.add_run("【题目内容暂缺，请参考答案解析】").italic = True
        no_question.paragraph_format.space_after = Pt(8)
    
    # 答案和解析
    if answer_data:
        # 正确答案（红色、加粗）
        answer_para = doc.add_paragraph()
        answer_para.paragraph_format.space_before = Pt(8)
        answer_run = answer_para.add_run(f"【正确答案】{answer_data['answer']}")
        answer_run.font.size = Pt(12)
        answer_run.bold = True
        answer_run.font.color.rgb = RGBColor(192, 0, 0)
        
        # 答案解析标题
        analysis_title = doc.add_paragraph()
        analysis_run = analysis_title.add_run("【答案解析】")
        analysis_run.font.size = Pt(12)
        analysis_run.bold = True
        analysis_run.font.color.rgb = RGBColor(0, 128, 0)
        analysis_title.paragraph_format.space_before = Pt(6)
        analysis_title.paragraph_format.space_after = Pt(4)
        
        # 解析内容
        analysis_lines = answer_data['content'].split('\n')
        for line in analysis_lines:
            line = line.strip()
            if line and not line.startswith('【答案解析】'):
                # 处理选项解析
                if re.match(r'^[A-D]\s*项?[：:]', line):
                    # 选项解析缩进
                    option_analysis = doc.add_paragraph()
                    option_analysis.paragraph_format.left_indent = Pt(20)
                    option_analysis.add_run(line)
                    option_analysis.paragraph_format.space_after = Pt(3)
                else:
                    # 普通解析内容
                    analysis_para = doc.add_paragraph(line)
                    analysis_para.paragraph_format.space_after = Pt(3)
    else:
        # 没有答案解析
        no_answer = doc.add_paragraph()
        no_answer_run = no_answer.add_run("【答案解析暂缺】")
        no_answer_run.font.color.rgb = RGBColor(128, 128, 128)
        no_answer.paragraph_format.space_before = Pt(8)
    
    # 分隔线
    separator = doc.add_paragraph("─" * 80)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    separator.paragraph_format.space_before = Pt(10)
    separator.paragraph_format.space_after = Pt(10)

def create_final_document(questions, answers, output_path):
    """创建最终的合并文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 文档标题
    title = doc.add_heading('2024年法考客观题真题及答案解析', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('完整版')
    subtitle_run.font.size = Pt(16)
    subtitle_run.bold = True
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    
    # 文档说明
    info = doc.add_paragraph()
    info.add_run("说明：本文档整合了2024年法考客观卷真题和官方答案解析，每道题目后紧跟正确答案及详细解析。")
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info.paragraph_format.space_after = Pt(10)
    
    # 统计信息
    all_numbers = sorted(set(questions.keys()) | set(answers.keys()))
    stats = doc.add_paragraph()
    stats.add_run(f"共收录 {len(all_numbers)} 道题目 | ")
    stats.add_run(f"题目完整: {len(questions)} 道 | ")
    stats.add_run(f"答案解析: {len(answers)} 道")
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(30)
    
    # 主分隔线
    doc.add_paragraph("═" * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    print(f"\n生成最终文档...")
    print(f"总题数: {len(all_numbers)}")
    
    # 按题号顺序添加
    for i, num in enumerate(all_numbers, 1):
        if i % 10 == 0:
            print(f"处理进度: {i}/{len(all_numbers)}")
        
        question_data = questions.get(num)
        answer_data = answers.get(num)
        
        format_question_with_answer(doc, num, question_data, answer_data)
    
    # 文档结尾
    doc.add_paragraph()
    doc.add_paragraph("═" * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    end_para = doc.add_paragraph()
    end_para.add_run("— 全部试题结束 —")
    end_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    end_para.paragraph_format.space_before = Pt(20)
    
    # 保存文档
    doc.save(output_path)
    print(f"\n文档生成完成！")
    print(f"保存位置: {output_path}")
    
    # 输出统计
    only_questions = set(questions.keys()) - set(answers.keys())
    only_answers = set(answers.keys()) - set(questions.keys())
    both = set(questions.keys()) & set(answers.keys())
    
    print(f"\n统计信息:")
    print(f"- 题目和答案都有: {len(both)} 道")
    print(f"- 只有题目无答案: {len(only_questions)} 道")
    print(f"- 只有答案无题目: {len(only_answers)} 道")
    
    if only_questions:
        print(f"\n只有题目的题号: {sorted(only_questions)}")
    if only_answers:
        print(f"只有答案的题号: {sorted(only_answers)[:10]}{'...' if len(only_answers) > 10 else ''}")

def main():
    """主函数"""
    # 文件路径
    questions_file = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）_扫描版.doc'
    answers_file = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    
    print("=" * 60)
    print("2024年法考真题完整版生成工具")
    print("=" * 60)
    
    # 检查文件
    if not os.path.exists(questions_file):
        print(f"错误：找不到题目文件 {questions_file}")
        # 尝试其他可能的题目文件
        alt_file = questions_file.replace('_扫描版.doc', '.docx')
        if os.path.exists(alt_file):
            print(f"使用替代文件: {alt_file}")
            questions_file = alt_file
        else:
            return
    
    if not os.path.exists(answers_file):
        print(f"错误：找不到答案文件 {answers_file}")
        return
    
    # 提取数据
    print("\n第一步：提取题目内容")
    questions = extract_questions_comprehensive(questions_file)
    
    print("\n第二步：提取答案解析")
    answers = extract_answers_comprehensive(answers_file)
    
    print("\n第三步：生成合并文档")
    create_final_document(questions, answers, output_file)
    
    print("\n✓ 处理完成！")

if __name__ == "__main__":
    main()