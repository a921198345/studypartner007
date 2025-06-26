#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能法考题分类系统
基于法条引用、专业术语和语境分析的多层分类算法
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 法条引用映射（最高优先级）
LAW_CITATIONS = {
    '刑法': [
        '《刑法》', '刑法第', '《中华人民共和国刑法》', '刑法条文', '刑法典',
        '《刑法修正案》', '刑法分则', '刑法总则'
    ],
    '民法': [
        '《民法典》', '民法典第', '《中华人民共和国民法典》', '《民法通则》',
        '《物权法》', '《合同法》', '《侵权责任法》', '《婚姻法》', '《继承法》'
    ],
    '民事诉讼法': [
        '《民事诉讼法》', '民诉法第', '《中华人民共和国民事诉讼法》', 
        '《最高人民法院关于适用〈中华人民共和国民事诉讼法〉》',
        '民事诉讼法解释', '民诉解释'
    ],
    '刑事诉讼法': [
        '《刑事诉讼法》', '刑诉法第', '《中华人民共和国刑事诉讼法》',
        '《最高人民法院关于适用〈中华人民共和国刑事诉讼法〉》',
        '刑事诉讼法解释', '刑诉解释'
    ],
    '行政法': [
        '《行政处罚法》', '《行政许可法》', '《行政强制法》', '《行政复议法》',
        '《行政诉讼法》', '《国家赔偿法》', '《公务员法》', '《政府信息公开条例》'
    ],
    '商经知': [
        '《公司法》', '《证券法》', '《保险法》', '《票据法》', '《破产法》',
        '《合伙企业法》', '《个人独资企业法》', '《反垄断法》', '《消费者权益保护法》',
        '《产品质量法》', '《劳动法》', '《劳动合同法》', '《税法》', '《建设工程质量管理条例》'
    ],
    '三国法': [
        '《联合国国际货物销售合同公约》', '《国际私法》', '《涉外民事关系法律适用法》',
        '《海商法》', '《维也纳外交关系公约》', '《维也纳领事关系公约》',
        'INCOTERMS', 'CIF', 'FOB', 'CIP'
    ],
    '理论法': [
        '《宪法》', '《立法法》', '《监察法》', '《选举法》', '《代表法》',
        '《法官法》', '《检察官法》', '《律师法》', '《仲裁法》', '宪法修正案'
    ]
}

# 专业术语（高权重关键词）
PROFESSIONAL_TERMS = {
    '刑法': [
        '犯罪构成', '主观要件', '客观要件', '共同犯罪', '正当防卫', '紧急避险',
        '数罪并罚', '累犯', '自首', '立功', '缓刑', '假释', '死刑缓期',
        '不作为犯', '过失致死', '故意伤害', '故意杀人', '盗窃罪', '抢劫罪',
        '诈骗罪', '贪污罪', '受贿罪', '职务侵占', '挪用公款', '渎职罪',
        '交通肇事罪', '危险驾驶罪', '寻衅滋事', '聚众斗殴', '非法拘禁',
        '强奸罪', '绑架罪', '敲诈勒索', '毒品犯罪', '走私罪', '洗钱罪'
    ],
    '民法': [
        '法人制度', '民事法律行为', '代理制度', '诉讼时效', '物权变动',
        '善意取得', '不动产登记', '用益物权', '担保物权', '债权债务',
        '不当得利', '无因管理', '缔约过失', '违约责任', '侵权责任',
        '婚姻家庭', '夫妻财产', '监护制度', '继承权', '遗嘱继承',
        '人格权', '肖像权', '隐私权', '名誉权'
    ],
    '民事诉讼法': [
        '管辖权异议', '级别管辖', '地域管辖', '专属管辖', '协议管辖',
        '举证责任', '证据规则', '财产保全', '先予执行', '简易程序',
        '小额诉讼', '缺席判决', '第三人撤销之诉', '案外人执行异议',
        '申请再审', '审判监督程序', '强制执行', '执行和解'
    ],
    '刑事诉讼法': [
        '侦查程序', '审查起诉', '审判程序', '强制措施', '取保候审',
        '监视居住', '刑事拘留', '逮捕', '搜查', '扣押', '查封', '冻结',
        '辩护权', '法律援助', '证人保护', '技侦措施', '认罪认罚',
        '速裁程序', '简易程序', '死刑复核', '刑事和解', '缺席审判'
    ],
    '行政法': [
        '行政主体', '行政相对人', '具体行政行为', '抽象行政行为',
        '行政许可', '行政处罚', '行政强制', '行政复议', '行政诉讼',
        '政府信息公开', '行政程序', '听证程序', '国家赔偿', '公务员'
    ],
    '商经知': [
        '股东大会', '董事会', '监事会', '股权转让', '公司治理',
        '有限责任公司', '股份有限公司', '合伙企业', '个人独资企业',
        '公司设立', '公司解散', '破产重整', '破产清算', '证券发行',
        '票据权利', '票据责任', '保险合同', '保险责任', '劳动合同',
        '社会保险', '工伤保险', '反垄断', '不正当竞争', '消费者权益'
    ],
    '三国法': [
        '国际私法', '法律冲突', '准据法', '反致', '识别', '先决问题',
        '公共秩序保留', '法律规避', '涉外合同', '涉外侵权', '涉外婚姻',
        '涉外继承', '国际商事仲裁', '司法协助', '外国判决承认执行',
        '国际贸易术语', '提单', '共同海损', '海上保险', '租船合同'
    ],
    '理论法': [
        '人民代表大会制度', '民主集中制', '基本权利', '公民权利',
        '政治权利', '人身自由', '宗教信仰自由', '选举权', '监督权',
        '国家机构', '全国人大', '国务院', '监察委员会', '人民法院',
        '人民检察院', '法律职业道德', '司法制度', '立法程序'
    ]
}

# 通用关键词（低权重）
GENERAL_KEYWORDS = {
    '刑法': ['刑罚', '犯罪', '有期徒刑', '拘役', '管制', '罚金', '没收财产'],
    '民法': ['合同', '协议', '债务', '赔偿', '损失', '责任'],
    '民事诉讼法': ['起诉', '上诉', '判决', '裁定', '执行'],
    '刑事诉讼法': ['起诉', '公诉', '辩护', '证据', '鉴定'],
    '行政法': ['行政', '政府', '行政机关', '复议', '行政诉讼'],
    '商经知': ['公司', '企业', '股东', '经营'],
    '三国法': ['国际', '涉外', '外国'],
    '理论法': ['宪法', '法律', '权利', '义务', '制度']
}

def classify_by_law_citation(text):
    """基于法条引用进行分类（最高优先级）"""
    for subject, citations in LAW_CITATIONS.items():
        for citation in citations:
            if citation in text:
                return subject, 10  # 最高权重
    return None, 0

def classify_by_professional_terms(text):
    """基于专业术语进行分类"""
    subject_scores = {}
    
    for subject, terms in PROFESSIONAL_TERMS.items():
        score = 0
        for term in terms:
            if term in text:
                score += 5  # 专业术语高权重
        
        if score > 0:
            subject_scores[subject] = score
    
    if subject_scores:
        best_subject = max(subject_scores.items(), key=lambda x: x[1])
        return best_subject[0], best_subject[1]
    
    return None, 0

def classify_by_general_keywords(text):
    """基于通用关键词进行分类"""
    subject_scores = {}
    
    for subject, keywords in GENERAL_KEYWORDS.items():
        score = 0
        for keyword in keywords:
            if keyword in text:
                score += 1  # 通用关键词低权重
        
        if score > 0:
            subject_scores[subject] = score
    
    if subject_scores:
        best_subject = max(subject_scores.items(), key=lambda x: x[1])
        return best_subject[0], best_subject[1]
    
    return None, 0

def intelligent_classify(question_text):
    """智能分类：多层次分析"""
    # 第一层：法条引用分析（最可靠）
    subject1, score1 = classify_by_law_citation(question_text)
    if subject1 and score1 >= 10:
        return subject1, f"法条引用:{score1}"
    
    # 第二层：专业术语分析
    subject2, score2 = classify_by_professional_terms(question_text)
    
    # 第三层：通用关键词分析
    subject3, score3 = classify_by_general_keywords(question_text)
    
    # 综合评分
    final_scores = {}
    
    # 法条引用得分
    if subject1:
        final_scores[subject1] = final_scores.get(subject1, 0) + score1
    
    # 专业术语得分
    if subject2:
        final_scores[subject2] = final_scores.get(subject2, 0) + score2
    
    # 通用关键词得分
    if subject3:
        final_scores[subject3] = final_scores.get(subject3, 0) + score3
    
    if final_scores:
        best_subject = max(final_scores.items(), key=lambda x: x[1])
        return best_subject[0], f"综合评分:{best_subject[1]}"
    
    return None, "未分类"

def extract_questions_simple(doc_path):
    """提取题目和答案"""
    doc = Document(doc_path)
    questions = {}
    current_question = None
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 跳过标题
        if text == '2023年法考客观题（真题+答案解析）':
            continue
        
        # 检查是否是分隔线
        if text.startswith('------'):
            if current_question and current_content:
                questions[current_question] = '\n'.join(current_content)
            current_question = None
            current_content = []
            continue
        
        # 匹配题号
        if text and text[0].isdigit():
            dot_pos = -1
            for i, char in enumerate(text):
                if char in '.．。':
                    dot_pos = i
                    break
            
            if dot_pos > 0 and dot_pos < 5:
                try:
                    q_num = int(text[:dot_pos])
                    if '正确答案' in text:
                        if current_question:
                            current_content.append(text)
                    else:
                        if current_question and current_content:
                            questions[current_question] = '\n'.join(current_content)
                        current_question = q_num
                        current_content = [text]
                except ValueError:
                    if current_question:
                        current_content.append(text)
            else:
                if current_question:
                    current_content.append(text)
        else:
            if current_question:
                current_content.append(text)
    
    # 保存最后一题
    if current_question and current_content:
        questions[current_question] = '\n'.join(current_content)
    
    return questions

def create_subject_document(subject, questions, output_dir, classification_details):
    """为特定科目创建文档"""
    doc = Document()
    
    # 设置文档的默认字体
    doc.styles['Normal'].font.name = '宋体'
    
    # 添加标题
    title = doc.add_heading(f'2023年法考真题 - {subject}', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph(f'共 {len(questions)} 道题')
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(12)
    
    # 添加分类说明
    if classification_details:
        details_para = doc.add_paragraph()
        details_para.add_run('分类依据：').bold = True
        details_text = []
        for method, count in classification_details.items():
            details_text.append(f"{method}: {count}题")
        details_para.add_run(' | '.join(details_text))
        details_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        details_para.paragraph_format.space_after = Pt(12)
    
    # 添加分隔线
    doc.add_paragraph("═" * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 添加题目
    for q_num, content in sorted(questions.items()):
        # 题目内容
        doc.add_paragraph(content)
        
        # 添加分隔线
        separator = doc.add_paragraph("─" * 80)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(12)
        separator.paragraph_format.space_after = Pt(12)
    
    # 保存文档
    filename = f"{subject}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filename

def main():
    # 输入文件路径
    input_file = '/Users/acheng/Downloads/law-exam-assistant/23年法考题.docx'
    
    # 检查文件是否存在
    if not os.path.exists(input_file):
        print(f"错误：找不到文件 {input_file}")
        return
    
    # 创建输出目录
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx_intelligent'
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"开始智能分类处理: {input_file}")
    print(f"输出目录: {output_dir}")
    
    # 提取所有题目
    print("\n提取题目中...")
    all_questions = extract_questions_simple(input_file)
    print(f"共提取到 {len(all_questions)} 道题目")
    
    if not all_questions:
        print("未能提取到任何题目，请检查文档格式")
        return
    
    # 智能分类
    subject_questions = {
        '民法': {}, '刑法': {}, '民事诉讼法': {}, '刑事诉讼法': {}, 
        '行政法': {}, '商经知': {}, '三国法': {}, '理论法': {}
    }
    unclassified = {}
    classification_methods = {subject: {} for subject in subject_questions.keys()}
    
    print("\n开始智能分类...")
    for q_num, content in all_questions.items():
        subject, method = intelligent_classify(content)
        
        if subject and subject in subject_questions:
            subject_questions[subject][q_num] = content
            # 统计分类方法
            method_type = method.split(':')[0]
            classification_methods[subject][method_type] = classification_methods[subject].get(method_type, 0) + 1
        else:
            unclassified[q_num] = content
        
        # 显示前10题的分类结果作为示例
        if q_num <= 10:
            print(f"题目 {q_num}: {subject or '未分类'} ({method})")
    
    # 生成各科目文档
    print("\n生成科目文档...")
    for subject, questions in subject_questions.items():
        if questions:
            filename = create_subject_document(subject, questions, output_dir, classification_methods[subject])
            print(f"{subject}: {len(questions)} 道题 -> {filename}")
    
    # 如果有未分类的题目，创建单独文档
    if unclassified:
        filename = create_subject_document("未分类", unclassified, output_dir, {})
        print(f"未分类: {len(unclassified)} 道题 -> {filename}")
    
    # 输出详细统计信息
    print("\n=== 智能分类统计 ===")
    total_classified = sum(len(q) for q in subject_questions.values())
    for subject, questions in subject_questions.items():
        if questions:
            print(f"{subject}: {len(questions)} 道题")
            methods = classification_methods[subject]
            if methods:
                method_info = ', '.join([f"{k}:{v}" for k, v in methods.items()])
                print(f"  分类依据: {method_info}")
    
    print(f"\n总计: 已分类 {total_classified} 道, 未分类 {len(unclassified)} 道")
    print("\n智能分类完成！")

if __name__ == "__main__":
    main()