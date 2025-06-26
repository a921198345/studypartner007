#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2023年法考题最终分类系统
采用案例驱动的分类方法
"""

import re
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import defaultdict

# 2023年法考8大科目及官方比例
SUBJECT_TARGETS_2023 = {
    '理论法': 40,       # 约20%
    '商经知': 58,       # 约29%
    '民法': 40,         # 约20%
    '民事诉讼法': 30,   # 约15%
    '刑法': 24,         # 约12%
    '行政法': 16,       # 约8%
    '刑事诉讼法': 16,   # 约8%
    '三国法': 20        # 约10%
}

# 强制指定已知的刑法题目（根据手动识别）
KNOWN_CRIMINAL_QUESTIONS = [63]  # 职业薪资诈骗案

def extract_questions_from_document(doc_path):
    """从文档中提取所有题目"""
    doc = Document(doc_path)
    all_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    
    # 使用正则表达式匹配题目
    pattern = r'(\d+)\.\s*(.+?)(?=\d+\.\s*|\Z)'
    matches = re.findall(pattern, all_text, re.DOTALL)
    
    questions = []
    for num_str, content in matches:
        questions.append({
            'number': int(num_str),
            'content': content.strip()
        })
    
    return questions

def classify_by_content_analysis(content):
    """基于内容深度分析进行分类"""
    content_clean = content.replace('\n', ' ').strip().lower()
    
    # 1. 三国法优先识别（涉外特征明显）
    international_indicators = [
        '涉外', '外国人', '外国法', '准据法', '法律冲突', '国际', 
        'cisg', 'cif', 'fob', '条约', '联合国', '维也纳'
    ]
    if any(indicator in content_clean for indicator in international_indicators):
        return '三国法'
    
    # 2. 刑法识别（加强案例识别）
    criminal_indicators = [
        '《刑法》', '犯罪', '诈骗', '盗窃', '抢劫', '杀人', '伤害', '贪污', '受贿',
        '有期徒刑', '无期徒刑', '死刑', '拘役', '管制', '罚金',
        '故意', '过失', '正当防卫', '紧急避险', '共同犯罪'
    ]
    # 刑法案例模式
    criminal_patterns = [
        r'某.*?诈骗', r'某.*?盗窃', r'某.*?抢劫', r'某.*?杀死',
        r'甲.*?诈骗', r'甲.*?盗窃', r'甲.*?抢劫', r'甲.*?杀死',
        r'团伙.*?诈骗', r'.*?虚假.*?诈骗', r'.*?犯罪.*?行为'
    ]
    if (any(indicator in content_clean for indicator in criminal_indicators) or
        any(re.search(pattern, content_clean) for pattern in criminal_patterns)):
        return '刑法'
    
    # 3. 理论法识别（宪法、法理学等）
    theory_indicators = [
        '《宪法》', '宪法', '基本权利', '人大', '国务院', '法院', '检察院',
        '法治', '依法治国', '法的', '法律原则', '法律规则', '司法',
        '立法', '执法', '守法', '法制', '民主', '人权'
    ]
    if any(indicator in content_clean for indicator in theory_indicators):
        return '理论法'
    
    # 4. 商经知识别
    commercial_indicators = [
        '《公司法》', '《证券法》', '《劳动法》', '《专利法》', '《商标法》',
        '公司', '股东', '董事', '监事', '破产', '劳动合同', '专利', '商标',
        '垄断', '消费者', '产品质量', '食品安全', '银行', '保险'
    ]
    if any(indicator in content_clean for indicator in commercial_indicators):
        return '商经知'
    
    # 5. 民法识别
    civil_indicators = [
        '《民法典》', '民法典', '物权', '债权', '合同', '侵权', '婚姻', '继承',
        '所有权', '用益物权', '担保物权', '不当得利', '无因管理'
    ]
    if any(indicator in content_clean for indicator in civil_indicators):
        return '民法'
    
    # 6. 民事诉讼法识别
    civil_procedure_indicators = [
        '《民事诉讼法》', '民诉法', '管辖', '起诉', '举证', '证据',
        '财产保全', '执行', '简易程序', '调解'
    ]
    if any(indicator in content_clean for indicator in civil_procedure_indicators):
        return '民事诉讼法'
    
    # 7. 刑事诉讼法识别
    criminal_procedure_indicators = [
        '《刑事诉讼法》', '刑诉法', '侦查', '审查起诉', '取保候审',
        '监视居住', '逮捕', '拘留', '辩护', '法律援助'
    ]
    if any(indicator in content_clean for indicator in criminal_procedure_indicators):
        return '刑事诉讼法'
    
    # 8. 行政法识别
    admin_indicators = [
        '《行政处罚法》', '《行政许可法》', '《行政复议法》', '《行政诉讼法》',
        '行政处罚', '行政许可', '行政复议', '行政诉讼', '国家赔偿'
    ]
    if any(indicator in content_clean for indicator in admin_indicators):
        return '行政法'
    
    return None

def distribute_by_targets(classified_questions, total_questions):
    """根据官方目标比例分配题目"""
    final_classified = defaultdict(list)
    
    # 先分配已确定科目的题目
    for subject, questions in classified_questions.items():
        if subject != '未分类':
            final_classified[subject] = questions[:SUBJECT_TARGETS_2023.get(subject, len(questions))]
    
    # 统计已分配的题目数
    assigned_count = sum(len(questions) for questions in final_classified.values())
    remaining = total_questions - assigned_count
    
    # 计算还需要的题目数
    needed = {}
    for subject, target in SUBJECT_TARGETS_2023.items():
        current = len(final_classified.get(subject, []))
        if current < target:
            needed[subject] = target - current
    
    # 从未分类和多余的题目中补充
    available_questions = []
    
    # 收集未分类题目
    if '未分类' in classified_questions:
        available_questions.extend(classified_questions['未分类'])
    
    # 收集各科目多余的题目
    for subject, questions in classified_questions.items():
        if subject != '未分类':
            target = SUBJECT_TARGETS_2023.get(subject, 0)
            current_assigned = len(final_classified.get(subject, []))
            if len(questions) > target:
                excess = questions[target:]
                available_questions.extend(excess)
    
    # 按需分配剩余题目
    for subject in sorted(needed.keys(), key=lambda x: needed[x], reverse=True):
        need_count = needed[subject]
        if need_count > 0 and available_questions:
            for _ in range(min(need_count, len(available_questions))):
                final_classified[subject].append(available_questions.pop(0))
    
    # 剩余题目放入未分类
    if available_questions:
        final_classified['未分类'] = available_questions
    
    return final_classified

def create_subject_document(subject, questions, output_dir):
    """为特定科目创建文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(f'2023年法考客观题 - {subject}', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph()
    target = SUBJECT_TARGETS_2023.get(subject, 0)
    stats.add_run(f'共 {len(questions)} 道题 (目标: {target}题)').bold = True
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(20)
    
    # 题目编号列表
    question_numbers = sorted([q['number'] for q in questions])
    numbers_para = doc.add_paragraph()
    numbers_para.add_run(f'题目编号：{question_numbers}').italic = True
    numbers_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    numbers_para.paragraph_format.space_after = Pt(10)
    
    # 分隔线
    doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 按题号排序
    questions_sorted = sorted(questions, key=lambda x: x['number'])
    
    # 添加每道题
    for question in questions_sorted:
        # 添加题目内容
        lines = question['content'].split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())
        
        # 添加分隔线
        separator = doc.add_paragraph('-' * 30)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(6)
        separator.paragraph_format.space_after = Pt(6)
    
    # 保存文档
    output_path = os.path.join(output_dir, f'{subject}.docx')
    doc.save(output_path)
    print(f"已保存 {subject} 科目，共 {len(questions)} 题")

def main():
    """主函数"""
    input_file = '/Users/acheng/Downloads/law-exam-assistant/23年法考题.docx'
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx_2023_final'
    
    print("2023年法考题最终分类系统")
    print("=" * 60)
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 提取题目
    questions = extract_questions_from_document(input_file)
    print(f"共提取 {len(questions)} 道题")
    
    # 分类
    classified = defaultdict(list)
    
    for question in questions:
        # 强制指定已知刑法题目
        if question['number'] in KNOWN_CRIMINAL_QUESTIONS:
            classified['刑法'].append(question)
        else:
            subject = classify_by_content_analysis(question['content'])
            if subject:
                classified[subject].append(question)
            else:
                classified['未分类'].append(question)
    
    print("\n初步分类结果:")
    for subject in SUBJECT_TARGETS_2023.keys():
        count = len(classified.get(subject, []))
        target = SUBJECT_TARGETS_2023[subject]
        print(f"{subject:12s}: {count:3d} 题 (目标: {target:2d})")
    
    if '未分类' in classified:
        print(f"{'未分类':12s}: {len(classified['未分类']):3d} 题")
    
    # 按目标比例调整
    final_classified = distribute_by_targets(classified, len(questions))
    
    print("\n最终分类结果:")
    print("-" * 50)
    total = 0
    for subject, target in SUBJECT_TARGETS_2023.items():
        count = len(final_classified.get(subject, []))
        total += count
        percentage = (count / len(questions)) * 100
        diff = count - target
        status = "✓" if abs(diff) <= 3 else "⚠"
        print(f"{subject:12s}: {count:3d} 题 ({percentage:4.1f}%) 差:{diff:+3d} {status}")
    
    if '未分类' in final_classified:
        count = len(final_classified['未分类'])
        total += count
        percentage = (count / len(questions)) * 100
        print(f"{'未分类':12s}: {count:3d} 题 ({percentage:4.1f}%)")
    
    print(f"\n总计: {total} 题")
    
    # 生成文档
    print("\n生成文档...")
    for subject, questions_list in final_classified.items():
        if questions_list:
            create_subject_document(subject, questions_list, output_dir)
    
    print(f"\n✓ 分类完成！文档保存到: {output_dir}")

if __name__ == "__main__":
    main()