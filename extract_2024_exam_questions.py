#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2024年法考真题提取脚本
用于从扫描版Word文档中提取完整的题目内容
"""

import os
import sys
import re
import json
from datetime import datetime

# 尝试不同的方法读取docx文件
try:
    import docx
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx未安装，将尝试其他方法")

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False
    print("警告: mammoth未安装")

try:
    from pydocx import PyDocX
    HAS_PYDOCX = True
except ImportError:
    HAS_PYDOCX = False
    print("警告: pydocx未安装")

try:
    import zipfile
    import xml.etree.ElementTree as ET
    HAS_XML = True
except ImportError:
    HAS_XML = False
    print("警告: xml解析库未安装")


class ExamQuestionExtractor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.questions = []
        self.raw_text = ""
        
    def extract_with_docx(self):
        """使用python-docx提取内容"""
        if not HAS_DOCX:
            return False
            
        try:
            print("尝试使用python-docx读取...")
            doc = Document(self.file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            self.raw_text = '\n'.join(paragraphs)
            print(f"成功提取 {len(paragraphs)} 个段落")
            return True
            
        except Exception as e:
            print(f"python-docx读取失败: {e}")
            return False
    
    def extract_with_mammoth(self):
        """使用mammoth提取内容"""
        if not HAS_MAMMOTH:
            return False
            
        try:
            print("尝试使用mammoth读取...")
            with open(self.file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                self.raw_text = result.value
                print(f"成功提取 {len(self.raw_text)} 字符")
                return True
                
        except Exception as e:
            print(f"mammoth读取失败: {e}")
            return False
    
    def extract_with_xml(self):
        """直接解析docx的XML内容"""
        if not HAS_XML:
            return False
            
        try:
            print("尝试使用XML解析读取...")
            with zipfile.ZipFile(self.file_path, 'r') as docx:
                # 读取document.xml
                xml_content = docx.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                # 定义命名空间
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # 提取所有文本
                paragraphs = []
                for p in tree.findall('.//w:p', namespaces):
                    texts = []
                    for t in p.findall('.//w:t', namespaces):
                        if t.text:
                            texts.append(t.text)
                    if texts:
                        paragraphs.append(''.join(texts))
                
                self.raw_text = '\n'.join(paragraphs)
                print(f"成功从XML提取 {len(paragraphs)} 个段落")
                return True
                
        except Exception as e:
            print(f"XML解析失败: {e}")
            return False
    
    def parse_questions(self):
        """解析提取的文本，识别题目结构"""
        if not self.raw_text:
            print("没有可解析的文本")
            return
        
        # 题目模式：寻找题号和选项
        lines = self.raw_text.split('\n')
        current_question = None
        current_options = []
        question_buffer = []
        
        # 多种题号模式
        question_patterns = [
            r'^(\d+)[．.、]\s*(.+)',  # 1. 或 1、
            r'^第(\d+)题[：:]\s*(.+)',  # 第1题：
            r'^题目(\d+)[：:]\s*(.+)',  # 题目1：
            r'^(\d+)\s+(.+)',  # 1 题干
        ]
        
        # 选项模式
        option_patterns = [
            r'^([A-D])[．.、]\s*(.+)',  # A. 或 A、
            r'^([A-D])[：:]\s*(.+)',     # A：
            r'^([A-D])\s+(.+)',          # A 选项内容
        ]
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # 检查是否是新题目
            is_new_question = False
            for pattern in question_patterns:
                match = re.match(pattern, line)
                if match:
                    # 保存上一题
                    if current_question and len(current_options) >= 4:
                        self.questions.append({
                            'number': current_question,
                            'stem': ' '.join(question_buffer),
                            'options': current_options[:4]  # 只取ABCD四个选项
                        })
                    
                    # 开始新题目
                    current_question = match.group(1)
                    question_buffer = [match.group(2) if len(match.groups()) > 1 else '']
                    current_options = []
                    is_new_question = True
                    break
            
            if is_new_question:
                i += 1
                continue
            
            # 检查是否是选项
            is_option = False
            for pattern in option_patterns:
                match = re.match(pattern, line)
                if match:
                    option_letter = match.group(1)
                    option_content = match.group(2)
                    current_options.append({
                        'letter': option_letter,
                        'content': option_content
                    })
                    is_option = True
                    break
            
            if not is_option and current_question and len(current_options) < 4:
                # 这可能是题干的延续
                question_buffer.append(line)
            
            i += 1
        
        # 保存最后一题
        if current_question and len(current_options) >= 4:
            self.questions.append({
                'number': current_question,
                'stem': ' '.join(question_buffer),
                'options': current_options[:4]
            })
        
        print(f"成功解析 {len(self.questions)} 道题目")
    
    def enhanced_parse(self):
        """增强版解析，处理更复杂的格式"""
        if not self.raw_text:
            return
        
        # 使用更宽松的分割策略
        # 首先尝试按题号分割整个文本
        question_split_pattern = r'\n+(?=\d+[．.、\s])'
        raw_questions = re.split(question_split_pattern, self.raw_text)
        
        for raw_q in raw_questions:
            if not raw_q.strip():
                continue
            
            # 提取题号
            number_match = re.match(r'^(\d+)[．.、\s]', raw_q)
            if not number_match:
                continue
            
            question_number = number_match.group(1)
            
            # 提取选项
            # 先找到所有选项的位置
            option_positions = []
            for letter in ['A', 'B', 'C', 'D']:
                patterns = [
                    f'{letter}[．.、]',
                    f'{letter}[：:]',
                    f'{letter}\\s+'
                ]
                for pattern in patterns:
                    matches = list(re.finditer(pattern, raw_q))
                    if matches:
                        option_positions.extend([(m.start(), letter, m.end()) for m in matches])
            
            if len(option_positions) < 4:
                continue
            
            # 按位置排序
            option_positions.sort(key=lambda x: x[0])
            
            # 提取题干（第一个选项之前的内容）
            first_option_pos = option_positions[0][0]
            stem = raw_q[number_match.end():first_option_pos].strip()
            
            # 提取各选项内容
            options = []
            for i, (start, letter, end) in enumerate(option_positions):
                if letter not in ['A', 'B', 'C', 'D']:
                    continue
                
                # 找下一个选项的开始位置
                if i + 1 < len(option_positions):
                    next_start = option_positions[i + 1][0]
                else:
                    next_start = len(raw_q)
                
                option_content = raw_q[end:next_start].strip()
                options.append({
                    'letter': letter,
                    'content': option_content
                })
            
            if len(options) >= 4:
                self.questions.append({
                    'number': question_number,
                    'stem': stem,
                    'options': options[:4]
                })
    
    def extract_all(self):
        """尝试所有方法提取内容"""
        # 尝试各种提取方法
        success = False
        
        if self.extract_with_docx():
            success = True
        elif self.extract_with_mammoth():
            success = True
        elif self.extract_with_xml():
            success = True
        else:
            print("所有提取方法都失败了")
            return False
        
        # 解析题目
        self.parse_questions()
        
        # 如果基础解析没有找到足够的题目，尝试增强解析
        if len(self.questions) < 10:
            print("基础解析结果不足，尝试增强解析...")
            self.questions = []  # 重置
            self.enhanced_parse()
        
        return len(self.questions) > 0
    
    def save_results(self):
        """保存提取结果"""
        # 保存JSON格式
        json_file = self.file_path.replace('.docx', '_extracted.json')
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump({
                'source': self.file_path,
                'extract_time': datetime.now().isoformat(),
                'total_questions': len(self.questions),
                'questions': self.questions
            }, f, ensure_ascii=False, indent=2)
        
        # 保存可读格式
        txt_file = self.file_path.replace('.docx', '_extracted.txt')
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(f"2024年法考真题提取结果\n")
            f.write(f"提取时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"总题数: {len(self.questions)}\n")
            f.write("=" * 80 + "\n\n")
            
            for q in self.questions:
                f.write(f"第{q['number']}题\n")
                f.write(f"题干: {q['stem']}\n")
                f.write("选项:\n")
                for opt in q['options']:
                    f.write(f"  {opt['letter']}. {opt['content']}\n")
                f.write("\n" + "-" * 60 + "\n\n")
        
        # 保存原始文本供调试
        raw_file = self.file_path.replace('.docx', '_raw.txt')
        with open(raw_file, 'w', encoding='utf-8') as f:
            f.write(self.raw_text)
        
        print(f"\n提取完成！")
        print(f"- JSON格式: {json_file}")
        print(f"- 文本格式: {txt_file}")
        print(f"- 原始文本: {raw_file}")
        print(f"- 共提取 {len(self.questions)} 道题目")


def main():
    file_path = "/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）_扫描版.docx"
    
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        return
    
    print(f"开始提取: {file_path}")
    print("=" * 80)
    
    extractor = ExamQuestionExtractor(file_path)
    
    if extractor.extract_all():
        extractor.save_results()
        
        # 显示前3道题作为示例
        print("\n\n提取示例（前3题）:")
        print("=" * 80)
        for i, q in enumerate(extractor.questions[:3]):
            print(f"\n第{q['number']}题")
            print(f"题干: {q['stem']}")
            print("选项:")
            for opt in q['options']:
                print(f"  {opt['letter']}. {opt['content']}")
    else:
        print("提取失败，请检查文档格式或安装必要的库")
        print("建议安装: pip install python-docx mammoth pydocx")


if __name__ == "__main__":
    main()