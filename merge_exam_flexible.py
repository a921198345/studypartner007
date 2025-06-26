#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
灵活合并2024年法考真题和答案解析
即使真题不完整，也能根据答案解析生成完整文档
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

def extract_questions_from_doc(doc_path):
    """从文档中提取题目，返回题号和内容的字典"""
    doc = Document(doc_path)
    questions = {}
    
    # 匹配题号的多种正则表达式
    patterns = [
        re.compile(r'^(\d+)[\.、．]\s*(.*)$'),  # 1. 或 1、或 1．
        re.compile(r'^第(\d+)题[：:]\s*(.*)$'),  # 第1题：
        re.compile(r'^【(\d+)】\s*(.*)$'),      # 【1】
    ]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # 尝试所有模式
        for pattern in patterns:
            match = pattern.match(text)
            if match:
                q_num = int(match.group(1))
                # 收集整个题目内容（包括后续段落直到下一题）
                questions[q_num] = text
                break
    
    return questions

def extract_answers_from_doc(doc_path):
    """从答案解析文档中提取答案，返回题号和解析的字典"""
    doc = Document(doc_path)
    answers = {}
    current_question = None
    current_content = []
    
    # 匹配题号的正则表达式（更灵活）
    question_patterns = [
        re.compile(r'^(\d+)[\.、．\s]+(.*)$'),  # 匹配 "1." "1、" "1．" "1 ."等
    ]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目
        is_new_question = False
        for pattern in question_patterns:
            match = pattern.match(text)
            if match:
                # 保存前一题
                if current_question:
                    answers[current_question] = '\n'.join(current_content)
                
                # 开始新题目
                current_question = int(match.group(1))
                current_content = [text]
                is_new_question = True
                break
        
        if not is_new_question and current_question:
            # 继续当前题目的内容
            current_content.append(text)
    
    # 保存最后一题
    if current_question:
        answers[current_question] = '\n'.join(current_content)
    
    return answers

def add_styled_heading(doc, text, level=1, color=None):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_question_section(doc, num, question_text=None, answer_text=None):
    """添加一个题目和答案的完整部分"""
    # 题目部分
    if question_text:
        # 如果有真题内容
        question_para = doc.add_paragraph()
        question_para.add_run(question_text).bold = True
    else:
        # 如果没有真题，从答案中提取
        question_para = doc.add_paragraph()
        question_para.add_run(f"{num}. [原题暂缺，请参考答案解析]").bold = True
        question_para.paragraph_format.space_after = Pt(6)
    
    # 添加答案解析
    if answer_text:
        # 答案标题
        answer_title = doc.add_paragraph()
        run = answer_title.add_run("【答案解析】")
        run.font.color.rgb = RGBColor(0, 112, 192)  # 深蓝色
        run.bold = True
        answer_title.paragraph_format.space_before = Pt(6)
        
        # 答案内容
        answer_lines = answer_text.split('\n')
        for line in answer_lines:
            if line.strip():
                answer_para = doc.add_paragraph(line)
                answer_para.paragraph_format.left_indent = Pt(20)
                answer_para.paragraph_format.space_after = Pt(3)
    else:
        # 如果没有答案
        no_answer = doc.add_paragraph()
        run = no_answer.add_run("【答案解析】暂无")
        run.font.color.rgb = RGBColor(192, 0, 0)  # 红色
        no_answer.paragraph_format.space_before = Pt(6)
    
    # 添加分隔线
    separator = doc.add_paragraph("─" * 80)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    separator.paragraph_format.space_before = Pt(12)
    separator.paragraph_format.space_after = Pt(12)

def merge_questions_and_answers(questions, answers, output_path):
    """合并题目和答案，生成新文档"""
    doc = Document()
    
    # 设置文档的默认字体
    doc.styles['Normal'].font.name = '宋体'
    
    # 添加标题
    add_styled_heading(doc, '2024年法考真题（含答案解析）', level=1)
    
    # 添加说明
    info_para = doc.add_paragraph()
    info_para.add_run("说明：本文档整合了2024年法考客观题真题和答案解析。部分题目原文可能暂缺，请以答案解析为参考。")
    info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info_para.paragraph_format.space_after = Pt(20)
    
    # 统计信息
    all_question_numbers = sorted(set(questions.keys()) | set(answers.keys()))
    total_questions = len(all_question_numbers)
    questions_with_text = len(questions)
    questions_with_answers = len(answers)
    
    # 添加统计信息
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"共收录 {total_questions} 道题目，其中 {questions_with_text} 道有原题，{questions_with_answers} 道有答案解析")
    stats_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats_para.paragraph_format.space_after = Pt(20)
    
    # 添加分割线
    doc.add_paragraph("═" * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    print(f"\n开始生成合并文档...")
    print(f"总题数: {total_questions}")
    print(f"有原题的: {questions_with_text}")
    print(f"有答案的: {questions_with_answers}")
    
    # 逐题添加
    for i, num in enumerate(all_question_numbers, 1):
        question_text = questions.get(num)
        answer_text = answers.get(num)
        
        # 显示进度
        if i % 10 == 0:
            print(f"处理进度: {i}/{total_questions}")
        
        add_question_section(doc, num, question_text, answer_text)
    
    # 保存文档
    doc.save(output_path)
    print(f"\n合并完成！文件已保存到: {output_path}")
    
    # 输出缺失信息
    questions_without_answers = set(questions.keys()) - set(answers.keys())
    if questions_without_answers:
        print(f"\n有题目但无答案的题号: {sorted(questions_without_answers)}")
    
    answers_without_questions = set(answers.keys()) - set(questions.keys())
    if answers_without_questions:
        print(f"\n有答案但无题目的题号共 {len(answers_without_questions)} 道")

def main():
    # 文件路径
    questions_path = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）.docx'
    answers_path = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_path = '/Users/acheng/Downloads/law-exam-assistant/24年法考题.docx'
    
    # 检查文件是否存在
    if not os.path.exists(questions_path):
        print(f"错误：找不到真题文件 {questions_path}")
        return
    
    if not os.path.exists(answers_path):
        print(f"错误：找不到答案解析文件 {answers_path}")
        return
    
    print("开始提取真题...")
    questions = extract_questions_from_doc(questions_path)
    print(f"提取到 {len(questions)} 道真题")
    
    print("\n开始提取答案解析...")
    answers = extract_answers_from_doc(answers_path)
    print(f"提取到 {len(answers)} 道题目的答案解析")
    
    # 合并文档
    merge_questions_and_answers(questions, answers, output_path)

if __name__ == "__main__":
    main()