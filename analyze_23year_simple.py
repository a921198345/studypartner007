#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简单分析23年法考题文档
"""

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

import os

def analyze_document_simple(doc_path):
    """简单分析文档内容"""
    if not HAS_DOCX:
        print("python-docx 未安装，无法分析 docx 文档")
        return
    
    if not os.path.exists(doc_path):
        print(f"文件不存在: {doc_path}")
        return
    
    try:
        doc = Document(doc_path)
        print(f"文档总段落数: {len(doc.paragraphs)}")
        
        # 收集所有非空段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        print(f"非空段落数: {len(paragraphs)}")
        
        # 显示前20个段落
        print("\n前20个段落:")
        for i, text in enumerate(paragraphs[:20]):
            print(f"{i+1:2d}. {text[:80]}{'...' if len(text) > 80 else ''}")
        
        # 查找题目模式
        print("\n寻找题目模式...")
        question_patterns = []
        for i, text in enumerate(paragraphs):
            # 检查是否以数字开头
            if text and text[0].isdigit():
                # 找到第一个点号
                dot_pos = -1
                for j, char in enumerate(text):
                    if char in '.．。':
                        dot_pos = j
                        break
                
                if dot_pos > 0 and dot_pos < 5:
                    try:
                        q_num = int(text[:dot_pos])
                        question_patterns.append((i, q_num, text[:100]))
                    except:
                        pass
        
        print(f"找到疑似题目: {len(question_patterns)} 个")
        for i, (line_num, q_num, preview) in enumerate(question_patterns[:10]):
            print(f"题目{q_num}: {preview}...")
        
        # 查找答案模式
        print("\n寻找答案模式...")
        answer_patterns = []
        for i, text in enumerate(paragraphs):
            if '正确答案' in text or '【答案解析】' in text or '【解析】' in text:
                answer_patterns.append((i, text[:100]))
        
        print(f"找到疑似答案: {len(answer_patterns)} 个")
        for i, (line_num, preview) in enumerate(answer_patterns[:5]):
            print(f"答案{i+1}: {preview}...")
        
        # 统计关键词
        print("\n关键词统计...")
        full_text = ' '.join(paragraphs).lower()
        
        keywords = {
            '民法': ['民法', '合同', '物权', '债权', '侵权', '婚姻', '继承'],
            '刑法': ['刑法', '犯罪', '刑罚', '故意', '盗窃', '诈骗', '抢劫'],
            '商经知': ['公司法', '公司', '股东', '董事', '证券', '保险', '劳动法'],
            '理论法': ['宪法', '法理', '法治', '立法', '司法', '基本权利'],
            '诉讼法': ['诉讼', '管辖', '证据', '审判', '执行', '起诉', '上诉']
        }
        
        for subject, words in keywords.items():
            count = sum(full_text.count(word) for word in words)
            print(f"{subject}: {count} 次")
        
    except Exception as e:
        print(f"分析文档时出错: {e}")

def analyze_txt_format(file_path):
    """分析文本格式的文档"""
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        return
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        print(f"文本文件总行数: {len(lines)}")
        
        # 显示前20行
        print("\n前20行:")
        for i, line in enumerate(lines[:20]):
            print(f"{i+1:2d}. {line[:80]}{'...' if len(line) > 80 else ''}")
        
    except Exception as e:
        print(f"读取文本文件时出错: {e}")

if __name__ == "__main__":
    # 分析23年法考题文档
    doc_path = '/Users/acheng/Downloads/law-exam-assistant/23年法考题.docx'
    print("=== 分析23年法考题.docx ===")
    analyze_document_simple(doc_path)
    
    # 如果有其他相关文档也分析
    other_docs = [
        '/Users/acheng/Downloads/law-exam-assistant/2023法考 客观回忆版真题.docx',
        '/Users/acheng/Downloads/law-exam-assistant/2023法考 客观回忆版真题-答案及解析.docx'
    ]
    
    for doc_path in other_docs:
        if os.path.exists(doc_path):
            print(f"\n=== 分析 {os.path.basename(doc_path)} ===")
            analyze_document_simple(doc_path)