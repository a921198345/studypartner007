#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析文档结构，找出题目格式
"""

from docx import Document
import re

def analyze_document(doc_path, doc_name):
    """分析文档结构"""
    print(f"\n{'='*60}")
    print(f"分析文档: {doc_name}")
    print(f"文件路径: {doc_path}")
    print(f"{'='*60}")
    
    try:
        doc = Document(doc_path)
        total_paras = len(doc.paragraphs)
        print(f"总段落数: {total_paras}")
        
        # 统计不同格式的题号
        patterns = [
            (r'^(\d+)[\.、．]\s*(.*)$', "数字+标点"),
            (r'^第(\d+)题[：:]\s*(.*)$', "第X题格式"),
            (r'^【(\d+)】\s*(.*)$', "【数字】格式"),
            (r'^(\d+)\s+(.*)$', "数字+空格"),
            (r'^\((\d+)\)\s*(.*)$', "(数字)格式"),
            (r'^试题(\d+)[：:]\s*(.*)$', "试题X格式"),
        ]
        
        # 查找匹配的题号
        found_questions = {}
        for pattern_str, pattern_name in patterns:
            pattern = re.compile(pattern_str)
            matches = []
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text and pattern.match(text):
                    match = pattern.match(text)
                    q_num = int(match.group(1))
                    matches.append((q_num, i, text[:50]))
            
            if matches:
                found_questions[pattern_name] = matches
                print(f"\n找到 {pattern_name} 格式的题目: {len(matches)} 个")
                print(f"题号范围: {min(m[0] for m in matches)} - {max(m[0] for m in matches)}")
                print("前3个示例:")
                for q_num, line_num, text in matches[:3]:
                    print(f"  第{line_num}行: {text}...")
        
        # 如果没有找到标准题号，显示前20个非空段落
        if not found_questions:
            print("\n未找到标准题号格式，显示前20个非空段落:")
            non_empty_paras = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
            for i, (line_num, text) in enumerate(non_empty_paras[:20]):
                print(f"第{line_num}行: {text[:80]}...")
        
        # 检查表格
        if doc.tables:
            print(f"\n文档包含 {len(doc.tables)} 个表格")
            for i, table in enumerate(doc.tables[:3]):
                print(f"\n表格{i+1}: {len(table.rows)}行 x {len(table.columns)}列")
                # 显示表格前几个单元格的内容
                for row_idx in range(min(3, len(table.rows))):
                    row_text = []
                    for col_idx in range(min(3, len(table.columns))):
                        cell_text = table.rows[row_idx].cells[col_idx].text.strip()[:20]
                        row_text.append(cell_text)
                    print(f"  第{row_idx+1}行: {' | '.join(row_text)}")
                    
    except Exception as e:
        print(f"分析文档时出错: {e}")

def main():
    # 分析两个文档
    questions_path = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）.docx'
    answers_path = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    
    analyze_document(questions_path, "真题文档")
    analyze_document(answers_path, "答案解析文档")

if __name__ == "__main__":
    main()