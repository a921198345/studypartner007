#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合并2024年法考真题和答案解析
将答案解析按题号匹配到对应的真题后面
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def extract_questions_from_doc(doc_path):
    """从文档中提取题目，返回题号和内容的字典"""
    doc = Document(doc_path)
    questions = {}
    current_question = None
    current_content = []
    
    # 匹配题号的正则表达式（支持多种格式）
    question_pattern = re.compile(r'^(\d+)[\.、．]\s*(.*)$')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # 检查是否是新题目
        match = question_pattern.match(text)
        if match:
            # 保存前一题
            if current_question:
                questions[current_question] = '\n'.join(current_content)
            
            # 开始新题目
            current_question = int(match.group(1))
            current_content = [text]
        else:
            # 继续当前题目的内容
            if current_question:
                current_content.append(text)
    
    # 保存最后一题
    if current_question:
        questions[current_question] = '\n'.join(current_content)
    
    return questions

def extract_answers_from_doc(doc_path):
    """从答案解析文档中提取答案，返回题号和解析的字典"""
    doc = Document(doc_path)
    answers = {}
    current_question = None
    current_content = []
    
    # 匹配题号的正则表达式
    question_pattern = re.compile(r'^(\d+)[\.、．]\s*(.*)$')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # 检查是否是新题目的答案
        match = question_pattern.match(text)
        if match:
            # 保存前一题的答案
            if current_question:
                answers[current_question] = '\n'.join(current_content)
            
            # 开始新题目的答案
            current_question = int(match.group(1))
            current_content = [text]
        else:
            # 继续当前答案的内容
            if current_question:
                current_content.append(text)
    
    # 保存最后一题的答案
    if current_question:
        answers[current_question] = '\n'.join(current_content)
    
    return answers

def merge_questions_and_answers(questions, answers, output_path):
    """合并题目和答案，生成新文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('2024年法考真题（含答案解析）', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 获取所有题号并排序
    all_question_numbers = sorted(set(questions.keys()) | set(answers.keys()))
    
    print(f"共找到 {len(all_question_numbers)} 道题目")
    
    for num in all_question_numbers:
        # 添加题目
        if num in questions:
            # 题目段落
            question_para = doc.add_paragraph()
            question_para.add_run(questions[num]).bold = True
            
            # 添加空行
            doc.add_paragraph()
        else:
            # 如果没有找到对应题目，添加提示
            para = doc.add_paragraph()
            para.add_run(f"{num}. [题目缺失]").bold = True
            doc.add_paragraph()
        
        # 添加答案解析
        if num in answers:
            # 答案标题
            answer_title = doc.add_paragraph()
            run = answer_title.add_run("【答案解析】")
            run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
            run.bold = True
            
            # 答案内容
            answer_para = doc.add_paragraph(answers[num])
            answer_para.paragraph_format.left_indent = Pt(20)  # 缩进
            
            # 添加分隔线
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
        else:
            # 如果没有找到对应答案，添加提示
            para = doc.add_paragraph()
            run = para.add_run("【答案解析】暂无")
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
    
    # 保存文档
    doc.save(output_path)
    print(f"合并完成！文件已保存到: {output_path}")

def main():
    # 文件路径
    questions_path = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）.docx'
    answers_path = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_path = '/Users/acheng/Downloads/law-exam-assistant/24年法考题.docx'
    
    # 检查文件是否存在
    if not os.path.exists(questions_path):
        print(f"错误：找不到真题文件 {questions_path}")
        return
    
    if not os.path.exists(answers_path):
        print(f"错误：找不到答案解析文件 {answers_path}")
        return
    
    print("开始提取真题...")
    questions = extract_questions_from_doc(questions_path)
    print(f"提取到 {len(questions)} 道真题")
    
    print("\n开始提取答案解析...")
    answers = extract_answers_from_doc(answers_path)
    print(f"提取到 {len(answers)} 道题目的答案解析")
    
    # 显示匹配情况
    matched = len(set(questions.keys()) & set(answers.keys()))
    print(f"\n成功匹配 {matched} 道题目")
    
    questions_only = set(questions.keys()) - set(answers.keys())
    if questions_only:
        print(f"只有题目没有答案的题号: {sorted(questions_only)}")
    
    answers_only = set(answers.keys()) - set(questions.keys())
    if answers_only:
        print(f"只有答案没有题目的题号: {sorted(answers_only)}")
    
    print("\n开始合并文档...")
    merge_questions_and_answers(questions, answers, output_path)

if __name__ == "__main__":
    main()