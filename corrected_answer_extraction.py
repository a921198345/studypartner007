#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
修正后的答案提取函数
直接替换你合并脚本中的答案提取部分
"""

import re
from docx import Document

def extract_answers_from_docx(answer_docx_path):
    """
    从答案文档中提取答案 - 修正版本
    
    修正内容：
    1. 使用正确的正则表达式模式
    2. 兼容中英文冒号
    3. 处理空格变化
    4. 支持多选题
    """
    answers = {}
    
    try:
        doc = Document(answer_docx_path)
        
        # 修正后的答案匹配模式
        # 格式：数字.正确答案： 字母(们)
        answer_pattern = r'^(\d+)\.正确答案[：:]\s*([ABCD]+)$'
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            match = re.match(answer_pattern, text)
            if match:
                question_num = int(match.group(1))
                answer = match.group(2)
                answers[question_num] = answer
        
        print(f"从答案文档提取到 {len(answers)} 个答案")
        
        # 统计单选和多选题
        single_choice = sum(1 for ans in answers.values() if len(ans) == 1)
        multi_choice = sum(1 for ans in answers.values() if len(ans) > 1)
        print(f"单选题: {single_choice} 个，多选题: {multi_choice} 个")
        
        return answers
        
    except Exception as e:
        print(f"提取答案时出错: {str(e)}")
        return {}

# 如果你的合并脚本中有其他答案格式，可以添加备用模式
def extract_answers_with_fallback(answer_docx_path):
    """
    带备用模式的答案提取函数
    """
    answers = {}
    
    try:
        doc = Document(answer_docx_path)
        
        # 主要模式和备用模式
        patterns = [
            r'^(\d+)\.正确答案[：:]\s*([ABCD]+)$',        # 主要模式
            r'^(\d+)\.\s*正确答案[：:]\s*([ABCD]+)$',     # 题号后有空格
            r'^\d+[.、]\s*([ABCD]+)$',                    # 简化格式：1. A
            r'^(\d+)[.、：:]\s*([ABCD]+)$',               # 更宽松的匹配
        ]
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 尝试各种模式
            for pattern in patterns:
                match = re.match(pattern, text)
                if match:
                    if len(match.groups()) == 2:
                        question_num = int(match.group(1))
                        answer = match.group(2)
                    else:
                        # 处理只有一个捕获组的情况
                        question_num = int(re.match(r'^\d+', text).group())
                        answer = match.group(1)
                    
                    answers[question_num] = answer
                    break
        
        print(f"使用备用模式提取到 {len(answers)} 个答案")
        return answers
        
    except Exception as e:
        print(f"提取答案时出错: {str(e)}")
        return {}

# 测试函数
if __name__ == "__main__":
    test_path = "/Users/acheng/Downloads/2023 年卷一答案.docx"
    
    print("测试标准答案提取:")
    answers1 = extract_answers_from_docx(test_path)
    print(f"提取结果：{len(answers1)} 个答案")
    
    if len(answers1) < 50:  # 如果提取结果不理想
        print("\n尝试备用模式:")
        answers2 = extract_answers_with_fallback(test_path)
        print(f"备用模式结果：{len(answers2)} 个答案")