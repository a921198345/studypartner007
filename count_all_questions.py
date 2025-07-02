#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
统计2024年卷一文档中的所有题目
"""

from docx import Document
import re

def count_all_questions():
    file_path = "/Users/acheng/Downloads/law-exam-assistant/2024年卷一完整版_修复版.docx"
    
    print("🔢 统计文档中的所有题目")
    print("=" * 50)
    
    doc = Document(file_path)
    
    # 查找所有"第X题"格式
    questions = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^第\s*\d+\s*题', text):
            match = re.search(r'第\s*(\d+)\s*题', text)
            if match:
                question_num = int(match.group(1))
                questions.append({
                    'number': question_num,
                    'paragraph_index': i,
                    'text': text
                })
    
    print(f"✅ 找到 {len(questions)} 个题目")
    
    if questions:
        question_numbers = [q['number'] for q in questions]
        question_numbers.sort()
        
        print(f"📊 题目编号范围: {min(question_numbers)} - {max(question_numbers)}")
        
        # 检查是否有缺失的题目
        missing = []
        for i in range(1, 101):  # 1-100题
            if i not in question_numbers:
                missing.append(i)
        
        if missing:
            print(f"⚠️  缺失的题目: {missing}")
        else:
            print(f"✅ 题目完整，1-100题全部存在")
        
        # 显示前10题和后10题
        print(f"\n📋 前10题:")
        for i, q in enumerate(questions[:10]):
            print(f"   {i+1}. 第{q['number']}题 (段落{q['paragraph_index']})")
        
        print(f"\n📋 后10题:")
        for i, q in enumerate(questions[-10:]):
            print(f"   {i+1}. 第{q['number']}题 (段落{q['paragraph_index']})")
    
    return questions

def create_correct_parser():
    """创建正确的解析脚本"""
    print(f"\n🔧 创建正确的解析脚本")
    print("=" * 50)
    
    script_content = '''#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2024年卷一完整版正确解析脚本
适配"第X题"格式
"""

from docx import Document
import re
import json

def parse_2024_paper1_correct():
    """正确解析2024年卷一题目"""
    file_path = "/Users/acheng/Downloads/law-exam-assistant/2024年卷一完整版_修复版.docx"
    
    print("🔄 正确解析2024年卷一题目")
    print("=" * 50)
    
    doc = Document(file_path)
    questions = []
    
    current_question = None
    current_content = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是题目标题 "第X题"
        if re.match(r'^第\\s*\\d+\\s*题', text):
            # 如果之前有题目，先保存
            if current_question:
                questions.append(parse_question_content(current_question, current_content))
            
            # 开始新题目
            match = re.search(r'第\\s*(\\d+)\\s*题', text)
            if match:
                current_question = int(match.group(1))
                current_content = []
        
        elif current_question:
            # 收集题目内容
            current_content.append(text)
            
            # 如果遇到分隔线，表示题目结束
            if '──────' in text:
                questions.append(parse_question_content(current_question, current_content))
                current_question = None
                current_content = []
    
    # 处理最后一题（如果没有分隔线）
    if current_question and current_content:
        questions.append(parse_question_content(current_question, current_content))
    
    print(f"✅ 成功解析 {len(questions)} 个题目")
    
    # 保存结果
    output_file = "2024年卷一题目_正确解析.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)
    
    print(f"💾 结果已保存到: {output_file}")
    return questions

def parse_question_content(question_num, content_lines):
    """解析单个题目的内容"""
    question_data = {
        'number': question_num,
        'question': '',
        'options': [],
        'answer': '',
        'explanation': ''
    }
    
    explanation_parts = []
    collecting_explanation = False
    
    for line in content_lines:
        line = line.strip()
        if not line or '──────' in line:
            continue
            
        # 题目内容（第一行非选项内容）
        if not question_data['question'] and not re.match(r'^[A-D][．.\s]', line) and not line.startswith('【答案解析】') and not line.startswith('正确答案：'):
            question_data['question'] = line
            
        # 选项
        elif re.match(r'^[A-D][．.\s]', line):
            question_data['options'].append(line)
            
        # 答案
        elif line.startswith('正确答案：'):
            question_data['answer'] = line.replace('正确答案：', '').strip()
            
        # 解析内容
        elif line.startswith('【答案解析】'):
            collecting_explanation = True
            
        elif collecting_explanation:
            explanation_parts.append(line)
    
    question_data['explanation'] = ' '.join(explanation_parts)
    return question_data

if __name__ == "__main__":
    questions = parse_2024_paper1_correct()
    
    # 显示前3题作为验证
    print(f"\\n📋 前3题验证:")
    for i, q in enumerate(questions[:3]):
        print(f"\\n题目 {q['number']}:")
        print(f"  题干: {q['question'][:50]}...")
        print(f"  选项数: {len(q['options'])}")
        print(f"  答案: {q['answer']}")
        print(f"  解析长度: {len(q['explanation'])} 字符")
'''
    
    with open('/Users/acheng/Downloads/law-exam-assistant/parse_2024_paper1_correct.py', 'w', encoding='utf-8') as f:
        f.write(script_content)
    
    print("✅ 正确的解析脚本已创建: parse_2024_paper1_correct.py")

if __name__ == "__main__":
    questions = count_all_questions()
    create_correct_parser()