#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2024年法考题智能分类系统（改进版）
特别优化刑法题目的识别
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import defaultdict

# 法考8大科目
SUBJECTS = ['民法', '刑法', '行政法', '民事诉讼法', '刑事诉讼法', '商经知', '三国法', '理论法']

# 改进的法条引用映射（特别加强刑法）
LAW_CITATIONS = {
    '刑法': [
        '《刑法》', '刑法第', '《中华人民共和国刑法》', '刑法条文', '刑法典',
        '《刑法修正案》', '刑法分则', '刑法总则', '《关于办理', '两高'
    ],
    '民法': [
        '《民法典》', '民法典第', '《中华人民共和国民法典》', '《民法通则》',
        '《物权法》', '《合同法》', '《侵权责任法》', '《婚姻法》', '《继承法》'
    ],
    '民事诉讼法': [
        '《民事诉讼法》', '民诉法第', '《中华人民共和国民事诉讼法》', 
        '《最高人民法院关于适用〈中华人民共和国民事诉讼法〉》',
        '民事诉讼法解释', '民诉解释', '《民事诉讼法解释》', '《民诉法解释》'
    ],
    '刑事诉讼法': [
        '《刑事诉讼法》', '刑诉法第', '《中华人民共和国刑事诉讼法》',
        '《最高人民法院关于适用〈中华人民共和国刑事诉讼法〉》',
        '刑事诉讼法解释', '刑诉解释'
    ],
    '行政法': [
        '《行政处罚法》', '《行政许可法》', '《行政强制法》', '《行政复议法》',
        '《行政诉讼法》', '《国家赔偿法》', '《公务员法》', '《政府信息公开条例》'
    ],
    '商经知': [
        '《公司法》', '《证券法》', '《保险法》', '《票据法》', '《破产法》',
        '《合伙企业法》', '《个人独资企业法》', '《反垄断法》', '《消费者权益保护法》',
        '《产品质量法》', '《劳动法》', '《劳动合同法》', '《知识产权法》', '《专利法》',
        '《商标法》', '《著作权法》', '《反不正当竞争法》', '《企业破产法》',
        '《建设工程施工合同解释》', '《公司法解释》'
    ],
    '三国法': [
        '《联合国国际货物销售合同公约》', '《国际私法》', '《涉外民事关系法律适用法》',
        '《海商法》', '《维也纳外交关系公约》', '《维也纳领事关系公约》',
        'INCOTERMS', 'CIF', 'FOB', 'CIP', '《纽约公约》', '《海牙公约》',
        '《出口管制法》'
    ],
    '理论法': [
        '《宪法》', '《立法法》', '《监察法》', '《选举法》', '《代表法》',
        '《法官法》', '《检察官法》', '《律师法》', '《仲裁法》', '宪法修正案',
        '《法理学》', '《法治理论》'
    ]
}

# 改进的专业术语（特别加强刑法）
PROFESSIONAL_TERMS = {
    '刑法': [
        # 犯罪构成
        '犯罪构成', '主观要件', '客观要件', '主体要件', '客体要件',
        '故意', '过失', '明知', '应当预见', '疏忽大意', '过于自信',
        # 具体罪名
        '故意杀人', '过失致人死亡', '故意伤害', '过失致人重伤',
        '强奸', '强制猥亵', '非法拘禁', '绑架', '拐卖',
        '盗窃', '抢劫', '抢夺', '诈骗', '敲诈勒索',
        '贪污', '受贿', '行贿', '挪用公款', '职务侵占',
        '交通肇事', '危险驾驶', '醉驾', '毒驾',
        '放火', '爆炸', '投毒', '危害公共安全',
        '走私', '贩毒', '制毒', '非法持有毒品',
        '聚众斗殴', '寻衅滋事', '妨害公务', '袭警',
        # 刑罚
        '死刑', '无期徒刑', '有期徒刑', '拘役', '管制',
        '罚金', '没收财产', '剥夺政治权利',
        # 量刑情节
        '从重', '从轻', '减轻', '免除处罚',
        '累犯', '自首', '立功', '坦白', '认罪认罚',
        '犯罪预备', '犯罪未遂', '犯罪中止', '犯罪既遂',
        # 特殊制度
        '正当防卫', '防卫过当', '紧急避险', '避险过当',
        '共同犯罪', '主犯', '从犯', '胁从犯', '教唆犯',
        '数罪并罚', '牵连犯', '想象竞合', '法条竞合',
        # 刑法特有表述
        '构成何罪', '如何定罪', '罪名认定', '定罪量刑',
        '承担刑事责任', '追究刑事责任', '刑事处罚'
    ],
    '民法': [
        '法律行为', '民事法律行为', '代理制度', '诉讼时效', '物权变动',
        '善意取得', '不动产登记', '用益物权', '担保物权', '债权债务',
        '不当得利', '无因管理', '缔约过失', '违约责任', '侵权责任',
        '婚姻家庭', '夫妻财产', '监护制度', '继承权', '遗嘱继承',
        '人格权', '肖像权', '隐私权', '名誉权', '收养关系', '离婚协议',
        '买卖合同', '租赁合同', '赠与合同', '所有权', '用益物权', '担保物权',
        '相邻关系', '共有', '遗产继承', '法定继承', '遗嘱', '遗赠'
    ],
    '民事诉讼法': [
        '管辖权异议', '级别管辖', '地域管辖', '专属管辖', '协议管辖',
        '举证责任', '证据规则', '财产保全', '先予执行', '简易程序',
        '小额诉讼', '缺席判决', '第三人撤销之诉', '案外人执行异议',
        '申请再审', '审判监督程序', '强制执行', '执行和解', '发回重审',
        '合议庭', '独任审判', '人民陪审员', '行为保全', '人身安全保护令',
        '诉讼费用', '诉讼代理', '申请复议'
    ],
    '刑事诉讼法': [
        '侦查', '审查起诉', '审判程序', '强制措施', '取保候审',
        '监视居住', '刑事拘留', '逮捕', '搜查', '扣押', '查封', '冻结',
        '辩护权', '法律援助', '证人保护', '技侦措施', '认罪认罚',
        '速裁程序', '简易程序', '死刑复核', '刑事和解', '缺席审判',
        '公诉', '自诉', '附带民事诉讼', '上诉', '抗诉'
    ],
    '行政法': [
        '行政主体', '行政相对人', '具体行政行为', '抽象行政行为',
        '行政许可', '行政处罚', '行政强制', '行政复议', '行政诉讼',
        '政府信息公开', '行政程序', '听证程序', '国家赔偿', '公务员',
        '市场监督管理局', '行政复议决定', '维持决定', '撤销决定',
        '行政机关', '行政裁决', '行政确认'
    ],
    '商经知': [
        '股东大会', '董事会', '监事会', '股权转让', '公司治理',
        '有限责任公司', '股份有限公司', '合伙企业', '个人独资企业',
        '公司设立', '公司解散', '破产重整', '破产清算', '证券发行',
        '票据权利', '票据责任', '保险合同', '保险责任', '劳动合同',
        '社会保险', '工伤保险', '反垄断', '不正当竞争', '消费者权益',
        '股东分红', '股权回购', '破产管理人', '债权申报', '清算程序',
        '合伙协议', '执行事务合伙人', '普通合伙企业', '建设工程合同',
        '知识产权', '专利权', '商标权', '著作权'
    ],
    '三国法': [
        '国际私法', '法律冲突', '准据法', '反致', '识别', '先决问题',
        '公共秩序保留', '法律规避', '涉外合同', '涉外侵权', '涉外婚姻',
        '涉外继承', '国际商事仲裁', '司法协助', '外国判决承认执行',
        '国际贸易术语', '提单', '共同海损', '海上保险', '租船合同',
        '涉外收养', '缔约国', '出口管制', '国际货物销售',
        '外国人', '外国法', '国际条约', '国际惯例'
    ],
    '理论法': [
        '法治理论', '法律原则', '法律规则', '法的效力', '法的适用',
        '立法程序', '法律解释', '法律冲突', '宪法权利', '基本权利',
        '法治国家', '依法治国', '党内法规', '实证主义', '非实证主义',
        '形式法治', '实质法治', '法律体系', '法律渊源', '法律位阶',
        '法律监督', '宪法监督', '违宪审查'
    ]
}

# 刑法特有的场景描述
CRIMINAL_SCENARIOS = [
    # 主体描述
    '甲为了', '甲因', '甲与乙', '甲伙同', '甲教唆', '甲指使',
    # 行为描述
    '杀死', '杀害', '致死', '致人死亡', '重伤', '轻伤',
    '殴打', '伤害', '强奸', '猥亵', '拐卖', '绑架',
    '盗窃', '抢劫', '抢夺', '诈骗', '敲诈',
    '贪污', '受贿', '行贿', '挪用',
    # 结果描述
    '死亡', '重伤', '轻伤', '财产损失', '经济损失',
    # 程序描述
    '公安机关', '抓获', '逮捕', '自首', '投案',
    '如实供述', '翻供', '串供',
    # 定罪量刑
    '构成', '不构成', '成立', '不成立',
    '应当认定', '应认定为', '定罪', '量刑'
]

def extract_questions_from_document(doc_path):
    """从文档中提取所有题目"""
    doc = Document(doc_path)
    questions = []
    current_question = None
    current_content = []
    question_pattern = re.compile(r'^(\d+)\.')
    in_question = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目的开始
        match = question_pattern.match(text)
        if match and not in_question:
            # 保存前一个题目
            if current_question is not None:
                questions.append({
                    'number': current_question,
                    'content': '\n'.join(current_content)
                })
            
            # 开始新题目
            current_question = int(match.group(1))
            current_content = [text]
            in_question = True
        elif in_question:
            # 继续收集当前题目内容
            if text.startswith('---'):
                # 题目结束
                questions.append({
                    'number': current_question,
                    'content': '\n'.join(current_content)
                })
                current_question = None
                current_content = []
                in_question = False
            else:
                current_content.append(text)
    
    # 保存最后一个题目
    if current_question is not None:
        questions.append({
            'number': current_question,
            'content': '\n'.join(current_content)
        })
    
    return questions

def is_criminal_law_question(content):
    """判断是否为刑法题目（加强版）"""
    score = 0
    
    # 1. 检查刑法特有场景（权重：50）
    for scenario in CRIMINAL_SCENARIOS:
        if scenario in content:
            score += 50
    
    # 2. 检查犯罪相关词汇（权重：30）
    crime_keywords = [
        '犯罪', '罪名', '定罪', '量刑', '刑罚', '刑事责任',
        '故意', '过失', '主观', '客观', '构成要件'
    ]
    for keyword in crime_keywords:
        if keyword in content:
            score += 30
    
    # 3. 检查具体罪名（权重：100）
    crime_names = [
        '杀人罪', '伤害罪', '强奸罪', '抢劫罪', '盗窃罪', '诈骗罪',
        '贪污罪', '受贿罪', '交通肇事罪', '危险驾驶罪'
    ]
    for crime in crime_names:
        if crime in content:
            score += 100
    
    # 4. 检查是否有"甲乙丙"式的案例描述
    if re.search(r'[甲乙丙丁][某]?(?:为了|因|与|伙同|教唆)', content):
        score += 40
    
    # 5. 排除明显的其他科目
    if any(word in content for word in ['合同', '物权', '婚姻', '继承', '民事']):
        score -= 50
    if any(word in content for word in ['公司', '股东', '破产', '保险', '劳动']):
        score -= 50
    if any(word in content for word in ['行政', '许可', '处罚', '复议']):
        score -= 50
    
    return score

def classify_question(question_content):
    """对单个题目进行分类（改进版）"""
    scores = {subject: 0 for subject in SUBJECTS}
    
    # 首先特别检查是否为刑法题目
    criminal_score = is_criminal_law_question(question_content)
    if criminal_score > 100:  # 如果刑法特征明显
        scores['刑法'] = criminal_score
    
    # 1. 法条引用（最高权重：100分）
    for subject, citations in LAW_CITATIONS.items():
        for citation in citations:
            if citation in question_content:
                scores[subject] += 100
    
    # 2. 专业术语（高权重：10分）
    for subject, terms in PROFESSIONAL_TERMS.items():
        for term in terms:
            if term in question_content:
                scores[subject] += 10
    
    # 3. 特殊规则调整
    # 如果同时涉及刑法和诉讼，根据重点判断
    if scores['刑法'] > 0 and scores['刑事诉讼法'] > 0:
        if any(word in question_content for word in ['定罪', '量刑', '罪名', '构成何罪']):
            scores['刑法'] += 100
        elif any(word in question_content for word in ['侦查', '起诉', '审判', '辩护']):
            scores['刑事诉讼法'] += 100
    
    # 三国法特殊处理
    if any(keyword in question_content for keyword in ['涉外', '国际', '外国人', '缔约国', 'CISG']):
        scores['三国法'] += 50
    
    # 返回得分最高的科目
    if max(scores.values()) == 0:
        # 如果都没有匹配，再次检查是否可能是刑法
        if criminal_score > 50:
            return '刑法'
        return '理论法'  # 默认分类
    
    return max(scores, key=scores.get)

def create_subject_document(subject, questions, output_dir):
    """为特定科目创建文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(f'2024年法考客观题 - {subject}', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加统计信息
    stats = doc.add_paragraph()
    stats.add_run(f'共 {len(questions)} 道题').bold = True
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(20)
    
    # 分隔线
    doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 添加每道题
    for question in questions:
        # 添加题目内容
        lines = question['content'].split('\n')
        for line in lines:
            doc.add_paragraph(line)
        
        # 添加分隔线
        separator = doc.add_paragraph('-' * 30)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(6)
        separator.paragraph_format.space_after = Pt(6)
    
    # 保存文档
    output_path = os.path.join(output_dir, f'{subject}.docx')
    doc.save(output_path)
    print(f"已保存 {subject} 科目，共 {len(questions)} 题到: {output_path}")

def main():
    """主函数"""
    # 文件路径
    input_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    output_dir = '/Users/acheng/Downloads/law-exam-assistant/docx'
    
    print("2024年法考题智能分类系统（改进版）")
    print("=" * 60)
    print("特别优化刑法题目识别")
    print("=" * 60)
    
    # 确保输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 提取所有题目
    print("\n提取题目...")
    questions = extract_questions_from_document(input_file)
    print(f"共提取 {len(questions)} 道题")
    
    # 对每道题进行分类
    print("\n开始分类...")
    classified_questions = defaultdict(list)
    
    for question in questions:
        subject = classify_question(question['content'])
        classified_questions[subject].append(question)
        
        # 如果是刑法题目，打印详细信息
        if subject == '刑法':
            print(f"题目 {question['number']:3d} -> {subject} ✓")
            # 打印题目前100个字符以供验证
            preview = question['content'][:100].replace('\n', ' ')
            print(f"    预览: {preview}...")
        else:
            print(f"题目 {question['number']:3d} -> {subject}")
    
    # 统计结果
    print("\n分类结果统计:")
    print("-" * 40)
    total = 0
    for subject in SUBJECTS:
        count = len(classified_questions[subject])
        total += count
        if subject == '刑法':
            print(f"{subject:10s}: {count:3d} 题 ⭐")  # 特别标记刑法
        else:
            print(f"{subject:10s}: {count:3d} 题")
    print("-" * 40)
    print(f"总计: {total} 题")
    
    # 删除旧文件
    print("\n删除旧文件...")
    for subject in SUBJECTS:
        old_file = os.path.join(output_dir, f'{subject}.docx')
        if os.path.exists(old_file):
            os.remove(old_file)
            print(f"已删除: {old_file}")
    
    # 为每个科目创建独立文档
    print("\n生成分科目文档...")
    for subject in SUBJECTS:
        if classified_questions[subject]:
            create_subject_document(subject, classified_questions[subject], output_dir)
    
    print("\n✓ 分类完成！")
    print(f"所有文档已保存到: {output_dir}")

if __name__ == "__main__":
    main()