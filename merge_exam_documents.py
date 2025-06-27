#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法考真题文档合并脚本
将扫描版原始文本中的题目题干补充到完整版文档中
"""

import re
import json
from docx import Document
from collections import OrderedDict


def read_scanned_text(file_path):
    """读取扫描版原始文本"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 解析题目
    questions = OrderedDict()
    lines = content.split('\n')
    
    current_question = None
    current_content = []
    
    i = 0
    while i < len(lines):
        # 去掉行号前缀
        line = re.sub(r'^\d+→\s*', '', lines[i]).strip()
        
        if not line:
            i += 1
            continue
            
        # 匹配题号（如：1.、2.、51.等）
        match = re.match(r'^(\d+)[.、]\s*(.+)', line)
        if match:
            # 保存上一题
            if current_question and current_content:
                questions[current_question] = '\n'.join(current_content)
            
            # 开始新题目
            current_question = int(match.group(1))
            current_content = [match.group(2)]
        elif current_question is not None:
            # 继续当前题目的内容
            current_content.append(line)
        
        i += 1
    
    # 保存最后一题
    if current_question and current_content:
        questions[current_question] = '\n'.join(current_content)
    
    return questions


def read_docx_content(file_path):
    """读取docx文档内容"""
    doc = Document(file_path)
    
    questions = OrderedDict()
    current_question = None
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 匹配题号
        match = re.match(r'^(\d+)[.、]\s*(.+)', text)
        if match:
            # 保存上一题
            if current_question and current_content:
                questions[current_question] = '\n'.join(current_content)
            
            # 开始新题目
            current_question = int(match.group(1))
            current_content = [text]
        elif current_question is not None:
            current_content.append(text)
    
    # 保存最后一题
    if current_question and current_content:
        questions[current_question] = '\n'.join(current_content)
    
    return questions


def analyze_missing_content(scanned_questions, docx_questions):
    """分析完整版文档中缺失的题目题干"""
    missing_stems = {}
    
    for q_num in docx_questions:
        docx_content = docx_questions[q_num]
        scanned_content = scanned_questions.get(q_num, "")
        
        # 检查是否缺少题干（通常题干会包含问号）
        if '?' not in docx_content and scanned_content and '?' in scanned_content:
            # 提取扫描版中的题干
            lines = scanned_content.split('\n')
            stem_lines = []
            for line in lines:
                stem_lines.append(line)
                if '?' in line:
                    break
            
            missing_stems[q_num] = '\n'.join(stem_lines)
    
    return missing_stems


def merge_documents(scanned_questions, docx_questions, missing_stems):
    """合并文档，补充缺失的题干"""
    merged_doc = Document()
    
    # 添加标题
    merged_doc.add_heading('2024年法律职业资格考试客观题真题（完整版）', 0)
    merged_doc.add_paragraph()
    
    # 按题号顺序处理
    for q_num in sorted(set(scanned_questions.keys()) | set(docx_questions.keys())):
        # 获取扫描版和完整版的内容
        scanned_content = scanned_questions.get(q_num, "")
        docx_content = docx_questions.get(q_num, "")
        
        # 如果该题需要补充题干
        if q_num in missing_stems:
            # 从扫描版提取完整题目
            merged_content = scanned_content
        else:
            # 优先使用完整版内容，如果没有则使用扫描版
            merged_content = docx_content if docx_content else scanned_content
        
        if merged_content:
            # 添加题目到文档
            para = merged_doc.add_paragraph()
            para.add_run(f"{q_num}. ").bold = True
            
            # 处理题目内容
            lines = merged_content.split('\n')
            for i, line in enumerate(lines):
                # 跳过题号行
                if i == 0 and re.match(r'^\d+[.、]', line):
                    line = re.sub(r'^\d+[.、]\s*', '', line)
                
                if line.strip():
                    if i > 0:
                        para.add_run('\n')
                    para.add_run(line)
            
            merged_doc.add_paragraph()  # 题目间空行
    
    return merged_doc


def main():
    """主函数"""
    # 文件路径
    scanned_file = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）_扫描版_raw_doc.txt'
    docx_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/2024年法考真题-补充完整版.docx'
    
    print("正在读取扫描版原始文本...")
    scanned_questions = read_scanned_text(scanned_file)
    print(f"从扫描版中提取了 {len(scanned_questions)} 道题目")
    
    print("\n正在读取完整版文档...")
    try:
        docx_questions = read_docx_content(docx_file)
        print(f"从完整版中读取了 {len(docx_questions)} 道题目")
    except Exception as e:
        print(f"读取docx文件出错: {e}")
        docx_questions = OrderedDict()
    
    print("\n分析缺失的题目题干...")
    missing_stems = analyze_missing_content(scanned_questions, docx_questions)
    print(f"发现 {len(missing_stems)} 道题目缺少题干")
    
    if missing_stems:
        print("\n缺少题干的题目编号：")
        for q_num in sorted(missing_stems.keys()):
            print(f"  第{q_num}题")
    
    print("\n正在合并文档...")
    merged_doc = merge_documents(scanned_questions, docx_questions, missing_stems)
    
    print(f"\n正在保存到: {output_file}")
    merged_doc.save(output_file)
    
    # 生成分析报告
    report_file = '/Users/acheng/Downloads/law-exam-assistant/merge_report.json'
    report = {
        'scanned_questions_count': len(scanned_questions),
        'docx_questions_count': len(docx_questions),
        'missing_stems_count': len(missing_stems),
        'missing_stems_numbers': sorted(missing_stems.keys()),
        'total_merged_questions': len(set(scanned_questions.keys()) | set(docx_questions.keys()))
    }
    
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    print(f"\n合并完成！分析报告已保存到: {report_file}")


if __name__ == '__main__':
    main()