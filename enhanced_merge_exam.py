#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版 2024年法考真题和答案解析合并工具
专门针对法考题目格式进行优化，确保题目和答案正确合并
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def extract_questions_from_doc(doc_path):
    """从文档中提取题目，使用多种策略确保完整提取"""
    doc = Document(doc_path)
    questions = {}
    current_question_num = None
    current_content = []
    
    # 多种题号匹配模式
    question_patterns = [
        re.compile(r'^(\d+)[．.、]\s*(.*)$'),        # 1．或1.或1、
        re.compile(r'^(\d+)\s+(.*)$'),              # 1 开头（数字后有空格）
        re.compile(r'^\s*(\d+)[．.、]\s*(.*)$'),    # 前面有空格的
        re.compile(r'^第(\d+)题[：:]\s*(.*)$'),      # 第1题：
        re.compile(r'^【(\d+)】\s*(.*)$'),          # 【1】
        re.compile(r'^题目(\d+)[：:]\s*(.*)$'),     # 题目1：
    ]
    
    print(f"开始解析题目文档：{doc_path}")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # 检查是否是新题目
        is_new_question = False
        for pattern in question_patterns:
            match = pattern.match(text)
            if match:
                # 保存前一题
                if current_question_num is not None and current_content:
                    questions[current_question_num] = '\n'.join(current_content)
                    print(f"题目 {current_question_num}: {len(current_content)} 段内容")
                
                # 开始新题目
                current_question_num = int(match.group(1))
                remaining_text = match.group(2) if len(match.groups()) > 1 else ""
                current_content = [text] if remaining_text else [text]
                is_new_question = True
                break
        
        if not is_new_question and current_question_num is not None:
            # 继续当前题目的内容
            current_content.append(text)
        elif not is_new_question:
            # 如果没有当前题目，检查是否可能是第一题
            if any(keyword in text for keyword in ['关于', '根据', '下列', '依据', 'A.', 'A、']):
                # 可能是题目内容，假设是第1题
                if 1 not in questions:
                    current_question_num = 1
                    current_content = [text]
                    print(f"推测第1题开始：{text[:50]}...")
    
    # 保存最后一题
    if current_question_num is not None and current_content:
        questions[current_question_num] = '\n'.join(current_content)
        print(f"题目 {current_question_num}: {len(current_content)} 段内容")
    
    print(f"从题目文档提取到 {len(questions)} 道题")
    return questions

def extract_answers_from_doc(doc_path):
    """从答案解析文档中提取答案"""
    doc = Document(doc_path)
    answers = {}
    current_question = None
    current_content = []
    
    # 答案题号匹配模式
    answer_patterns = [
        re.compile(r'^(\d+)[．.、]\s*正确答案[:：]\s*([A-Z])', re.IGNORECASE),
        re.compile(r'^(\d+)[．.、]\s*(.*)$'),
        re.compile(r'^\s*(\d+)[．.、]\s*(.*)$'),
    ]
    
    print(f"开始解析答案文档：{doc_path}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目的答案
        is_new_answer = False
        for pattern in answer_patterns:
            match = pattern.match(text)
            if match:
                # 保存前一题的答案
                if current_question is not None and current_content:
                    answers[current_question] = '\n'.join(current_content)
                
                # 开始新题目的答案
                current_question = int(match.group(1))
                current_content = [text]
                is_new_answer = True
                break
        
        if not is_new_answer and current_question is not None:
            current_content.append(text)
    
    # 保存最后一题的答案
    if current_question is not None and current_content:
        answers[current_question] = '\n'.join(current_content)
    
    print(f"从答案文档提取到 {len(answers)} 道题的答案解析")
    return answers

def format_question_content(content):
    """格式化题目内容，确保选项清晰显示"""
    lines = content.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 检查是否是选项（A. B. C. D. 或 A、B、C、D、）
        if re.match(r'^[A-D][．.、]\s*', line):
            formatted_lines.append(f"    {line}")  # 选项缩进
        else:
            formatted_lines.append(line)
    
    return '\n'.join(formatted_lines)

def add_question_with_answer(doc, num, question_text=None, answer_text=None):
    """添加一个完整的题目和答案"""
    # 题目标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"第 {num} 题")
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 128)  # 深蓝色
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(6)
    
    # 题目内容
    if question_text:
        formatted_question = format_question_content(question_text)
        question_para = doc.add_paragraph(formatted_question)
        question_para.paragraph_format.left_indent = Pt(0)
        question_para.paragraph_format.space_after = Pt(6)
    else:
        # 如果没有题目内容，从答案中尝试提取
        if answer_text:
            # 尝试从答案中提取题目信息
            answer_lines = answer_text.split('\n')
            question_hint = f"【题目内容请参考答案解析】"
            question_para = doc.add_paragraph(question_hint)
            question_para.paragraph_format.space_after = Pt(6)
        else:
            no_question = doc.add_paragraph("【题目内容暂缺】")
            no_question.paragraph_format.space_after = Pt(6)
    
    # 答案和解析
    if answer_text:
        # 答案解析标题
        answer_title = doc.add_paragraph()
        answer_run = answer_title.add_run("【答案解析】")
        answer_run.font.size = Pt(12)
        answer_run.bold = True
        answer_run.font.color.rgb = RGBColor(128, 0, 0)  # 深红色
        answer_title.paragraph_format.space_before = Pt(8)
        answer_title.paragraph_format.space_after = Pt(4)
        
        # 答案内容
        answer_lines = answer_text.split('\n')
        for line in answer_lines:
            if line.strip():
                answer_para = doc.add_paragraph(line.strip())
                answer_para.paragraph_format.left_indent = Pt(20)
                answer_para.paragraph_format.space_after = Pt(2)
    else:
        # 没有答案
        no_answer = doc.add_paragraph()
        no_answer_run = no_answer.add_run("【答案解析】暂无")
        no_answer_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
        no_answer.paragraph_format.space_before = Pt(8)
    
    # 分隔线
    separator = doc.add_paragraph("─" * 60)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    separator.paragraph_format.space_before = Pt(10)
    separator.paragraph_format.space_after = Pt(10)

def create_merged_document(questions, answers, output_path):
    """创建合并的文档"""
    doc = Document()
    
    # 文档标题
    title = doc.add_heading('2024年法考真题（含完整答案解析）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 文档说明
    description = doc.add_paragraph()
    description.add_run("本文档整合了2024年法考客观卷真题和详细答案解析，每道题目后直接跟随对应的答案和解析内容。")
    description.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    description.paragraph_format.space_after = Pt(20)
    
    # 统计信息
    all_numbers = sorted(set(questions.keys()) | set(answers.keys()))
    total_count = len(all_numbers)
    question_count = len(questions)
    answer_count = len(answers)
    
    stats = doc.add_paragraph()
    stats.add_run(f"共收录 {total_count} 道题目，其中 {question_count} 道有完整题目内容，{answer_count} 道有答案解析")
    stats.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats.paragraph_format.space_after = Pt(30)
    
    # 分隔符
    doc.add_paragraph("═" * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    print(f"\n开始生成合并文档...")
    print(f"总题数: {total_count}")
    print(f"有题目内容: {question_count}")
    print(f"有答案解析: {answer_count}")
    
    # 逐题添加
    for i, num in enumerate(all_numbers, 1):
        question_text = questions.get(num)
        answer_text = answers.get(num)
        
        if i % 10 == 0:
            print(f"处理进度: {i}/{total_count}")
        
        add_question_with_answer(doc, num, question_text, answer_text)
    
    # 保存文档
    doc.save(output_path)
    print(f"\n文档生成完成！")
    print(f"保存路径: {output_path}")
    
    # 统计信息
    only_questions = set(questions.keys()) - set(answers.keys())
    only_answers = set(answers.keys()) - set(questions.keys())
    
    if only_questions:
        print(f"\n仅有题目无答案: {sorted(only_questions)}")
    if only_answers:
        print(f"仅有答案无题目: {len(only_answers)} 道")

def main():
    """主函数"""
    # 文件路径
    questions_file = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）.docx'
    answers_file = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    
    # 检查文件
    if not os.path.exists(questions_file):
        print(f"错误：找不到题目文件 {questions_file}")
        return
    
    if not os.path.exists(answers_file):
        print(f"错误：找不到答案文件 {answers_file}")
        return
    
    print("=" * 60)
    print("2024年法考真题合并工具")
    print("=" * 60)
    
    # 提取题目和答案
    questions = extract_questions_from_doc(questions_file)
    answers = extract_answers_from_doc(answers_file)
    
    # 创建合并文档
    create_merged_document(questions, answers, output_file)
    
    print("\n处理完成！")

if __name__ == "__main__":
    main()