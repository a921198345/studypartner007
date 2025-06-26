#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析23年法考题文档的结构
"""

from docx import Document

def analyze_document(doc_path):
    """分析文档结构"""
    print(f"分析文档: {doc_path}")
    
    try:
        doc = Document(doc_path)
        print(f"总段落数: {len(doc.paragraphs)}")
        
        # 显示前30个非空段落
        print("\n前30个非空段落:")
        non_empty = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                non_empty.append((i, text))
                if len(non_empty) <= 30:
                    print(f"第{i}行: {text[:100]}{'...' if len(text) > 100 else ''}")
        
        print(f"\n总共找到 {len(non_empty)} 个非空段落")
        
        # 检查表格
        if doc.tables:
            print(f"\n文档包含 {len(doc.tables)} 个表格")
            for i, table in enumerate(doc.tables[:3]):
                print(f"\n表格{i+1}: {len(table.rows)}行 x {len(table.columns)}列")
        
    except Exception as e:
        print(f"分析文档时出错: {e}")

if __name__ == "__main__":
    analyze_document('/Users/acheng/Downloads/law-exam-assistant/23年法考题.docx')