import docx
import re
import sys

QUESTIONS_FILE = '1.2024年回忆版真题（客观卷）.docx'
ANSWERS_FILE = '3.2024年回忆版真题答案解析（客观卷）.docx'
OUTPUT_FILE = '24年法考题.docx'


def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        import pip
        pip.main(['install', package])
        globals()[package] = __import__(package)


def parse_questions(path: str):
    doc = docx.Document(path)
    questions = {}
    current_num = None
    buffer = []
    for para in doc.paragraphs:
        text = para.text
        match = re.match(r'^\s*(\d+)\s*[．\.]', text)
        if match:
            if current_num is not None:
                questions[current_num] = "".join(buffer).strip()
            current_num = int(match.group(1))
            buffer = [text + '\n']
        elif current_num is not None:
            buffer.append(text + '\n')
    if current_num is not None:
        questions[current_num] = "".join(buffer).strip()
    return questions


def parse_answers(path: str):
    doc = docx.Document(path)
    answers = {}
    current_num = None
    buffer = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = re.match(r'^\s*(\d+)\s*\.\s*正确答案\s*[:：]\s*[A-Z]', text)
        if match:
            if current_num is not None:
                answers[current_num] = "\n".join(buffer).strip()
            current_num = int(match.group(1))
            buffer = [text]
        elif current_num is not None:
            buffer.append(text)
    if current_num is not None:
        answers[current_num] = "\n".join(buffer).strip()
    return answers


def merge_docs():
    install_and_import('docx')
    qs = parse_questions(QUESTIONS_FILE)
    ans = parse_answers(ANSWERS_FILE)

    doc = docx.Document()
    doc.add_heading('2024年法考客观题（真题+答案解析）', 0)
    for num in sorted(qs.keys()):
        q_text = qs.get(num)
        a_text = ans.get(num)
        doc.add_paragraph(q_text)
        if a_text:
            a_lines = a_text.split('\n')
            p = doc.add_paragraph()
            p.add_run(a_lines[0]).bold = True
            if len(a_lines) > 1:
                p.add_run('\n' + '\n'.join(a_lines[1:]))
        else:
            doc.add_paragraph('【答案缺失】')
        doc.add_paragraph('-'*30)
    doc.save(OUTPUT_FILE)
    print('合并完成，输出文件:', OUTPUT_FILE)


if __name__ == '__main__':
    merge_docs() 