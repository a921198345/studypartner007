#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
手动合并助手 - 帮助你按正确格式合并题目
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_template_document():
    """创建一个模板文档，展示正确的格式"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('2024年法考客观题（真题+答案解析）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 示例题目1（根据你的图片）
    # 题干
    q1_stem = doc.add_paragraph()
    q1_stem.add_run("1.").bold = True
    q1_stem.add_run(" 基于以下选项分析，请选择正确答案。")
    
    # 选项
    doc.add_paragraph("【选项】").bold = True
    doc.add_paragraph("A. （错误选项 - 根据解析：错误选项...）")
    doc.add_paragraph("B. （正确选项 - 根据解析：正确选项...）")  
    doc.add_paragraph("C. （错误选项 - 根据解析：错误选项...）")
    doc.add_paragraph("D. （错误选项 - 根据解析：错误选项...）")
    
    # 答案
    doc.add_paragraph()
    answer_para = doc.add_paragraph()
    answer_run = answer_para.add_run("【正确答案】B")
    answer_run.bold = True
    answer_run.font.color.rgb = RGBColor(255, 0, 0)
    
    # 解析
    doc.add_paragraph()
    analysis_title = doc.add_paragraph()
    analysis_title.add_run("【答案解析】").bold = True
    
    doc.add_paragraph("选项分析：")
    doc.add_paragraph("A项：错误选项").paragraph_format.left_indent = Pt(20)
    doc.add_paragraph("B项：正确选项").paragraph_format.left_indent = Pt(20)
    doc.add_paragraph("C项：错误选项").paragraph_format.left_indent = Pt(20)
    doc.add_paragraph("D项：错误选项").paragraph_format.left_indent = Pt(20)
    
    # 分隔线
    separator = doc.add_paragraph("-" * 60)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    separator.paragraph_format.space_before = Pt(10)
    separator.paragraph_format.space_after = Pt(10)
    
    # 保存模板
    doc.save('/Users/acheng/Downloads/law-exam-assistant/24年法考题-模板.docx')
    print("已创建模板文档：24年法考题-模板.docx")
    print("请参考此模板格式，手动输入题目内容")

def extract_and_format_answers(answers_file, output_file):
    """提取答案并创建待填充的文档"""
    doc_in = Document(answers_file)
    doc_out = Document()
    
    # 标题
    title = doc_out.add_heading('2024年法考客观题（真题+答案解析）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc_out.add_paragraph()
    
    # 处理每道题
    current_num = None
    for para in doc_in.paragraphs:
        text = para.text.strip()
        
        # 检测题号
        import re
        match = re.match(r'^(\d+)[．.、]\s*正确答案[：:]\s*([A-Z])', text)
        if match:
            current_num = int(match.group(1))
            answer = match.group(2)
            
            # 题干占位
            q_para = doc_out.add_paragraph()
            q_para.add_run(f"{current_num}.").bold = True
            placeholder_run = q_para.add_run(" 【请补充题干内容】")
            placeholder_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # 选项占位
            doc_out.add_paragraph()
            option_para = doc_out.add_paragraph()
            option_run = option_para.add_run("【请补充选项】")
            option_run.font.color.rgb = RGBColor(255, 0, 0)
            doc_out.add_paragraph("A. ...")
            doc_out.add_paragraph("B. ...")
            doc_out.add_paragraph("C. ...")
            doc_out.add_paragraph("D. ...")
            
            doc_out.add_paragraph()
            
        # 复制答案和解析内容
        if text and current_num:
            doc_out.add_paragraph(text)
            
        # 添加分隔线
        if text.startswith('综上所述'):
            separator = doc_out.add_paragraph("-" * 30)
            separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            separator.paragraph_format.space_before = Pt(8)
            separator.paragraph_format.space_after = Pt(8)
    
    doc_out.save(output_file)
    print(f"已创建待补充文档：{output_file}")
    print("请手动补充标记为红色的题干和选项内容")

if __name__ == "__main__":
    print("手动合并助手")
    print("=" * 50)
    
    # 创建模板
    create_template_document()
    
    # 创建待填充文档
    answers_file = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-待补充.docx'
    
    extract_and_format_answers(answers_file, output_file)