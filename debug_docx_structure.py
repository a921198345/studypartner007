#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试docx文档结构，查看实际内容格式
"""

from docx import Document
import re

def analyze_document(doc_path, max_paragraphs=20):
    """分析文档结构"""
    print(f"\n分析文档: {doc_path}")
    print("=" * 80)
    
    doc = Document(doc_path)
    
    print(f"总段落数: {len(doc.paragraphs)}")
    print("\n前 {max_paragraphs} 段内容:")
    print("-" * 80)
    
    for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
        text = para.text
        if text.strip():
            print(f"段落 {i+1:2d}: '{text}'")
            print(f"        长度: {len(text)}")
            
            # 检查是否匹配题号
            patterns = [
                r'^(\d+)[．.、]\s*',
                r'^(\d+)\s+',
                r'^\s*(\d+)[．.、]\s*',
            ]
            
            for j, pattern in enumerate(patterns):
                match = re.match(pattern, text)
                if match:
                    print(f"        匹配模式 {j+1}: 题号 {match.group(1)}")
            
            print()
        else:
            print(f"段落 {i+1:2d}: [空段落]")
    
    print("\n查找所有可能的题号...")
    print("-" * 40)
    
    question_numbers = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 尝试各种题号模式
            patterns = [
                r'^(\d+)[．.、]\s*',
                r'^(\d+)\s+[^\d]',
                r'^\s*(\d+)[．.、]\s*',
                r'第(\d+)题',
                r'题(\d+)',
            ]
            
            for pattern in patterns:
                match = re.search(pattern, text)
                if match:
                    num = int(match.group(1))
                    if num <= 200:  # 合理的题号范围
                        question_numbers.append((num, text[:100]))
    
    # 去重并排序
    unique_questions = {}
    for num, text in question_numbers:
        if num not in unique_questions:
            unique_questions[num] = text
    
    print(f"发现的题号: {sorted(unique_questions.keys())}")
    print(f"题号数量: {len(unique_questions)}")
    
    if unique_questions:
        print("\n题号示例:")
        for num in sorted(unique_questions.keys())[:10]:
            print(f"  {num}: {unique_questions[num]}")

def main():
    """主函数"""
    # 分析题目文档
    questions_file = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）.docx'
    analyze_document(questions_file, 30)
    
    # 分析答案文档
    answers_file = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    analyze_document(answers_file, 30)

if __name__ == "__main__":
    main()