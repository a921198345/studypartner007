#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简化版本：基于答案解析生成完整的法考试题文档
专注于答案文档处理，确保输出质量
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def extract_answers_with_details(doc_path):
    """提取答案解析，包含更多细节"""
    doc = Document(doc_path)
    answers = {}
    current_question = None
    current_answer = None
    current_content = []
    
    print(f"开始解析答案文档: {doc_path}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 匹配题目开始 "数字. 正确答案：X"
        answer_match = re.match(r'^(\d+)[．.、]\s*正确答案[：:]\s*([A-Z])', text)
        if answer_match:
            # 保存前一题
            if current_question is not None:
                answers[current_question] = {
                    'answer': current_answer,
                    'content': '\n'.join(current_content),
                    'raw_content': current_content[:]
                }
            
            # 开始新题目
            current_question = int(answer_match.group(1))
            current_answer = answer_match.group(2).upper()
            current_content = []
            
        elif current_question is not None:
            # 收集解析内容，跳过总结行
            if not (text.startswith('综上所述') and '本题答案' in text):
                if text != '【答案解析】':  # 跳过标题
                    current_content.append(text)
    
    # 保存最后一题
    if current_question is not None:
        answers[current_question] = {
            'answer': current_answer,
            'content': '\n'.join(current_content),
            'raw_content': current_content[:]
        }
    
    print(f"提取了 {len(answers)} 道题的答案解析")
    return answers

def extract_topic_from_analysis(content_lines):
    """从答案解析中推断题目类型和内容"""
    topic_hints = []
    
    for line in content_lines[:3]:  # 只看前3行
        # 寻找可能的题目线索
        if any(keyword in line for keyword in ['根据', '依据', '关于', '某', '下列', '以下']):
            topic_hints.append(line)
    
    return topic_hints

def create_comprehensive_document(answers, output_path):
    """创建全面的题目文档"""
    doc = Document()
    
    # 文档标题设置
    title = doc.add_heading('2024年法考客观题真题及答案解析', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 文档说明
    intro_para = doc.add_paragraph()
    intro_text = "本文档基于2024年法考官方答案解析整理，每道题包含正确答案和详细解析。"
    intro_para.add_run(intro_text).italic = True
    intro_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    intro_para.paragraph_format.space_after = Pt(20)
    
    # 统计信息
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"共收录 {len(answers)} 道题目，每题均含官方答案解析")
    stats_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats_para.paragraph_format.space_after = Pt(30)
    
    # 主分隔线
    separator_main = doc.add_paragraph("═" * 80)
    separator_main.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    separator_main.paragraph_format.space_after = Pt(20)
    
    print(f"\n开始生成文档，共 {len(answers)} 题...")
    
    # 按题号顺序处理
    for i, num in enumerate(sorted(answers.keys()), 1):
        data = answers[num]
        
        if i % 10 == 0:
            print(f"处理进度: {i}/{len(answers)}")
        
        # 题目编号和答案
        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_before = Pt(16)
        
        # 题号（蓝色加粗）
        title_run = header_para.add_run(f"第 {num} 题　")
        title_run.font.size = Pt(14)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0, 102, 204)
        
        # 正确答案（红色加粗）
        answer_run = header_para.add_run(f"正确答案：{data['answer']}")
        answer_run.font.size = Pt(12)
        answer_run.bold = True
        answer_run.font.color.rgb = RGBColor(192, 0, 0)
        
        # 题目内容提示（基于解析推断）
        topic_hints = extract_topic_from_analysis(data['raw_content'])
        if topic_hints:
            hint_para = doc.add_paragraph()
            hint_para.add_run("【题目线索】").bold = True
            hint_para.paragraph_format.space_before = Pt(6)
            hint_para.paragraph_format.space_after = Pt(4)
            
            for hint in topic_hints[:2]:  # 最多显示2个线索
                hint_content = doc.add_paragraph()
                hint_content.add_run(f"• {hint}")
                hint_content.paragraph_format.left_indent = Pt(20)
                hint_content.paragraph_format.space_after = Pt(2)
        
        # 答案解析标题
        analysis_title = doc.add_paragraph()
        analysis_run = analysis_title.add_run("【答案解析】")
        analysis_run.font.size = Pt(12)
        analysis_run.bold = True
        analysis_run.font.color.rgb = RGBColor(0, 128, 0)
        analysis_title.paragraph_format.space_before = Pt(8)
        analysis_title.paragraph_format.space_after = Pt(6)
        
        # 解析内容格式化
        for line in data['raw_content']:
            line = line.strip()
            if not line:
                continue
            
            if re.match(r'^[A-D]\s*项?[：:]', line):
                # 选项分析（缩进显示）
                option_para = doc.add_paragraph()
                option_para.paragraph_format.left_indent = Pt(25)
                option_para.paragraph_format.space_after = Pt(4)
                
                # 选项标记加粗
                parts = re.split(r'([A-D]\s*项?[：:])', line, 1)
                if len(parts) >= 3:
                    option_run = option_para.add_run(parts[1])
                    option_run.bold = True
                    option_run.font.color.rgb = RGBColor(0, 102, 153)
                    option_para.add_run(parts[2])
                else:
                    option_para.add_run(line)
                    
            elif re.match(r'^[A-D]\s*选项', line):
                # 选项总结（缩进显示）
                summary_para = doc.add_paragraph()
                summary_para.paragraph_format.left_indent = Pt(25)
                summary_para.paragraph_format.space_after = Pt(4)
                summary_para.add_run(line)
                
            else:
                # 普通解析内容
                normal_para = doc.add_paragraph(line)
                normal_para.paragraph_format.space_after = Pt(4)
        
        # 题目分隔线
        separator = doc.add_paragraph("─" * 70)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(12)
        separator.paragraph_format.space_after = Pt(12)
    
    # 文档结尾
    end_separator = doc.add_paragraph("═" * 80)
    end_separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    end_separator.paragraph_format.space_before = Pt(20)
    
    end_text = doc.add_paragraph()
    end_text.add_run("— 全部试题结束 —").bold = True
    end_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    end_text.paragraph_format.space_before = Pt(10)
    
    # 保存文档
    doc.save(output_path)
    print(f"\n✓ 文档生成完成！")
    print(f"保存位置: {output_path}")
    
    # 答案统计
    answer_stats = {}
    for data in answers.values():
        answer = data['answer']
        answer_stats[answer] = answer_stats.get(answer, 0) + 1
    
    print(f"\n答案分布统计:")
    for answer in sorted(answer_stats.keys()):
        print(f"  {answer}: {answer_stats[answer]} 题")

def main():
    """主函数"""
    # 文件路径
    answers_file = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    
    print("2024年法考题完整版生成工具")
    print("=" * 50)
    
    if not os.path.exists(answers_file):
        print(f"错误：找不到答案文件 {answers_file}")
        return
    
    # 提取答案解析
    answers = extract_answers_with_details(answers_file)
    
    if not answers:
        print("错误：没有提取到答案数据")
        return
    
    # 生成完整文档
    create_comprehensive_document(answers, output_file)
    
    print(f"\n✓ 处理完成！共生成 {len(answers)} 道题目的完整文档")

if __name__ == "__main__":
    main()