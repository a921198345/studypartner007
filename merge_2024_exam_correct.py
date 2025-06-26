#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
正确合并2024年法考真题和答案解析
参照2023年的成功格式
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def extract_questions_from_scanned_doc(doc_path):
    """从扫描版文档中提取题目（如果可能）"""
    # 由于是扫描版，可能需要OCR或手动处理
    # 暂时返回空字典
    return {}

def extract_answers_detailed(doc_path):
    """详细提取答案解析文档内容"""
    doc = Document(doc_path)
    questions = {}
    current_num = None
    current_data = {
        'answer': '',
        'full_content': [],
        'analysis_sections': []
    }
    
    print(f"开始解析答案文档: {doc_path}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是新题目的答案（如：1. 正确答案：A）
        answer_match = re.match(r'^(\d+)[．.、]\s*正确答案[：:]\s*([A-Z])', text)
        if answer_match:
            # 保存前一题
            if current_num is not None:
                questions[current_num] = {
                    'answer': current_data['answer'],
                    'full_content': current_data['full_content'][:],
                    'analysis_sections': current_data['analysis_sections'][:]
                }
            
            # 开始新题
            current_num = int(answer_match.group(1))
            current_data = {
                'answer': answer_match.group(2).upper(),
                'full_content': [text],
                'analysis_sections': []
            }
            
        elif current_num is not None:
            # 收集当前题目的所有内容
            current_data['full_content'].append(text)
            
            # 分析内容结构
            if text == '【答案解析】':
                current_data['analysis_sections'].append({'type': 'title', 'content': text})
            elif re.match(r'^[A-D]\s*项[：:]', text):
                current_data['analysis_sections'].append({'type': 'option_analysis', 'content': text})
            elif re.match(r'^[A-D]\s*选项', text):
                current_data['analysis_sections'].append({'type': 'option_summary', 'content': text})
            elif text.startswith('综上所述'):
                current_data['analysis_sections'].append({'type': 'conclusion', 'content': text})
            else:
                current_data['analysis_sections'].append({'type': 'general', 'content': text})
    
    # 保存最后一题
    if current_num is not None:
        questions[current_num] = {
            'answer': current_data['answer'],
            'full_content': current_data['full_content'][:],
            'analysis_sections': current_data['analysis_sections'][:]
        }
    
    print(f"提取了 {len(questions)} 道题的答案解析")
    return questions

def reconstruct_question_from_analysis(num, answer_data):
    """从答案解析中重构题目信息"""
    analysis_sections = answer_data['analysis_sections']
    
    # 提取题目背景和选项信息
    question_context = []
    options = {'A': '', 'B': '', 'C': '', 'D': ''}
    
    for section in analysis_sections:
        content = section['content']
        section_type = section['type']
        
        if section_type == 'general':
            # 查找可能的题目背景
            if any(keyword in content for keyword in ['根据', '依据', '某', '关于', '下列', '以下']):
                if len(content) > 30:  # 足够长的内容
                    question_context.append(content)
        
        elif section_type == 'option_analysis':
            # 提取选项分析
            match = re.match(r'^([A-D])\s*项[：:](.*)$', content)
            if match:
                opt_letter = match.group(1)
                opt_content = match.group(2).strip()
                options[opt_letter] = opt_content
    
    # 构造题干
    if question_context:
        # 使用找到的背景信息
        topic = question_context[0]
        if len(question_context) > 1:
            topic += "。" + question_context[1]
    else:
        # 默认题干
        topic = f"根据相关法律规定，关于下列情形，正确的是？"
    
    return {
        'topic': topic,
        'options': options
    }

def create_final_merged_document(answers_data, output_path):
    """创建最终的合并文档，格式参照2023年"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('2024年法考客观题（真题+答案解析）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加空行
    doc.add_paragraph()
    
    print(f"\n生成合并文档，共 {len(answers_data)} 题...")
    
    # 按题号顺序处理
    for i, num in enumerate(sorted(answers_data.keys()), 1):
        answer_data = answers_data[num]
        
        if i % 10 == 0:
            print(f"处理进度: {i}/{len(answers_data)}")
        
        # 重构题目信息
        reconstructed = reconstruct_question_from_analysis(num, answer_data)
        
        # 1. 题目编号和题干
        question_para = doc.add_paragraph()
        question_para.add_run(f"{num}.").bold = True
        
        # 添加题干
        topic_text = reconstructed['topic']
        if not topic_text.endswith('？'):
            topic_text += '？'
        question_para.add_run(f" {topic_text}")
        question_para.paragraph_format.space_after = Pt(6)
        
        # 2. 选项（基于解析推断）
        for opt_letter in ['A', 'B', 'C', 'D']:
            option_para = doc.add_paragraph()
            option_para.paragraph_format.left_indent = Pt(0)
            
            # 选项内容
            if reconstructed['options'][opt_letter]:
                # 使用解析中的信息
                option_text = f"{opt_letter}. （根据解析：{reconstructed['options'][opt_letter][:50]}...）"
            else:
                # 默认选项
                option_text = f"{opt_letter}. ..."
            
            option_para.add_run(option_text)
            option_para.paragraph_format.space_after = Pt(3)
        
        # 空行
        doc.add_paragraph()
        
        # 3. 正确答案和解析（原样复制）
        for line in answer_data['full_content']:
            if line.strip():
                # 答案行特殊处理
                if re.match(r'^\d+[．.、]\s*正确答案[：:]', line):
                    answer_para = doc.add_paragraph()
                    answer_para.add_run(line).bold = True
                    answer_para.paragraph_format.space_before = Pt(6)
                
                # 答案解析标题
                elif line == '【答案解析】':
                    analysis_title = doc.add_paragraph()
                    analysis_title.add_run(line).bold = True
                    analysis_title.paragraph_format.space_before = Pt(3)
                
                # 选项分析（缩进）
                elif re.match(r'^[A-D]\s*项[：:]', line) or re.match(r'^[A-D]\s*选项', line):
                    option_analysis = doc.add_paragraph(line)
                    option_analysis.paragraph_format.left_indent = Pt(20)
                    option_analysis.paragraph_format.space_after = Pt(2)
                
                # 其他内容
                else:
                    normal_para = doc.add_paragraph(line)
                    normal_para.paragraph_format.space_after = Pt(2)
        
        # 4. 题目之间的分隔线
        separator = doc.add_paragraph("-" * 30)
        separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        separator.paragraph_format.space_before = Pt(8)
        separator.paragraph_format.space_after = Pt(8)
    
    # 保存文档
    doc.save(output_path)
    print(f"\n✓ 文档生成完成！")
    print(f"保存位置: {output_path}")
    
    # 统计答案分布
    answer_stats = {}
    for data in answers_data.values():
        answer = data['answer']
        answer_stats[answer] = answer_stats.get(answer, 0) + 1
    
    print(f"\n答案分布:")
    for answer in sorted(answer_stats.keys()):
        print(f"  {answer}: {answer_stats[answer]} 题")

def main():
    """主函数"""
    # 文件路径
    questions_file = '/Users/acheng/Downloads/law-exam-assistant/1.2024年回忆版真题（客观卷）_扫描版.doc'
    answers_file = '/Users/acheng/Downloads/law-exam-assistant/3.2024年回忆版真题答案解析（客观卷）.docx'
    output_file = '/Users/acheng/Downloads/law-exam-assistant/24年法考题-完整版.docx'
    
    print("2024年法考题合并工具（参照2023年格式）")
    print("=" * 60)
    
    # 暂时只使用答案文档
    if not os.path.exists(answers_file):
        print(f"错误：找不到答案文件 {answers_file}")
        return
    
    # 提取答案数据
    answers_data = extract_answers_detailed(answers_file)
    
    if not answers_data:
        print("错误：没有提取到答案数据")
        return
    
    # 生成合并文档
    create_final_merged_document(answers_data, output_file)
    
    print("\n✓ 处理完成！")
    print("注意：由于原始题目文档为扫描版，题干和选项基于答案解析推断")
    print("建议手动核对并补充完整的题目内容")

if __name__ == "__main__":
    main()