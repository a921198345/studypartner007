#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
从Word文档提取文本的模块
用于法考助手AI问答知识库构建
"""

import os
import docx
import traceback
from docx.opc.exceptions import PackageNotFoundError
import subprocess
import tempfile

def extract_text_from_doc_with_textutil(doc_file_path):
    """
    使用macOS的textutil工具从.doc文件中提取文本
    
    参数:
        doc_file_path (str): .doc文件的路径
        
    返回:
        str: 提取的文本内容
    """
    try:
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp_file:
            tmp_path = tmp_file.name
        
        # 使用textutil转换doc到txt
        cmd = ['textutil', '-convert', 'txt', '-output', tmp_path, doc_file_path]
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode != 0:
            raise Exception(f"textutil转换失败: {result.stderr}")
        
        # 读取转换后的文本
        with open(tmp_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # 删除临时文件
        os.unlink(tmp_path)
        
        return text
    except Exception as e:
        print(f"使用textutil提取文本失败: {str(e)}")
        raise


def extract_text_from_word(word_file_path):
    """
    从Word文档中提取文本内容
    
    参数:
        word_file_path (str): Word文档的路径
        
    返回:
        str: 提取的文本内容
        
    异常:
        FileNotFoundError: 如果文件不存在
        ValueError: 如果文件格式不支持
        Exception: 其他错误
    """
    # 检查文件是否存在
    if not os.path.exists(word_file_path):
        raise FileNotFoundError(f"文件不存在: {word_file_path}")
    
    # 检查文件扩展名
    file_extension = os.path.splitext(word_file_path)[1].lower()
    if file_extension not in ['.docx', '.doc']:
        raise ValueError(f"不支持的文件格式: {file_extension}，仅支持.doc和.docx格式")
    
    # 检查文件的实际格式
    # 通过检查文件头来判断
    with open(word_file_path, 'rb') as f:
        header = f.read(8)
    
    # .docx文件是zip格式，以PK开头
    is_docx = header.startswith(b'PK')
    
    # 如果是真正的.docx文件，使用python-docx
    if is_docx:
        try:
            # 使用python-docx库打开文档
            doc = docx.Document(word_file_path)
            
            # 提取所有段落文本
            full_text = []
            
            # 提取段落
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text.append(paragraph.text)
            
            # 将所有文本段落合并成一个字符串，用换行符分隔
            return '\n'.join(full_text)
        
        except Exception as e:
            # 记录详细错误信息
            error_msg = f"提取文本时发生错误: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)  # 输出到控制台方便调试
            raise Exception(f"提取文本失败: {str(e)}")
    else:
        # 如果是.doc格式（或误标为.docx的.doc文件）
        print(f"检测到旧版Word文档格式，尝试使用textutil工具提取...")
        try:
            # 在macOS上使用textutil
            if os.name == 'posix' and subprocess.run(['which', 'textutil'], capture_output=True).returncode == 0:
                return extract_text_from_doc_with_textutil(word_file_path)
            else:
                raise ValueError("当前系统不支持处理旧版.doc格式文件，请转换为.docx格式后重试")
        except Exception as e:
            raise ValueError(f"无法处理旧版Word文档: {str(e)}")

if __name__ == "__main__":
    # 简单的测试代码
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            text = extract_text_from_word(file_path)
            print(f"成功提取文本，共 {len(text)} 个字符")
            print("文本预览（前500个字符）:")
            print(text[:500] + "...")
        except Exception as e:
            print(f"错误: {str(e)}")
    else:
        print("使用方法: python parse_docx.py <word文件路径>") 